import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = docx.Document()

# --- LAMPIRAN 2 ---
doc.add_heading("Lampiran 2 – Lembar Bimbingan Tesis", level=1)
p = doc.add_paragraph("Lembar bimbingan tesis ini merupakan dokumen resmi yang memuat rekam jejak konsultasi dan diskusi antara peneliti dengan Dosen Pembimbing selama proses penyusunan tesis. Bimbingan dilakukan secara berkala dan mencakup berbagai tahapan krusial, mulai dari perumusan masalah, eksplorasi data, implementasi algoritma (IndoBERT, GMM, dan LLM), hingga penyempurnaan penulisan laporan akhir Bab I sampai Bab V.")
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(12)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Hari/Tanggal'
hdr_cells[2].text = 'Materi / Catatan Bimbingan'
hdr_cells[3].text = 'Paraf Pembimbing'

# Set column widths (approximate)
for row in table.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(4.0)
    row.cells[2].width = Cm(8.0)
    row.cells[3].width = Cm(3.5)

# Add 8 empty rows for the student to fill
for i in range(1, 9):
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    for j in range(1, 4):
        row_cells[j].text = ""

doc.add_page_break()

# --- LAMPIRAN 1 ---
doc.add_heading("Lampiran 1 – Kode Python untuk olahdata", level=1)
doc.add_paragraph("Berikut adalah kumpulan kode sumber (source code) utama yang digunakan dalam penelitian ini, mencakup proses ekstraksi data, pemodelan IndoBERT, klasterisasi Gaussian Mixture Model (GMM), dan antarmuka dashboard berbasis Streamlit.")

src_dir = "DATASET/OLAH DATA/src"
files_to_include = ["pmb_pipeline.py", "app.py"]

for fname in files_to_include:
    fpath = os.path.join(src_dir, fname)
    if os.path.exists(fpath):
        doc.add_heading(f"File: {fname}", level=2)
        with open(fpath, 'r', encoding='utf-8') as f:
            code = f.read()
        
        # Add code block using a small monospaced font
        p_code = doc.add_paragraph(code)
        p_code.paragraph_format.space_after = Pt(0)
        p_code.paragraph_format.line_spacing = 1
        for run in p_code.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(8)

doc.save("FULL TESIS/Bahan_Lampiran_1_dan_2.docx")
print("Bahan_Lampiran_1_dan_2.docx created successfully!")
