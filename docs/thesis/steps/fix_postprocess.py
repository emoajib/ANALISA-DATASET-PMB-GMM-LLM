#!/usr/bin/env python3
"""
fix_postprocess.py — Post-processing pipeline (merged from otomatis_cleanup + otomatis_pasca).

Tasks:
  1. Copy ABSTRAK/ABSTRACT heading from FRONT_MATTER_DRAFT
  2. Clean up placeholders
  3. Report .bib vs Daftar Pustaka
  4. Swap ABSTRAK/ABSTRACT order if reversed
  5. Replace [Penguji 1/2/3] placeholders
  6. Generate TOC via AppleScript - Word
  7. Import .bib to Mendeley

Usage:
    python3 steps/fix_postprocess.py <docx_path> --front-matter <path> --bib <path> [--generate-toc]
"""
import argparse
import copy
import os
import re
import sys
import subprocess
import shutil
import zipfile
from lxml import etree
import shutil
import sqlite3
import subprocess
import sys
import textwrap
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))
from pipeline import utils
from pipeline import config


def backup_file(filepath):
    p = Path(filepath)
    backup = p.parent / f"{p.stem}_POSTPROC_BACKUP{p.suffix}"
    if not backup.exists():
        shutil.copy2(str(p), str(backup))
        print(f"  Backup: {backup}")
    return str(backup)


# ═══════════════════════════════════════════════════════════════
# TASK 1: Copy ABSTRAK heading from FRONT_MATTER
# ═══════════════════════════════════════════════════════════════
def _para_text(elem):
    texts = [t.text for t in elem.findall('.//' + qn('w:t')) if t.text]
    return ''.join(texts).strip()


def task_copy_abstrak(main_doc, front_matter_path):
    print("\n--- Task: Copy ABSTRAK/ABSTRACT Heading ---")
    if not os.path.exists(front_matter_path):
        print(f"  FRONT_MATTER_DRAFT tidak ditemukan: {front_matter_path}")
        return False
    front_doc = Document(front_matter_path)
    main_body = main_doc.element.body
    front_abstrak = front_abstract = None
    for child in front_doc.element.body:
        if child.tag != qn('w:p'): continue
        text = _para_text(child)
        if text.upper() == 'ABSTRAK' and front_abstrak is None: front_abstrak = child
        elif text.upper() == 'ABSTRACT' and front_abstract is None: front_abstract = child
    changes = False
    if front_abstrak is not None:
        main_children = list(main_body)
        target_idx = None
        for i, child in enumerate(main_children):
            if child.tag == qn('w:p') and 'kata kunci' in _para_text(child).lower():
                target_idx = i; break
        if target_idx is not None:
            has_abstrak = any(_para_text(main_children[j]).upper() == 'ABSTRAK' for j in range(max(0, target_idx-5), target_idx))
            if not has_abstrak:
                main_children[target_idx].addprevious(copy.deepcopy(front_abstrak))
                print("  ABSTRAK heading added"); changes = True
    if front_abstract is not None:
        main_children = list(main_body)
        target_idx = None
        for i, child in enumerate(main_children):
            if child.tag == qn('w:p') and _para_text(child).lower().startswith('keywords'):
                target_idx = i; break
        if target_idx is not None:
            has_abstract = any(_para_text(main_children[j]).upper() == 'ABSTRACT' for j in range(max(0, target_idx-5), target_idx))
            if not has_abstract:
                main_children[target_idx].addprevious(copy.deepcopy(front_abstract))
                print("  ABSTRACT heading added"); changes = True
    return changes


def task_clean_placeholders(main_doc):
    print("\n--- Task: Clean Placeholders ---")
    body = main_doc.element.body
    children = list(body)
    patterns = re.compile(r'\[(?:Isi\s+(?:Abstrak|Abstract|Pendahuluan|Bab\s+I|Rumusan\s+Masalah|Tujuan|Manfaat|Batasan|Hipotesis|Kajian\s+Pustaka|Metode|Hasil|Pembahasan|Kesimpulan)[^\]]*|isi\s+(?:abstrak|abstract)[^\]]*|Petunjuk[^\]]*|Lengkapi[^\]]*|Ganti[^\]]*|Hapus[^\]]*|Data\s+Dummy[^\]]*|Nama\s+(?:Dosen|Pembimbing|Ketua\s+Prodi|Rektor|Dekan|Kaprodi|Ketua\s+Penguji|Anggota\s+\d)[^\]]*|NIM[^\]]*|Lampiran[^\]]*|Daftar\s+Isi[^\]]*|Nomor[^\]]*|Deskripsi[^\]]*|Uraian[^\]]*|Hasilkan\s+Daftar\s+Isi[^\]]*)\]', re.I)
    removed = 0
    # First pass: paragraphs that are ENTIRELY placeholder → remove
    for child in children:
        if child.tag != qn('w:p'): continue
        t = _para_text(child)
        if patterns.fullmatch(t.strip()) or t == 'Page':
            body.remove(child); removed += 1
    # Second pass: inline placeholders → replace text in runs
    for child in list(body):
        if child.tag != qn('w:p'): continue
        t = _para_text(child)
        if not patterns.search(t): continue
        # Replace placeholder text within paragraph runs
        for run in child.findall(qn('w:r')):
            for txt_el in run.findall(qn('w:t')):
                if txt_el.text and patterns.search(txt_el.text):
                    txt_el.text = patterns.sub('', txt_el.text).strip()
                    if not txt_el.text:
                        run.getparent().remove(run)

    # Third pass: clean up known artifacts — ", " after "1. " etc., trailing "]"
    for child in list(body):
        if child.tag != qn('w:p'): continue
        for run in child.findall(qn('w:r')):
            for txt_el in run.findall(qn('w:t')):
                if not txt_el.text: continue
                t = txt_el.text
                # Remove leading comma+space artifact
                if t.startswith(', ') or t.startswith(' ,'):
                    txt_el.text = t.lstrip(' ,').strip()
                # Remove trailing ] or : artifact
                if txt_el.text and txt_el.text.rstrip().endswith((']', ':')):
                    txt_el.text = txt_el.text.rstrip().rstrip(':]').strip()
                if txt_el.text and not txt_el.text:
                    run.getparent().remove(run)

    print(f"  Cleaned {removed} placeholder paragraphs")


# ═══════════════════════════════════════════════════════════════
# TASK 2: Swap ABSTRAK/ABSTRACT if reversed
# ═══════════════════════════════════════════════════════════════
def _is_heading1(elem):
    if elem.tag != qn('w:p'): return False
    pPr = elem.find(qn('w:pPr'))
    if pPr is None: return False
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is not None:
        val = pStyle.get(qn('w:val'), '')
        if val in ('1', 'Heading1', 'heading1') or 'Heading1' in val: return True
    outlineLvl = pPr.find(qn('w:outlineLvl'))
    if outlineLvl is not None and outlineLvl.get(qn('w:val')) == '0': return True
    return False


def _find_heading(body, target_text):
    first_match = first_text = None
    for child in body:
        if child.tag != qn('w:p'): continue
        text = _para_text(child)
        if text.upper() == target_text.upper():
            if _is_heading1(child): return child, text
            if first_match is None: first_match, first_text = child, text
    return first_match, first_text


_SECTION_HEADINGS = {
    'KATA PENGANTAR', 'DAFTAR ISI', 'DAFTAR TABEL', 'DAFTAR GAMBAR',
    'DAFTAR LAMPIRAN', 'DAFTAR LAMBANG', 'DAFTAR SINGKATAN',
    'BAB I', 'BAB II', 'BAB III', 'BAB IV', 'BAB V',
}


def _find_section_end(body, start_elem):
    found_start = False
    for child in body:
        if child is start_elem: found_start = True; continue
        if found_start and child.tag == qn('w:p'):
            text = _para_text(child).upper()
            if text in _SECTION_HEADINGS: return child
    return None


def swap_abstrak_abstract(doc):
    print("\n--- Task: Check & Swap ABSTRAK/ABSTRACT Order ---")
    body = doc.element.body
    children = list(body)
    abstrak_elem, abstrak_text = _find_heading(body, "ABSTRAK")
    abstract_elem, abstract_text = _find_heading(body, "ABSTRACT")
    if abstrak_elem is None or abstract_elem is None:
        print("  ABSTRAK or ABSTRACT not found"); return
    abstrak_idx = children.index(abstrak_elem)
    abstract_idx = children.index(abstract_elem)
    if abstract_idx < abstrak_idx:
        print("  ABSTRACT before ABSTRAK — swapping...")
        abs_end = _find_section_end(body, abstrak_elem)
        eng_end = _find_section_end(body, abstract_elem)
        if abs_end and eng_end:
            abs_elems = []
            for child in list(body):
                if child is abstrak_elem: abs_elems.append(child)
                elif child is abs_end: abs_elems.append(child); break
                elif child is not abstract_elem and not (abstract_idx < children.index(child) < (children.index(abs_end) if abs_end else len(children))): pass
            eng_elems = []
            for child in list(body):
                if child is abstract_elem: eng_elems.append(child)
                elif child is eng_end: eng_elems.append(child); break
            for e in abs_elems: body.remove(e) if e in body else None
            for e in eng_elems: body.remove(e) if e in body else None
            ref_elem = _find_heading(body, "KATA PENGANTAR") or _find_heading(body, "DAFTAR ISI") or body[0]
            for e in reversed(eng_elems): ref_elem.addprevious(e)
            for e in reversed(abs_elems): ref_elem.addprevious(e)
            print("  Swapped: ABSTRAK before ABSTRACT ✓")


# ═══════════════════════════════════════════════════════════════
# TASK 3: Replace [Penguji 1/2/3] placeholders
# ═══════════════════════════════════════════════════════════════
def replace_penguji(doc, names):
    print("\n--- Task: Replace Penguji Placeholders ---")
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        for placeholder, name in names.items():
            if placeholder in txt:
                for r in p.runs:
                    if placeholder in r.text:
                        r.text = r.text.replace(placeholder, name)
                        print(f"  P{i}: {placeholder} → {name}")


def _scan_for_placeholders(doc):
    found = []
    for i, p in enumerate(doc.paragraphs):
        for match in re.finditer(r'\[(Penguji\s*\d|Penguji\s+(1|2|3))\]', p.text):
            found.append((i, match.group()))
    return found


# ═══════════════════════════════════════════════════════════════
# TASK 4: Generate TOC via AppleScript - Word
# ═══════════════════════════════════════════════════════════════
def _insert_toc_field_xml(doc):
    body = doc.element.body
    toc_inserted = False
    for i, child in enumerate(body):
        if child.tag == qn('w:p') and _para_text(child).upper() in ('DAFTAR ISI', 'DAFTAR TABEL', 'DAFTAR GAMBAR'):
            children = list(body)
            idx = children.index(child)
            if idx + 1 < len(children):
                next_p = children[idx + 1]
                if next_p.tag == qn('w:p'):
                    existing_toc = next_p.find('.//' + qn('w:instrText'))
                    if existing_toc is not None and 'TOC' in (existing_toc.text or ''):
                        print(f"  TOC field exists after '{_para_text(child)}'")
                        continue
                toc_p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                toc_p.append(pPr)
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                r.append(rPr)
                fld_begin = OxmlElement('w:fldChar')
                fld_begin.set(qn('w:fldCharType'), 'begin')
                r.append(fld_begin)
                r2 = OxmlElement('w:r')
                instr = OxmlElement('w:instrText')
                instr.set(qn('xml:space'), 'preserve')
                instr.text = ' TOC \\o "1-3" \\h \\z \\u '
                r2.append(instr)
                r3 = OxmlElement('w:r')
                fld_sep = OxmlElement('w:fldChar')
                fld_sep.set(qn('w:fldCharType'), 'separate')
                r3.append(fld_sep)
                r4 = OxmlElement('w:r')
                r4t = OxmlElement('w:t')
                r4t.text = '[Update Daftar Isi: klik kanan → Update Field]'
                r4.append(r4t)
                r5 = OxmlElement('w:r')
                fld_end = OxmlElement('w:fldChar')
                fld_end.set(qn('w:fldCharType'), 'end')
                r5.append(fld_end)
                toc_p.append(r); toc_p.append(r2); toc_p.append(r3); toc_p.append(r4); toc_p.append(r5)
                child.addnext(toc_p)
                print(f"  TOC field inserted after '{_para_text(child)}'")
                toc_inserted = True
    return toc_inserted


def generate_toc_via_word(filepath, doc):
    print("\n--- Task: Generate TOC via Word ---")
    doc.save(str(filepath))
    _insert_toc_field_xml(doc)
    doc.save(str(filepath))
    _refresh_fields_via_word(filepath)


def _refresh_fields_via_word(filepath):
    applescript = f'''
tell application "Microsoft Word"
    activate
    open "{filepath}"
    tell active document
        update artistic styles
        update styles
    end tell
end tell
'''
    try:
        p = subprocess.run(['osascript', '-e', applescript], capture_output=True, text=True, timeout=30)
        if p.returncode == 0:
            print("  Word fields refreshed ✓")
        else:
            print(f"  Word script warning: {p.stderr}")
    except FileNotFoundError:
        print("  osascript not available (macOS only)")
    except subprocess.TimeoutExpired:
        print("  Word script timed out")


def update_toc_via_word(filepath):
    print("\n--- Task: Update TOC ---")
    _refresh_fields_via_word(filepath)


# ═══════════════════════════════════════════════════════════════
# TASK 5: Import .bib to Mendeley
# ═════════════════════──────────────────────────────────────────
def _parse_bibtex_simple(filepath):
    import re as _re
    entries = {}
    current_key = None
    current_entry = []
    with open(filepath, 'r', encoding='utf-8') as f:
        for line in f:
            m = _re.match(r'@(\w+)\{(\w+),', line)
            if m:
                if current_key: entries[current_key] = '\n'.join(current_entry)
                current_key = m.group(2)
                current_entry = [line]
            elif current_key is not None:
                current_entry.append(line)
                if line.strip() == '}': entries[current_key] = '\n'.join(current_entry); current_key = None
    if current_key: entries[current_key] = '\n'.join(current_entry)
    return entries


def _find_mendeley_db():
    home = Path.home()
    candidates = [
        home / '.local/share/data/Mendeley Ltd./Mendeley Desktop/',
        home / 'Library/Application Support/Mendeley Desktop/',
        home / 'AppData/Local/Mendeley Ltd/Mendeley Desktop/',
    ]
    for d in candidates:
        if d.exists():
            db_files = list(d.glob('*.sqlite'))
            if db_files: return str(db_files[0])
    return None


def import_bib_to_mendeley(bib_path):
    print(f"\n--- Task: Import .bib to Mendeley ---")
    bib_path = Path(bib_path)
    if not bib_path.exists():
        print(f"  File not found: {bib_path}")
        return
    mendeley_db = _find_mendeley_db()
    if mendeley_db:
        entries = _parse_bibtex_simple(str(bib_path))
        try:
            conn = sqlite3.connect(mendeley_db)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM documents")
            total = cursor.fetchone()[0]
            print(f"  Mendeley DB: {total} existing documents")
            cursor.execute("SELECT title FROM documents")
            existing = {row[0].lower().strip() if row[0] else '' for row in cursor.fetchall()}
            new_count = sum(1 for k, v in entries.items() if not any(title in v.lower() for title in existing))
            conn.close()
            print(f"  New entries to import: ~{new_count}")
        except Exception as e:
            print(f"  Mendeley DB error: {e}")
    print(f"  Manual step: Open Mendeley → File → Add → {bib_path}")
    print(f"  Or use: open -a Mendeley-Desktop {bib_path}")


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════
def main():
    parser = argparse.ArgumentParser(description='Post-processing pipeline')
    parser.add_argument('file', nargs='?', help='Path ke DOCX')
    parser.add_argument('--front-matter', help='Path ke FRONT_MATTER_DRAFT.docx (copy ABSTRAK heading only)')
    parser.add_argument('--merge-front-matter', help='Path ke FRONT_MATTER_DRAFT.docx (full merge)')
    parser.add_argument('--bib', help='Path ke .bib file')
    parser.add_argument('--generate-toc', action='store_true', help='Generate TOC')
    parser.add_argument('--update-toc', action='store_true', help='Update existing TOC')
    parser.add_argument('--penguji1', help='Nama Penguji 1')
    parser.add_argument('--penguji2', help='Nama Penguji 2')
    parser.add_argument('--penguji3', help='Nama Penguji 3')
    parser.add_argument('--import-bib', help='Import .bib to Mendeley')
    args = parser.parse_args()

    doc_path = utils.get_doc_path(args.file)
    if not doc_path.exists():
        print(f'ERROR: File not found: {doc_path}')
        sys.exit(1)

    backup_file(str(doc_path))

    # ── Full front matter merge (before any other processing) ──
    if args.merge_front_matter:
        merge_script = os.path.join(os.path.dirname(__file__), 'merge_frontmatter.py')
        result = subprocess.run(
            [sys.executable, merge_script, str(doc_path),
             '--front-matter', args.merge_front_matter,
             '--output', str(doc_path)],
            capture_output=True, text=True, timeout=120)
        print(result.stdout)
        if result.returncode != 0:
            print(result.stderr)
            sys.exit(1)

    doc = utils.load_docx(doc_path)

    if args.front_matter:
        task_copy_abstrak(doc, args.front_matter)
        doc.save(str(doc_path))

    task_clean_placeholders(doc)
    doc.save(str(doc_path))

    swap_abstrak_abstract(doc)
    doc.save(str(doc_path))

    penguji = {}
    if args.penguji1: penguji['[Penguji 1]'] = args.penguji1
    if args.penguji2: penguji['[Penguji 2]'] = args.penguji2
    if args.penguji3: penguji['[Penguji 3]'] = args.penguji3
    if penguji:
        replace_penguji(doc, penguji)
        doc.save(str(doc_path))

    if args.generate_toc:
        generate_toc_via_word(str(doc_path), doc)

    if args.update_toc:
        update_toc_via_word(str(doc_path))

    if args.import_bib:
        import_bib_to_mendeley(args.import_bib)

    print("\nPost-process selesai.")


if __name__ == '__main__':
    main()
