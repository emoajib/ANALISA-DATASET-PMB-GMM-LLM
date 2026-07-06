#!/usr/bin/env python3
"""Fix tables: seragamkan border + gabung caption terpisah."""
import sys, os, re, shutil
from lxml import etree
import docx

NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

DOC = sys.argv[1] if len(sys.argv) > 1 else os.environ.get('TESIS_DOC', 'Tesis_ITSNU_v10_Final.docx')
BACKUP = DOC.replace('.docx', '_BEFORE_TABLEFIX.docx')

# Standard borders: top/bottom/insideH, single/4pt/auto (swiss style)
STD_BORDERS = ['top', 'bottom', 'insideH']

BORDER_PROPS = {
    'top': {'val': 'single', 'sz': '4', 'space': '0', 'color': 'auto'},
    'left': {'val': 'single', 'sz': '4', 'space': '0', 'color': 'auto'},
    'bottom': {'val': 'single', 'sz': '4', 'space': '0', 'color': 'auto'},
    'right': {'val': 'single', 'sz': '4', 'space': '0', 'color': 'auto'},
    'insideH': {'val': 'single', 'sz': '4', 'space': '0', 'color': 'auto'},
    'insideV': {'val': 'single', 'sz': '4', 'space': '0', 'color': 'auto'},
}


def make_border(name):
    """Create a w:border element."""
    b = etree.Element(f'{{{NS_W}}}{name}')
    for attr, val in BORDER_PROPS[name].items():
        b.set(f'{{{NS_W}}}{attr}', val)
    return b


def standardize_borders(tbl):
    """Set standard borders on a table."""
    tblPr = tbl.find(f'{{{NS_W}}}tblPr')
    if tblPr is None:
        tblPr = etree.SubElement(tbl, f'{{{NS_W}}}tblPr')
        tbl.insert(0, tblPr)
    
    # Remove existing borders
    existing = tblPr.find(f'{{{NS_W}}}tblBorders')
    if existing is not None:
        tblPr.remove(existing)
    
    # Create new borders
    borders = etree.SubElement(tblPr, f'{{{NS_W}}}tblBorders')
    for name in STD_BORDERS:
        borders.append(make_border(name))


def merge_caption_paragraphs(doc, num_para, title_para):
    """Merge two caption paragraphs into one."""
    # Get title text
    title_text = ''
    for r in doc.paragraphs[title_para].runs:
        title_text += r.text
    
    # Get number text
    num_text = ''
    for r in doc.paragraphs[num_para].runs:
        num_text += r.text
    
    # Merge: "Tabel X.Y Title"
    merged = f'{num_text} {title_text.strip()}'
    
    # Set merged text in the number paragraph
    # Clear existing runs first
    p_elem = doc.paragraphs[num_para]._element
    for r in p_elem.findall(f'{{{NS_W}}}r'):
        p_elem.remove(r)
    
    # Create new run with merged text
    rPr = etree.SubElement(etree.Element('dummy'), f'{{{NS_W}}}rPr')
    rFonts = etree.SubElement(rPr, f'{{{NS_W}}}rFonts')
    rFonts.set(f'{{{NS_W}}}ascii', 'Times New Roman')
    rFonts.set(f'{{{NS_W}}}hAnsi', 'Times New Roman')
    rFonts.set(f'{{{NS_W}}}cs', 'Times New Roman')
    sz = etree.SubElement(rPr, f'{{{NS_W}}}sz')
    sz.set(f'{{{NS_W}}}val', '20')  # 10pt
    szCs = etree.SubElement(rPr, f'{{{NS_W}}}szCs')
    szCs.set(f'{{{NS_W}}}val', '20')
    etree.SubElement(rPr, f'{{{NS_W}}}b')  # bold
    
    r = etree.SubElement(p_elem, f'{{{NS_W}}}r')
    r.append(rPr)
    t = etree.SubElement(r, f'{{{NS_W}}}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = merged
    
    # Clear the title paragraph
    for r in doc.paragraphs[title_para]._element.findall(f'{{{NS_W}}}r'):
        r.text = ''


def main():
    shutil.copy2(DOC, BACKUP)
    doc = docx.Document(DOC)
    
    # === M1: Standardize table borders ===
    for i, table in enumerate(doc.tables):
        standardize_borders(table._tbl)
    print(f'  ✓ Table borders standardized (top/bottom/insideH): {len(doc.tables)} tables')
    
    # === L3: Merge caption terpisah (dynamic detection) ===
    # Detect split captions: "Tabel X.Y" on one paragraph, title on the next
    caption_splits = []
    for i in range(len(doc.paragraphs) - 1):
        p1 = doc.paragraphs[i].text.strip()
        p2 = doc.paragraphs[i + 1].text.strip()
        if re.match(r'^Tabel\s+\d+\.\d+$', p1) and p2 and p2[0].isupper():
            caption_splits.append((i, i + 1))
    
    for num_para, title_para in caption_splits:
        try:
            merge_caption_paragraphs(doc, num_para, title_para)
        except Exception as e:
            print(f'  ✗ Merge failed for [{num_para}]+[{title_para}]: {e}')
    
    print(f'  ✓ Caption paragraphs merged: {len(caption_splits)}')
    
    doc.save(DOC)
    print(f'Saved: {DOC}')


if __name__ == '__main__':
    main()
