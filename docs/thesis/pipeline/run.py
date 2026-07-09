#!/usr/bin/env python3
"""
pipeline/run.py — Unified master pipeline for ITSNU thesis formatting.

Supports both targets:
  - Tesis_ITSNU_v10_Final.docx (full thesis, subprocess steps)
  - BAB I - BAB IV.docx (WIP, inline steps)

Usage:
    python3 pipeline/run.py Tesis_ITSNU_v10_Final.docx
    python3 pipeline/run.py BAB\ I\ -\ BAB\ IV.docx
    python3 pipeline/run.py  (uses default: BAB I - BAB IV.docx)
"""
import argparse
import os
import shutil
import subprocess
import sys
import re
import time
from datetime import datetime
from pathlib import Path
from lxml import etree

sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))
from pipeline import config
from pipeline import utils

SCRIPT_DIR = Path(__file__).parent.parent.resolve()


def step_subprocess(name, script, *args):
    print(f'\n{"="*60}')
    print(f'  STEP: {name}')
    print(f'  Script: {script}')
    print(f'{"="*60}')
    script_path = SCRIPT_DIR / 'steps' / script
    if not script_path.exists():
        print(f'  ERROR: {script_path} tidak ditemukan!')
        return False
    cmd = [sys.executable, str(script_path)] + list(args)
    result = subprocess.run(cmd, capture_output=False, text=True)
    return result.returncode == 0


def step_z_direct_xml(doc_path):
    print(f'\n{"="*60}')
    print(f'  STEP Z: Post-fix (margins + outline + spacing)')
    print(f'{"="*60}')
    try:
        import docx as dx_lib
        from docx.shared import Cm as DxCm
        z_doc = dx_lib.Document(str(doc_path))

        for si in range(len(z_doc.sections)):
            sec = z_doc.sections[si]
            sec.top_margin = DxCm(config.MARGIN_TOP)
            sec.bottom_margin = DxCm(config.MARGIN_BOTTOM)
            sec.left_margin = DxCm(config.MARGIN_LEFT)
            sec.right_margin = DxCm(config.MARGIN_RIGHT)
        print(f'  ✅ {len(z_doc.sections)} sections: margins {config.MARGIN_TOP}-{config.MARGIN_BOTTOM}-{config.MARGIN_LEFT}-{config.MARGIN_RIGHT} cm')

        h1_count = h2_count = h3_count = 0
        for p in z_doc.paragraphs:
            s = p.style.name if p.style else ''
            p_elem = p._element
            pPr = p_elem.find(f'{{{utils.NS_W}}}pPr')
            if pPr is None:
                pPr = etree.SubElement(p_elem, f'{{{utils.NS_W}}}pPr')
                p_elem.insert(0, pPr)
            expected_lvl = None
            if s == 'Heading 1': expected_lvl = '1'
            elif s == 'Heading 2': expected_lvl = '2'
            elif s == 'Heading 3': expected_lvl = '3'
            if expected_lvl:
                for ol in pPr.findall(f'{{{utils.NS_W}}}outlineLvl'):
                    pPr.remove(ol)
                ol = etree.SubElement(pPr, f'{{{utils.NS_W}}}outlineLvl')
                ol.set(f'{{{utils.NS_W}}}val', expected_lvl)
                if expected_lvl == '1': h1_count += 1
                elif expected_lvl == '2': h2_count += 1
                elif expected_lvl == '3': h3_count += 1
        print(f'  ✅ Outline levels: H1={h1_count} H2={h2_count} H3={h3_count}')

        bib_start = None
        for i, p in enumerate(z_doc.paragraphs):
            if p.text.strip() == 'DAFTAR PUSTAKA' or p.text.strip().startswith('DAFTAR PUSTAKA\n'):
                bib_start = i
                break
        bib_fixed = 0
        if bib_start:
            for j in range(bib_start + 1, len(z_doc.paragraphs)):
                pj = z_doc.paragraphs[j]
                if not pj.text.strip():
                    continue
                pj_elem = pj._element
                pPr = pj_elem.find(f'{{{utils.NS_W}}}pPr')
                if pPr is None:
                    pPr = etree.SubElement(pj_elem, f'{{{utils.NS_W}}}pPr')
                    pj_elem.insert(0, pPr)
                for sp in pPr.findall(f'{{{utils.NS_W}}}spacing'):
                    pPr.remove(sp)
                sp = etree.SubElement(pPr, f'{{{utils.NS_W}}}spacing')
                sp.set(f'{{{utils.NS_W}}}line', '240')
                sp.set(f'{{{utils.NS_W}}}lineRule', 'auto')
                sp.set(f'{{{utils.NS_W}}}after', '0')
                sp.set(f'{{{utils.NS_W}}}before', '0')
                bib_fixed += 1
        print(f'  ✅ {bib_fixed} bibliography entries: line=240')

        abs_removed = 0
        abs_seen = 0
        za_remove = []
        for i, p in enumerate(z_doc.paragraphs):
            t = p.text.strip()
            if t in ('ABSTRAK', 'ABSTRACT'):
                abs_seen += 1
                if abs_seen > 2:
                    za_remove.append(p._element)
                    abs_removed += 1
        for elem in za_remove:
            elem.getparent().remove(elem)
        if abs_removed:
            print(f'  ✅ {abs_removed} duplicate ABSTRAK/ABSTRACT removed')

        cap_fixed = 0
        for p in z_doc.paragraphs:
            txt = p.text.strip()
            if re.match(r'^(Tabel|Gambar)\s+\d+\.\d+\s+[A-Z]', txt) and len(txt) > 90:
                p_elem = p._element
                pPr = p_elem.find(f'{{{utils.NS_W}}}pPr')
                if pPr is None:
                    pPr = etree.SubElement(p_elem, f'{{{utils.NS_W}}}pPr')
                    p_elem.insert(0, pPr)
                for sp in pPr.findall(f'{{{utils.NS_W}}}spacing'):
                    pPr.remove(sp)
                sp = etree.SubElement(pPr, f'{{{utils.NS_W}}}spacing')
                sp.set(f'{{{utils.NS_W}}}line', '240')
                sp.set(f'{{{utils.NS_W}}}lineRule', 'auto')
                sp.set(f'{{{utils.NS_W}}}after', '240')
                sp.set(f'{{{utils.NS_W}}}before', '0')
                cap_fixed += 1
        print(f'  ✅ {cap_fixed} multi-line captions: line=240')

        z_doc.save(str(doc_path))
    except Exception as e:
        import traceback
        print(f'  ⚠  Step Z error: {e}')
        traceback.print_exc()


def run_pipeline_thesis(doc_path):
    """Full thesis pipeline (Tesis_ITSNU_vXX_Final.docx)."""
    backup_path = doc_path.parent / f"{doc_path.stem}_PIPELINE_BACKUP{doc_path.suffix}"
    shutil.copy2(str(doc_path), str(backup_path))
    print(f'Backup: {backup_path}')

    env = os.environ.copy()
    env['TESIS_DOC'] = str(doc_path)

    pipeline = [
        ('A. Bibliografi (clean + sort + hanging indent)', 'fix_bibliography.py'),
        ('B. Struktur (heading styles, outline, alignment)', 'fix_structure.py'),
        ('C. Tabel (borders, caption merge)', 'fix_tables.py'),
        ('D. Remaining (line spacing, italic, table font)', 'fix_remaining.py'),
        ('E. Postprocess (ABSTRAK, placeholders, TOC)', 'fix_postprocess.py',
         '--front-matter', str(config.FRONT_MATTER_PATH),
         '--bib', str(config.BIB_PATH)),
        ('F. Re-fix structure (post-TOC)', 'fix_structure.py'),
        ('G. Re-fix remaining (post-TOC)', 'fix_remaining.py'),
        ('H. Final compliance (margins, justify, headings)', 'fix_compliance.py'),
    ]

    for step_info in pipeline:
        name = step_info[0]
        script = step_info[1]
        extra_args = step_info[2:] if len(step_info) > 2 else []
        ok = step_subprocess(name, script, str(doc_path), *extra_args)
        if not ok:
            print(f'  ⚠  Step gagal: {name}')
            print(f'  Pipeline berhenti.')
            sys.exit(1)

    step_z_direct_xml(doc_path)

    print(f'\n{"="*60}')
    print(f'  COMPLIANCE CHECK')
    print(f'{"="*60}')
    subprocess.run([sys.executable, str(SCRIPT_DIR / 'check' / 'check_pedoman.py')])

    print(f'\n{"="*60}')
    print(f'  PIPELINE SELESAI')
    print(f'  Input: {doc_path}')
    print(f'  Backup: {backup_path}')
    print(f'{"="*60}')


def run_pipeline_inline(doc_path):
    """BAB I-IV inline pipeline."""
    from pipeline import config as cfg

    bak_path = doc_path.parent / f"{doc_path.stem}.bak"
    shutil.copy2(str(doc_path), str(bak_path))
    size_mb = os.path.getsize(str(bak_path)) / (1024 * 1024)
    print(f'Backup: {bak_path.name} ({size_mb:.1f} MB)')

    start_time = time.time()

    print('\n[1/7] GENERATE DAFTAR PUSTAKA')
    step_subprocess('Generate References', 'generate_references.py', str(doc_path))

    doc = utils.load_docx(doc_path)

    print('\n[2/7] FIX PARAGRAPHS (bold, spaces, indent, spacing, outline)')
    step_subprocess('Fix Paragraphs', 'fix_paragraphs.py', str(doc_path), '--step', 'all')

    print('\n[3/7] FIX REMAINING (italic, table font)')
    step_subprocess('Fix Remaining', 'fix_remaining.py', str(doc_path))

    print('\n[4/7] FIX MERGE PARAGRAPHS (BAB I)')
    step_subprocess('Merge Paragraphs', 'fix_merge_paragraphs.py', str(doc_path))

    print('\n[5/7] FIX PAKAR LABEL')
    step_subprocess('Fix Pakar', 'fix_pakar_label.py', str(doc_path))

    print('\n[6/7] FIX ITALIC FOREIGN')
    step_subprocess('Fix Italic', 'fix_italic_foreign.py', str(doc_path))

    doc.save(str(doc_path))

    print('\n[7/7] VALIDASI')
    doc_v = utils.load_docx(doc_path)
    dp_idx, _ = utils.find_daftar_pustaka(doc_v)
    if dp_idx is not None:
        ref_count = sum(1 for j in range(dp_idx + 1, len(doc_v.paragraphs)) if doc_v.paragraphs[j].text.strip())
        print(f'Daftar Pustaka: {ref_count} entries')
    h1 = sum(1 for p in doc_v.paragraphs if p.style and p.style.name == 'Heading 1')
    h2 = sum(1 for p in doc_v.paragraphs if p.style and p.style.name == 'Heading 2')
    h3 = sum(1 for p in doc_v.paragraphs if p.style and p.style.name == 'Heading 3')
    print(f'Headings: H1={h1} H2={h2} H3={h3}')

    elapsed = time.time() - start_time
    print(f'\nPIPELINE SELESAI ({elapsed:.1f}s)')


def main():
    parser = argparse.ArgumentParser(description='Unified pipeline Tesis ITSNU')
    parser.add_argument('file', nargs='?', help='Path ke DOCX')
    parser.add_argument('--mode', choices=['auto', 'thesis', 'bab1_4'], default='auto',
                        help='Pipeline mode (auto: detect from filename)')
    args = parser.parse_args()

    doc_path = utils.get_doc_path(args.file)
    if not doc_path.exists():
        print(f'ERROR: File tidak ditemukan: {doc_path}')
        sys.exit(1)

    print(f'Pipeline: {doc_path.name}')
    print(f'Mode: {args.mode}')

    if args.mode == 'thesis' or (args.mode == 'auto' and 'Tesis' in doc_path.name):
        run_pipeline_thesis(doc_path)
    else:
        run_pipeline_inline(doc_path)


if __name__ == '__main__':
    main()
