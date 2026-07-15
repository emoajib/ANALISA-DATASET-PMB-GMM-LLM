#!/usr/bin/env python3
"""
Sync draft docs to canonical Set-A values + fix K-Means-vs-GMM comparison GMM column.
Cell-scoped, schema-aware. Source of truth: src/outputs CSVs.
"""
import os, re
from docx import Document
from docx.oxml.ns import qn
import pandas as pd

ROOT = "/Volumes/WORK/MTI UNSIBANK/TESIS"
OUT = os.path.join(ROOT, "src/outputs")
MASTER = os.path.join(ROOT, "FULL TESIS/FULL_TESIS_FIXED_v2.docx")
BABIV = os.path.join(ROOT, "PENGERJAAN TESIS BAB I - BAB V/BAB I - BAB IV.docx")
BABV = os.path.join(ROOT, "PENGERJAAN TESIS BAB I - BAB V/BAB V.docx")
RABIT = os.path.join(ROOT, "PUBLIKASI/ARTIKEL_RABIT.docx")
RABITL = os.path.join(ROOT, "PUBLIKASI/Artikel_RABIT_LENGKAP.docx")

kscan = pd.read_csv(os.path.join(OUT, "tabel_4_5_kscan.csv"))
evalc = pd.read_csv(os.path.join(OUT, "tabel_4_7_evaluasi_internal.csv"))

kscan["BIC"] = kscan["BIC"].astype(float)
kscan["Sil"] = kscan["Sil"].astype(float)
kscan["LL"] = kscan["LL"].astype(float)
kcanon = {}
for y, g in kscan.groupby("Tahun"):
    r = g.loc[g["BIC"].idxmin()]
    kcanon[int(y)] = dict(K=int(r["K"]), BIC=float(r["BIC"]), Sil=float(r["Sil"]), LL=float(r["LL"]))
ecanon = {}
for _, r in evalc.iterrows():
    ecanon[int(r["Tahun"])] = dict(K=int(r["K"]), Sil=float(r["Silhouette"]),
                                     CH=float(r["Calinski-Harabasz"]), DB=float(r["Davies-Bouldin"]))

def f4(x): return f"{x:.4f}".replace('.', ',')
def f2(x): return f"{x:.2f}".replace('.', ',')
def f0(x): return str(int(round(x)))

def set_cell(cell, text):
    cell.text = ""
    cell.paragraphs[0].add_run(str(text))

def set_para_text(p, text):
    for r in p.runs:
        r._element.getparent().remove(r._element)
    p.add_run(text)

def rep(cell, old, new):
    t = cell.text
    if old in t:
        cell.text = t.replace(old, new)

def patch_kscan(doc, ti):
    tbl = doc.tables[ti]
    for r in tbl.rows[1:]:
        ym = re.match(r'(\d{4})', r.cells[0].text.strip())
        if not ym or int(ym.group(1)) not in kcanon:
            continue
        k = kcanon[int(ym.group(1))]
        set_cell(r.cells[2], k["K"])
        set_cell(r.cells[3], f0(k["BIC"]))
        set_cell(r.cells[4], f4(k["Sil"]))
        set_cell(r.cells[5], f2(k["LL"]))
    print(f"  patched kscan table {ti}")

def patch_eval(doc, ti):
    tbl = doc.tables[ti]
    for r in tbl.rows[1:]:
        ym = re.match(r'(\d{4})', r.cells[0].text.strip())
        if not ym or int(ym.group(1)) not in ecanon:
            continue
        e = ecanon[int(ym.group(1))]
        set_cell(r.cells[2], e["K"])
        set_cell(r.cells[3], f4(e["Sil"]))
        set_cell(r.cells[4], f2(e["CH"]))
        set_cell(r.cells[5], f4(e["DB"]))
    print(f"  patched eval table {ti}")

def patch_comparison(doc, ti, schema):
    tbl = doc.tables[ti]
    for r in tbl.rows[1:]:
        cells = r.cells
        metric = cells[0].text.strip()
        if schema == 'km_gmm_sel':
            gmm, sel = cells[2], cells[3]
            if metric.startswith("Silhouette"):
                rep(gmm, "0,0279", "0,0585"); rep(sel, "−0,0443", "−0,0137"); rep(sel, "-0,0443", "-0,0137")
            elif "Calinski" in metric:
                rep(gmm, "18,01", "27,19"); rep(sel, "−22,74", "−13,56"); rep(sel, "-22,74", "-13,56")
            elif "Davies" in metric:
                rep(gmm, "4,19", "3,1211"); rep(sel, "+1,22", "+0,1511"); rep(sel, "−1,22", "−0,1511")
        elif schema == 'gmm_km':
            gmm = cells[1]
            if "Silhouette" in metric and "2023" in metric: rep(gmm, "0,0905", "0,0691")
            elif "Silhouette" in metric and "2024" in metric: rep(gmm, "0,0279", "0,0585")
            elif "Calinski" in metric: rep(gmm, "18,01", "27,19")
            elif "Davies" in metric: rep(gmm, "4,19", "3,1211")
        elif schema == 'km_gmm_nosel':
            gmm = cells[2]
            if "Silhouette" in metric: rep(gmm, "0,0279", "0,0585")
            elif "Calinski" in metric: rep(gmm, "18,01", "27,19")
            elif "Davies" in metric: rep(gmm, "4,19", "3,1211")
    print(f"  patched comparison table {ti} ({schema})")

def fix_prose(doc):
    a, b, c = "0,0722 vs 0,0279", "40,75 vs 18,01", "2,97 vs 4,19"
    A, B, C = "0,0722 vs 0,0585", "40,75 vs 27,19", "2,97 vs 3,1211"
    n = 0
    for p in doc.paragraphs:
        t = p.text
        nt = t.replace(a, A).replace(b, B).replace(c, C)
        if nt != t:
            set_para_text(p, nt); n += 1
    if n: print(f"  fixed {n} comparison prose paragraphs")

def backup(path):
    import datetime
    bk = path.replace('.docx', f'_pre_cmpfix_backup_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    if not os.path.exists(bk):
        import shutil
        shutil.copy2(path, bk)
    return bk

def process(path, ops):
    print("###", os.path.basename(path))
    backup(path)
    doc = Document(path)
    for op in ops:
        op(doc)
    out = path
    doc.save(out)
    print("  saved", out)

process(MASTER, [lambda d: patch_comparison(d, 22, 'km_gmm_sel'), fix_prose])
process(BABIV, [lambda d: patch_kscan(d, 17),
                  lambda d: patch_comparison(d, 20, 'km_gmm_sel'),
                  fix_prose])
process(BABV, [])
process(RABIT, [lambda d: patch_kscan(d, 1),
                  lambda d: patch_comparison(d, 3, 'gmm_km'),
                  fix_prose])
process(RABITL, [lambda d: patch_eval(d, 2),
                   lambda d: patch_comparison(d, 3, 'km_gmm_nosel'),
                   fix_prose])
print("\nDONE")
