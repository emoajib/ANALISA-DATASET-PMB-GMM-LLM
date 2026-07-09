import docx
from lxml import etree
import shutil

NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def set_hanging_indent(pPr):
    ind = pPr.find(f'{{{NS_W}}}ind')
    if ind is None:
        ind = etree.SubElement(pPr, f'{{{NS_W}}}ind')
    if ind.get(f'{{{NS_W}}}first') is not None:
        del ind.attrib[f'{{{NS_W}}}first']
    ind.set(f'{{{NS_W}}}left', '720')
    ind.set(f'{{{NS_W}}}hanging', '720')

def set_1_spasi(pPr):
    sp = pPr.find(f'{{{NS_W}}}spacing')
    if sp is not None:
        pPr.remove(sp)
    sp = etree.SubElement(pPr, f'{{{NS_W}}}spacing')
    sp.set(f'{{{NS_W}}}line', '240')
    sp.set(f'{{{NS_W}}}lineRule', 'auto')
    sp.set(f'{{{NS_W}}}after', '0')
    sp.set(f'{{{NS_W}}}before', '0')

def main():
    doc_path = "FULL TESIS/FULL TESIS FINAL.docx"
    shutil.copy2(doc_path, doc_path + ".bak")
    doc = docx.Document(doc_path)
    
    # 1. Fix Empty Headings
    for p in doc.paragraphs:
        if p.style and 'Heading' in p.style.name:
            if not p.text.strip():
                p.style = 'Normal'
                
    # 2. Fix Abstrak Font Size
    in_abstrak = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in ('ABSTRAK', 'ABSTRACT'):
            in_abstrak = True
            continue
        if in_abstrak and p.style and 'Heading' in p.style.name and text:
            in_abstrak = False
            
        if in_abstrak and text:
            for r in p.runs:
                r.font.size = docx.shared.Pt(12)
                
    # 3. Fix Bibliography Spacing and Indent
    in_bib = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith('DAFTAR PUSTAKA'):
            in_bib = True
            continue
        if text.startswith('DAFTAR LAMPIRAN'):
            in_bib = False
            break
            
        if in_bib and text:
            p_elem = p._element
            pPr = p_elem.find(f'{{{NS_W}}}pPr')
            if pPr is None:
                pPr = etree.SubElement(p_elem, f'{{{NS_W}}}pPr')
                p_elem.insert(0, pPr)
            set_hanging_indent(pPr)
            set_1_spasi(pPr)
            
    doc.save(doc_path)
    print("Manual fixes applied successfully.")

if __name__ == '__main__':
    main()
