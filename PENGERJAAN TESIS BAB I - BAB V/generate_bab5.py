import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(run, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic

doc = docx.Document()

# Styles
for style in doc.styles:
    if hasattr(style, 'font'):
        style.font.name = 'Times New Roman'

p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p1.add_run("BAB V")
set_font(r1, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("KESIMPULAN DAN SARAN")
set_font(r2, bold=True)

# 5.1 Kesimpulan
p3 = doc.add_paragraph()
r3 = p3.add_run("5.1 Kesimpulan")
set_font(r3, bold=True)

# Intro to Kesimpulan
p_intro = doc.add_paragraph("Berdasarkan hasil analisis data segmentasi probabilistik menggunakan Gaussian Mixture Model (GMM) dan ekstraksi fitur semantik melalui IndoBERT pada populasi pendaftar mahasiswa baru ITSNU Pekalongan periode 2019-2024, penelitian ini berhasil mengisolasi karakteristik pendaftar secara multidimensional. Penggabungan otomasi penalaran (reasoning) berskala besar menggunakan Hybrid Cognitive Pipeline dengan model OpenRouter (meta-llama/llama-3.3-70b-instruct, nvidia/nemotron-3-ultra-550b-a55b, dan nvidia/nemotron-3-super-120b-a12b) juga telah menunjukkan kapabilitas AI generatif dalam mengubah data kuantitatif yang kompleks menjadi narasi manajerial yang dapat ditindaklanjuti. Mengacu pada rumusan masalah dan hipotesis yang diajukan, berikut adalah tiga kesimpulan utama dari penelitian ini:")
p_intro.paragraph_format.first_line_indent = Pt(36)
for r in p_intro.runs: set_font(r)

# Point 1
p4 = doc.add_paragraph("1. Efektivitas Integrasi Segmentasi Probabilistik dan Ekstraksi Semantik\nPenelitian ini secara empiris membuktikan bahwa integrasi antara representasi ruang vektor berdimensi 768 dari model bahasa besar berbahasa Indonesia (IndoBERT) dengan soft clustering probabilistik (GMM) sangat efektif dalam memetakan profil calon mahasiswa secara presisi. Dari dataset 2.362 pendaftar, pendekatan ini sukses mengidentifikasi K optimal yang bervariasi antara 2 hingga 6 per periode (masing-masing 6, 6, 6, 5, 2, 3 untuk tahun 2019-2024), dengan konsolidasi segmen yang semakin homogen dari 6 klaster pada 2019 menjadi 2 klaster pada 2023. Nilai Silhouette Score yang dihasilkan (misalnya 0,0905 pada tahun 2023) berada dalam rentang yang bermakna untuk model Gaussian Mixture berdasarkan Saqr & Lopez-Pernas (2024), yang mengonfirmasi bahwa tumpang tindih antar segmen dapat diukur secara eksak sebagai probabilitas posterior, bukan sekadar ketidakpastian. Hal ini membuktikan efektivitas metode GMM dalam memetakan probabilitas setiap pendaftar tanpa paksaan ke dalam satu kategori absolut, yang sebelumnya menjadi titik buta pada algoritma hard clustering seperti K-Means.")
for r in p4.runs: set_font(r)

# Point 2
p5 = doc.add_paragraph("2. Dinamika Perubahan Struktural Pendaftar (Structural Breaks)\nAnalisis stabilitas lintas waktu menggunakan metrik Adjusted Rand Index (ARI) membuktikan bahwa asumsi segmentasi pasar yang statis di lingkungan pendidikan tinggi adalah keliru. Seluruh 5 transisi antar periode (2019→2024) terkonfirmasi sebagai structural break dengan ARI < 0,30, di mana dua transisi menunjukkan disrupsi paling ekstrem dengan ARI negatif. Pada transisi Baseline ke COVID Crisis (2019→2020), nilai ARI tercatat sebesar -0,0036, yang mengindikasikan terjadinya perombakan mendasar dalam preferensi pendaftar akibat pandemi. Selanjutnya, pada fase pemulihan akhir (2022→2023), kembali terjadi guncangan struktural dengan nilai ARI sebesar -0,0036, di mana dominasi program studi vokasional dan teknologi informasi mengubah konfigurasi demografi pendaftar secara fundamental. Temuan bahwa setiap tahun terjadi perubahan struktural menggarisbawahi urgensi otomasi segmentasi tahunan, karena pola historis tahun sebelumnya belum tentu relevan untuk digunakan pada tahun ajaran berikutnya.")
for r in p5.runs: set_font(r)

# Point 3
p6 = doc.add_paragraph("3. Implikasi Praktis Hybrid Cognitive Pipeline untuk Kebijakan Rekrutmen\nPemanfaatan agen kecerdasan buatan (Hybrid Cognitive Pipeline) pada penelitian ini secara sukses menjembatani kesenjangan antara keluaran analisis data teknis dengan perumusan strategi rekrutmen. Model Ollama lokal memastikan keamanan privasi data (PII) sebelum metadata yang diagregasi diproses oleh model kognitif tingkat lanjut di cloud. Implikasi manajerial yang dihasilkan adalah tersedianya peta prioritas segmen dan rekomendasi strategi komunikasi secara instan bagi tim manajemen penerimaan mahasiswa baru. Hasil validasi sistem menunjukkan bahwa rekomendasi diferensiasi channel pemasaran (seperti Instagram/TikTok untuk segmen IT dan jalur bimbingan konseling untuk segmen lainnya) telah selaras secara akurat dengan pola perilaku klaster yang ditemukan di GMM, memastikan alokasi anggaran promosi institusi menjadi jauh lebih tepat sasaran.")
for r in p6.runs: set_font(r)

# 5.2 Saran
p7 = doc.add_paragraph()
r7 = p7.add_run("5.2 Saran")
set_font(r7, bold=True)

# Intro to Saran
p_saran = doc.add_paragraph("Berdasarkan rumusan kesimpulan serta proses validasi akhir terhadap arsitektur penelitian yang telah dibangun, peneliti mengajukan beberapa saran yang diklasifikasikan ke dalam ranah praktis untuk institusi dan ranah akademis untuk pengembangan keilmuan selanjutnya:")
p_saran.paragraph_format.first_line_indent = Pt(36)
for r in p_saran.runs: set_font(r)

# Practical
p8 = doc.add_paragraph("1. Saran Praktis (Bagi Manajemen ITSNU Pekalongan)\nTim Manajemen Penerimaan Mahasiswa Baru (PMB) ITSNU Pekalongan sangat direkomendasikan untuk mulai mengoperasionalkan proyeksi klaster tahun 2025 yang dihasilkan oleh pipeline ini sebagai basis penyusunan Rencana Kerja dan Anggaran (RKA) promosi. Mengingat klaster dengan karakteristik vokasi dan penguasaan teknologi digital terbukti memiliki momentum pertumbuhan terbesar pasca-pandemi, institusi perlu menyesuaikan bauran komunikasi pemasarannya (Kotler & Keller, 2016). Khususnya, alokasi sumber daya promosi di wilayah Kabupaten Pemalang dan Batang perlu ditingkatkan menggunakan pendekatan media sosial interaktif, yang menurut profil klaster GMM lebih diminati oleh lulusan SMK dan sekolah berbasis kejuruan di wilayah tersebut.")
for r in p8.runs: set_font(r)

# Academic
p9 = doc.add_paragraph("2. Saran Akademis (Bagi Penelitian Lanjutan)\nBagi akademisi maupun praktisi data yang berminat melanjutkan pengembangan arsitektur ini, peneliti menyarankan dua arah penyempurnaan utama. Pertama, eksplorasi penggunaan metode Dirichlet Process Mixture Models (DPMM) sebagai substitusi dari Gaussian Mixture Model (GMM). Berbeda dengan GMM yang membutuhkan pemindaian rentang nilai K secara heuristik (K-scan), DPMM adalah metode pemodelan Bayesian non-parametrik yang dapat menentukan jumlah klaster optimal secara inferensial dan otomatis sesuai dengan kompleksitas data (Scrucca et al., 2024). Kedua, penelitian lanjutan disarankan untuk melakukan evaluasi Fine-Tuning terhadap Small Language Models (SLM) seperti Llama-3-8B secara mandiri (on-premise). Hal ini bertujuan agar penalaran strategis dapat sepenuhnya berjalan tanpa ketergantungan pada API cloud eksternal seperti OpenRouter, sehingga memperkuat aspek skalabilitas dan kerahasiaan institusi jangka panjang.")
for r in p9.runs: set_font(r)


doc.save("BAB V.docx")
print("BAB V.docx created.")
