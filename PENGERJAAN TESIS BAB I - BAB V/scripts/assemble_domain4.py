import docx
from docxcompose.composer import Composer
import sys
import os
import glob

def delete_paragraphs_from(doc, start_idx):
    # Iterate in reverse to avoid index shifting
    for i in range(len(doc.paragraphs)-1, start_idx-1, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

def delete_paragraphs_until(doc, end_idx):
    for i in range(end_idx-1, -1, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

def assemble():
    print("Assembling FULL TESIS FINAL...")
    base_path = "BAB I - BAB IV.docx"
    bab5_path = "BAB V.docx"
    output_path = "../FULL TESIS/FULL TESIS FINAL.docx"
    
    # 1. Split BAB I - BAB IV into Base and Bib
    doc = docx.Document(base_path)
    bib_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "DAFTAR PUSTAKA" in p.text.upper():
            bib_idx = i
            break
            
    if bib_idx == -1:
        print("Error: Could not find DAFTAR PUSTAKA in BAB I - BAB IV.docx")
        return
        
    print(f"Found DAFTAR PUSTAKA at index {bib_idx}")
    
    # Create temp_base.docx
    doc_base = docx.Document(base_path)
    delete_paragraphs_from(doc_base, bib_idx)
    doc_base.save("temp_base.docx")
    
    # Create temp_bib.docx
    doc_bib = docx.Document(base_path)
    delete_paragraphs_until(doc_bib, bib_idx)
    # Also add a page break before DAFTAR PUSTAKA if needed (handled by pipeline)
    doc_bib.save("temp_bib.docx")
    
    # Create temp_lampiran.docx
    print("Generating LAMPIRAN from Python source codes...")
    doc_lampiran = docx.Document()
    doc_lampiran.add_heading("LAMPIRAN", level=1)
    doc_lampiran.add_paragraph("Lampiran 1 – Kumpulan Kode Pemrograman Python", style="Heading 2")
    
    src_dir = "../../DATASET/OLAH DATA/src"
    for py_file in glob.glob(os.path.join(src_dir, "*.py")):
        filename = os.path.basename(py_file)
        doc_lampiran.add_heading(f"Source Code: {filename}", level=3)
        try:
            with open(py_file, 'r', encoding='utf-8') as f:
                code = f.read()
            p = doc_lampiran.add_paragraph(code)
            # Use Courier New for code, size 9pt
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = docx.shared.Pt(9)
        except Exception as e:
            print(f"Skipping {filename}: {e}")
            
    doc_lampiran.save("temp_lampiran.docx")
    
    # Compose
    print("Composing documents...")
    master = docx.Document("temp_base.docx")
    composer = Composer(master)
    
    doc5 = docx.Document(bab5_path)
    composer.append(doc5)
    
    doc_b = docx.Document("temp_bib.docx")
    composer.append(doc_b)
    
    doc_l = docx.Document("temp_lampiran.docx")
    composer.append(doc_l)
    
    composer.save(output_path)
    print(f"Successfully assembled into {output_path}")
    
    # Cleanup
    os.remove("temp_base.docx")
    os.remove("temp_bib.docx")
    os.remove("temp_lampiran.docx")

if __name__ == "__main__":
    assemble()
