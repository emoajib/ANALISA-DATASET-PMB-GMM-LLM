#!/usr/bin/env python3
from docx import Document

def analyze_duplicate_bab5_issue():
    print("=== ANALYZING DUPLICATE BAB V ISSUE ===\n")
    
    # Load the current document
    doc = Document('Tesis_ITSNU_v11_Final.docx')
    
    print(f"Current document analysis:")
    print(f"  Total paragraphs: {len(doc.paragraphs)}")
    
    # Find all BAB V sections
    bab5_positions = []
    for i, para in enumerate(doc.paragraphs):
        if 'BAB V' in para.text:
            bab5_positions.append(i)
    
    print(f"  BAB V sections found at positions: {bab5_positions}")
    print(f"  Total BAB V sections: {len(bab5_positions)}")
    
    # Analyze each BAB V section
    print(f"\n=== Analyzing each BAB V section ===")
    
    for section_num, pos in enumerate(bab5_positions):
        print(f"\n--- BAB V Section {section_num + 1} (P{pos}) ---")
        
        # Show the first few paragraphs of this section
        section_content = []
        for i in range(pos, min(pos + 30, len(doc.paragraphs))):
            para_text = doc.paragraphs[i].text.strip()
            if para_text:
                section_content.append((i, para_text[:100]))
        
        # Look for key structural elements
        has_kesimpulan = any('KESIMPULAN' in text for _, text in section_content)
        has_51 = any('5.1 Simpulan' in text for _, text in section_content)
        has_52 = any('5.2 Implikasi' in text for _, text in section_content)
        has_53 = any('5.3 Keterbatasan' in text for _, text in section_content)
        has_54 = any('5.4 Saran' in text for _, text in section_content)
        has_daftar_pustaka = any('DAFTAR PUSTAKA' in text for _, text in section_content)
        
        print(f"  Key elements found:")
        print(f"    - KESIMPULAN: {'✓' if has_kesimpulan else '✗'}")
        print(f"    - 5.1 Simpulan: {'✓' if has_51 else '✗'}")
        print(f"    - 5.2 Implikasi: {'✓' if has_52 else '✗'}")
        print(f"    - 5.3 Keterbatasan: {'✓' if has_53 else '✗'}")
        print(f"    - 5.4 Saran: {'✓' if has_54 else '✗'}")
        print(f"    - DAFTAR PUSTAKA: {'✓' if has_daftar_pustaka else '✗'}")
        
        # Check for duplicate content patterns
        simpulan_count = sum(1 for _, text in section_content if 'KESIMPULAN' in text)
        print(f"  KESIMPULAN occurrences: {simpulan_count}")
        
        if section_num > 0:
            # Check if this is a duplicate of the first section
            first_section = bab5_positions[0]
            
            # Compare first few paragraphs
            first_match = True
            for i in range(min(10, len(section_content))):
                if section_num > 0:
                    # Get corresponding paragraphs from first section
                    first_para_text = doc.paragraphs[first_section + i].text.strip() if first_section + i < len(doc.paragraphs) else ""
                    current_para_text = section_content[i][1]
                    
                    if first_para_text and current_para_text:
                        # Compare normalized text
                        first_norm = first_para_text.lower().strip()
                        current_norm = current_para_text.lower().strip()
                        
                        if len(first_norm) > 10 and len(current_norm) > 10:
                            if first_norm in current_norm or current_norm in first_norm:
                                continue
                            else:
                                first_match = False
                                break
            
            if first_match:
                print(f"  ❌ This appears to be a DUPLICATION of section {first_bab5_section + 1}")
            else:
                print(f"  ? This section has different content but appears in wrong position")
    
    # Check document structure integrity
    print(f"\n=== Document Structure Analysis ===")
    
    # Expected structure
    expected_structure = [
        'BAB V',
        'KESIMPULAN DAN SARAN', 
        '5.1 Simpulan',
        '5.2 Implikasi',
        '5.3 Keterbatasan',
        '5.4 Saran',
        'DAFTAR PUSTAKA'
    ]
    
    # Find positions of key sections
    found_sections = {}
    for section in expected_structure:
        for i, para in enumerate(doc.paragraphs):
            if section in para.text:
                found_sections[section] = i
                break
    
    print(f"Found sections at positions:")
    for section, pos in sorted(found_sections.items()):
        print(f"  {section}: P{pos}")
    
    # Check if the structure is correct
    print(f"\n=== Structure Validation ===")
    
    if 'BAB V' in found_sections and '5.4 Saran' in found_sections:
        bab5_pos = found_sections['BAB V']
        saran_pos = found_sections['5.4 Saran']
        
        if bab5_pos < saran_pos:
            print(f"✓ BAB V (P{bab5_pos}) comes before 5.4 Saran (P{saran_pos})")
            
            # Check for other expected sections
            issues = []
            if 'DAFTAR PUSTAKA' not in found_sections:
                issues.append("DAFTAR PUSTAKA not found")
                
            if issues:
                print(f"❌ Structural issues:")
                for issue in issues:
                    print(f"   - {issue}")
            else:
                print(f"✅ Structure looks good!")
        else:
            print(f"❌ BAB V (P{bab5_pos}) comes after 5.4 Saran (P{saran_pos}) - structure incorrect")
    else:
        print(f"❌ Missing key sections in structure")
    
    return doc, bab5_positions, found_sections

if __name__ == "__main__":
    doc, bab5_positions, found_sections = analyze_duplicate_bab5_issue()
