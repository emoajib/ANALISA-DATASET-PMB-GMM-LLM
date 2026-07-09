import docx
import sys

def update_frontmatter(path):
    print("Updating Front Matter...")
    doc = docx.Document(path)
    
    # Text replacements
    replacements = {
        "[Nama Pembimbing]": "Dr. Drs. Eri Zuliarso, M.Kom., MCF",
        "[Nama Ketua Prodi]": "Dr. Eka Ardhianto, S.Kom., M.Cs., MTA.",
        "[Penguji 1]": "Dr. Kristiawan Nugroho, M.Kom.",
        "[Penguji 2]": "Dr. Kristiawan Nugroho, M.Kom.",  # Fallback if different label
        "[Penguji 3]": "Dr. Kristiawan Nugroho, M.Kom."
    }
    
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, val)
                
    # Find KATA PENGANTAR index
    kp_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "KATA PENGANTAR":
            kp_idx = i
            break
            
    if kp_idx != -1:
        # Clear existing KATA PENGANTAR content up to the next heading or empty lines
        for i in range(kp_idx + 1, min(kp_idx + 25, len(doc.paragraphs))):
            text_upper = doc.paragraphs[i].text.upper()
            if "DAFTAR" in text_upper or "ABSTRAK" in text_upper:
                break
            doc.paragraphs[i].text = ""

        # The new Kata Pengantar paragraphs
        paragraphs = [
            "",
            "Puji syukur ke hadirat Allah SWT atas segala limpahan rahmat, karunia, serta hidayah-Nya sehingga penulis dapat menyelesaikan tesis dengan judul “Strategi Segmentasi Probabilistik Calon Mahasiswa Menggunakan Gaussian Mixture Model Dan Otomasi Analisis Large Language Model Untuk Optimalisasi Rekrutmen Di Itsnu Pekalongan”. Tesis ini disusun untuk memenuhi salah satu syarat guna memperoleh gelar Magister pada Program Studi Magister Teknologi Informasi, Fakultas Teknologi Informasi dan Industri, Universitas Stikubank (UNISBANK).",
            "",
            "Dalam proses penyusunan tesis ini, penulis menyadari sepenuhnya bahwa terselesaikannya penelitian ini tidak terlepas dari bantuan, bimbingan, dukungan, dan doa dari berbagai pihak. Oleh karena itu, pada kesempatan ini penulis menyampaikan penghargaan dan terima kasih yang sebesar-besarnya kepada:",
            "",
            "1.\tDr. Drs. Eri Zuliarso, M.Kom., MCF, selaku Dekan Fakultas Fakultas Teknologi Informasi dan Industri Universitas Universitas Stikubank (UNISBANK) dan Dosen Pembimbing Tesis yang telah memberikan bimbingan, arahan, dan motivasi kepada penulis.",
            "2.\tDr. Eka Ardhianto, S.Kom., M.Cs., MTA., selaku Ketua Program Studi Magister Teknologi Informasi Universitas Universitas Stikubank (UNISBANK).",
            "3.\tDr. Kristiawan Nugroho, M.Kom., selaku Dosen Penguji Tesis yang telah memberikan masukan kepada penulis.",
            "4.\tSeluruh dosen Program Studi Magister Teknologi Informasi Universitas Universitas Stikubank (UNISBANK) yang telah membekali penulis dengan ilmu pengetahuan selama masa studi.",
            "5.\tKedua orang tua tercinta, keluarga besar, dan sahabat yang senantiasa memberikan doa, dorongan, kasih sayang, serta semangat yang tiada henti.",
            "6.\tSemua pihak yang telah membantu baik secara langsung maupun tidak langsung hingga tersusunnya tesis ini.",
            "",
            "Penulis menyadari bahwa dalam penyusunan tesis ini masih terdapat banyak kekurangan. Oleh karena itu, kritik dan saran yang membangun sangat penulis harapkan demi kesempurnaan karya ini. Semoga tesis ini dapat memberikan manfaat bagi pengembangan ilmu pengetahuan, khususnya dalam bidang analisis sentimen dan kecerdasan buatan.",
            "",
            "Semarang, 10 Juli 2026",
            ""
        ]
        
        # Insert new paragraphs
        p_idx = kp_idx + 1
        for text in paragraphs:
            if p_idx < len(doc.paragraphs):
                doc.paragraphs[p_idx].text = text
            else:
                doc.add_paragraph(text)
            p_idx += 1
            
    doc.save(path)
    print("Done updating Front Matter.")

if __name__ == "__main__":
    update_frontmatter(sys.argv[1])
