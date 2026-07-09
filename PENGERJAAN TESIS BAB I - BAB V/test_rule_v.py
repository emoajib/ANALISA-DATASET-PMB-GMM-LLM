import docx, re
from lxml import etree
doc = docx.Document('../FULL TESIS/FULL TESIS FINAL.docx')
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

bab_start = 0
bib_start = 0
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'BAB I': bab_start = i
    if p.text.strip().startswith('DAFTAR PUSTAKA'): bib_start = i

for i in range(bab_start, bib_start):
    p = doc.paragraphs[i]
    if p.style.name in ('Heading 1','Heading 2','Heading 3','List Paragraph'): continue
    txt = p.text.strip()
    if not txt or len(txt) < 30: continue
    if txt.startswith(('Sumber:','Tabel ','Gambar ','Keterangan:','Catatan:')): continue
    if p.style.name and 'toc' in p.style.name.lower(): continue
    if re.match(r'^(Tabel|Gambar)\s+\d+\.\d+\s+', txt): continue
    if re.match(r'^\(\w+\.\d+\)', txt): continue
    if re.match(r'^\d+\.\d+\s+[A-Z]', txt): continue

    pPr = p._element.find(f'{{{NS_W}}}pPr')
    if pPr is None: continue
    sp = pPr.find(f'{{{NS_W}}}spacing')
    after = sp.get(f'{{{NS_W}}}after') if sp is not None else None
    if after not in ('0', None):
        print(f"FAILED (after={after}): {txt[:50]}")
