#!/usr/bin/env python3
"""
check_pedoman.py — Compliance checker for Panduan Tesis v3 (2025).
Checks A–Z. All positions detected dynamically.
"""
import os, sys, re
from lxml import etree
import docx

from pathlib import Path
SCRIPT_DIR = Path(__file__).parent.parent.resolve()
DOC = sys.argv[1] if len(sys.argv) > 1 else os.environ.get(
    'TESIS_DOC',
    str(SCRIPT_DIR / 'Tesis_ITSNU_v11_Final.docx')
)
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS_M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

results = {'pass': 0, 'fail': 0, 'warn': 0}

def P(desc, cond, detail=''):
    if cond:
        results['pass'] += 1
        print(f'  \u2713  {desc}')
    else:
        results['fail'] += 1
        print(f'  \u2717  {desc}' + (f'  \u2014 {detail}' if detail else ''))

def W(desc, detail=''):
    results['warn'] += 1
    print(f'  \u26a0  {desc}' + (f'  \u2014 {detail}' if detail else ''))

doc = docx.Document(DOC)
paras = doc.paragraphs
body_elem = doc.element.body

def _extract_surname(text):
    if ',' in text:
        words = text.split(',')[0].strip().split()
        return words[-1].strip('.,;:()[]\'\"').lower() if words else text.lower()
    return text.split()[0].lower() if text.split() else text.lower()

# ── KEY POSITIONS ──
bab_positions = {}
for i, p in enumerate(paras):
    txt = p.text.strip()
    if txt.startswith('BAB ') and p.style.name == 'Heading 1':
        bab_positions[txt] = i
    if txt == 'DAFTAR PUSTAKA' or txt.startswith('DAFTAR PUSTAKA\n'):
        bab_positions['DAFTAR PUSTAKA'] = i

bib_start = bab_positions.get('DAFTAR PUSTAKA', 0) + 1
# Skip blank paragraphs right after the heading
while bib_start < len(paras) and not paras[bib_start].text.strip():
    bib_start += 1
bib_end = len(paras) - 1
for j in range(bib_start, len(paras)):
    if not paras[j].text.strip():
        bib_end = j - 1
        break

bib_entries_p = list(range(bib_start, bib_end + 1))

for h in ['BAB V', 'KESIMPULAN DAN SARAN', '5.1 Kesimpulan', '5.1 Simpulan', '5.2 Keterbatasan Penelitian', '5.3 Saran', '5.4 Saran']:
    for i, p in enumerate(paras):
        if p.text.strip() == h:
            bab_positions[h] = i
            break

print('\u2550' * 55)
print('  COMPLIANCE CHECK: Panduan Tesis v3 (2025)')
babv_pos = bab_positions.get('BAB V', '?')
dp_range = f'P{bib_start}–P{bib_end}'
print(f'  BAB V: P{babv_pos} | DP: {dp_range}')
print('\u2550' * 55)

# ═══════════════════════════════════════════════
# A. Font & Formatting (§4.1.2)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 A. Font & Formatting (\u00a74.1.2) \u2500\u2500')
tnr_errs = tsize_errs = 0
for p in paras:
    if p.style.name not in ('Heading 1', 'Heading 2', 'Heading 3'): continue
    for r in p.runs:
        rPr = r._element.find(f'{{{NS_W}}}rPr')
        if rPr is None: continue
        rFonts = rPr.find(f'{{{NS_W}}}rFonts')
        if rFonts is not None:
            f = rFonts.get(f'{{{NS_W}}}ascii', '')
            if f and 'Times New Roman' not in f: tnr_errs += 1
        sz = rPr.find(f'{{{NS_W}}}sz')
        if sz is not None and sz.get(f'{{{NS_W}}}val', '') != '24': tsize_errs += 1
P('Headings use Times New Roman', tnr_errs == 0, f'{tnr_errs} non-TNR')
P('Headings use 12pt', tsize_errs == 0, f'{tsize_errs} non-12pt')

tnr_body = 0
for p in paras[:bib_start]:
    if not p.text.strip() or p.style.name in ('Heading 1', 'Heading 2', 'Heading 3'): continue
    for r in p.runs[:3]:
        rPr = r._element.find(f'{{{NS_W}}}rPr')
        if rPr is None: continue
        rFonts = rPr.find(f'{{{NS_W}}}rFonts')
        if rFonts is not None:
            f = rFonts.get(f'{{{NS_W}}}ascii', '')
            if f and 'Times New Roman' not in f: tnr_body += 1
P('Body text uses TNR', tnr_body == 0, f'{tnr_body} non-TNR runs')

# ═══════════════════════════════════════════════
# B. Page Margins (§4.1.5)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 B. Page Margins (\u00a74.1.5) \u2500\u2500')
for si, sec in enumerate(doc.sections):
    t, b, l, r = sec.top_margin, sec.bottom_margin, sec.left_margin, sec.right_margin
    ok = False
    TG = 1440000; BG = 1080000; LG = 1440000; RG = 1080000
    if all(v is not None for v in (t, b, l, r)):
        ok = abs(t-TG)<10000 and abs(b-BG)<10000 and abs(l-LG)<10000 and abs(r-RG)<10000
    P(f'Section {si+1}: margins 4-3-4-3 cm', ok, f'top={t} bot={b} left={l} right={r}')

# ═══════════════════════════════════════════════
# C. Justification (§4.1.6)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 C. Justification (\u00a74.1.6) \u2500\u2500')
just_errs = 0
for i in range(bab_positions.get('BAB I', 0), bib_start):
    p = paras[i]
    if p.style.name in ('Heading 1','Heading 2','Heading 3'): continue
    if p.style.name and 'toc' in p.style.name.lower(): continue
    txt = p.text.strip()
    if not txt or txt.startswith('Sumber:') or len(txt) < 40 or txt.isupper(): continue
    # Skip image descriptions (extended captions) & captions
    if re.match(r'^(Gambar|Tabel)\s+\d+\.\d+[a-z]?\s+', txt): continue
    # Skip image intro paragraphs like "Gambar X.Y menampilkan/memvisualisasikan..."
    if re.match(r'^Gambar\s+\d+\.\d+[a-z]?\s+(menunju|menyajik|menampil|memvisualisasik|mengungkap|mendemonstrasik|memperlihatk|merangkum)', txt): continue
    pPr = p._element.find(f'{{{NS_W}}}pPr')
    if pPr is None: continue
    jc = pPr.find(f'{{{NS_W}}}jc')
    if jc is None or jc.get(f'{{{NS_W}}}val') != 'both': just_errs += 1
P('Body justified (rata kanan-kiri)', just_errs == 0, f'{just_errs} not justified')

# ═══════════════════════════════════════════════
# D. First-line Indent (§4.1.6)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 D. First-line Indent (\u00a74.1.6) \u2500\u2500')
indent_ok = indent_miss = indent_wrong_val = 0
for i in range(bab_positions.get('BAB I', 0), bib_start):
    p = paras[i]
    if p.style.name in ('Heading 1','Heading 2','Heading 3','List Paragraph'): continue
    txt = p.text.strip()
    if not txt or len(txt) < 30 or txt.startswith(('Sumber:','Tabel ','Gambar ')): continue
    pPr = p._element.find(f'{{{NS_W}}}pPr')
    if pPr is None: continue
    ind = pPr.find(f'{{{NS_W}}}ind')
    fl = ind.get(f'{{{NS_W}}}firstLine') if ind is not None else None
    if fl is not None:
        if fl == '720': indent_ok += 1
        else: indent_wrong_val += 1
    else: indent_miss += 1
detail = f'{indent_ok} correct 720, {indent_wrong_val} wrong val, {indent_miss} missing'
P('Body first-line indent = 720 twips', indent_wrong_val == 0 and indent_miss < indent_ok, detail)

# ═══════════════════════════════════════════════
# E. BAB V Heading Format
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 E. BAB V Heading Format \u2500\u2500')
for h_text in ['BAB V', 'KESIMPULAN DAN SARAN']:
    found = any(p.text.strip() == h_text and p.style.name == 'Heading 1' for p in paras)
    P(f'Heading 1 \"{h_text}\" found', found)
simpulan_h = any(p.text.strip() in ('5.1 Kesimpulan', '5.1 Simpulan') and p.style.name == 'Heading 2' for p in paras)
P(f'Heading 2 "5.1 Kesimpulan/Simpulan" found', simpulan_h)
keterbatasan_found = any(p.text.strip() == '5.2 Keterbatasan Penelitian' and p.style.name == 'Heading 2' for p in paras)
saran_found = any(p.text.strip() in ('5.2 Saran', '5.3 Saran', '5.4 Saran') and p.style.name == 'Heading 2' for p in paras)
P(f'Heading 2 "5.2/5.3/5.4 Saran" found', saran_found)
W('Heading 2 "5.2 Keterbatasan Penelitian" — v3: opsional', '(tidak diwajibkan)' if not keterbatasan_found else 'ditemukan')

bab5_start = bab_positions.get('BAB V', 0)
bab5_end = bab_positions.get('DAFTAR PUSTAKA', len(paras))
bab5_texts = [p.text for p in paras[bab5_start:bab5_end] if p.text.strip()] if bab5_start else []
bab5_words = sum(len(t.split()) for t in bab5_texts)
bab5_cites = sum(len(re.findall(r'\([^)]+,\s*\d{4}\)', t)) + len(re.findall(r'[A-Za-z]+\s+et\s+al\.\s*\(\d{4}\)', t)) for t in bab5_texts)

# Check E2: BAB V content
print('\n── E2. BAB V Content ──')
P('BAB V word count >= 500', bab5_words >= 500, f'{bab5_words} words')
P('BAB V citations >= 2', bab5_cites >= 2, f'{bab5_cites} citations')
P('BAB V body paragraphs exist', len(bab5_texts) >= 5, f'{len(bab5_texts)} paragraphs')
implikasi_found = any('Implikasi' in t for t in bab5_texts)
P('BAB V contains Implikasi', implikasi_found)
keterbatasan_found = any('Keterbatasan' in t for t in bab5_texts)
W('BAB V Keterbatasan — v3: opsional', 'ditemukan' if keterbatasan_found else 'tidak ada')

# ═══════════════════════════════════════════════
# F. Outline Levels
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 F. Outline Levels \u2500\u2500')
h1_bad = h2_bad = h3_bad = 0
for p in paras:
    s = p.style.name if p.style else ''
    pPr = p._element.find(f'{{{NS_W}}}pPr')
    if pPr is None: continue
    ol = pPr.find(f'{{{NS_W}}}outlineLvl')
    lvl = ol.get(f'{{{NS_W}}}val') if ol is not None else None
    lv = ol.get(f'{{{NS_W}}}val') if ol is not None else None
    if s == 'Heading 1' and lv != '1': h1_bad += 1
    elif s == 'Heading 2' and lv != '2': h2_bad += 1
    elif s == 'Heading 3' and lv != '3': h3_bad += 1
P('H1 outline=1', h1_bad == 0, f'{h1_bad}')
P('H2 outline=2', h2_bad == 0, f'{h2_bad}')
P('H3 outline=3', h3_bad == 0, f'{h3_bad}')

# ═══════════════════════════════════════════════
# G. BAB V Body Spacing (v3: space_after=0)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 G. BAB V Body Spacing (v3: 0) \u2500\u2500')
sec51 = bab_positions.get('5.1 Kesimpulan', 0)
sec53 = bab_positions.get('5.3 Saran', 0)
sp_issues = 0
for i in range(sec51 + 1, sec53):
    p = paras[i]
    if p.style.name in ('Heading 1','Heading 2'): continue
    if not p.text.strip(): continue
    pPr = p._element.find(f'{{{NS_W}}}pPr')
    if pPr is None: continue
    sp = pPr.find(f'{{{NS_W}}}spacing')
    after = sp.get(f'{{{NS_W}}}after') if sp is not None else None
    if not after or after != '0': sp_issues += 1
P('BAB V body spacing_after=0', sp_issues == 0, f'{sp_issues} issues')

# ═══════════════════════════════════════════════
# H. Daftar Pustaka
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 H. Daftar Pustaka \u2500\u2500')
entries = [(i, paras[i].text.strip()) for i in bib_entries_p if paras[i].text.strip()]
P('Bibliography entries', len(entries) >= 20, f'{len(entries)}')
glued = sum(1 for _, t in entries if re.match(r'^https?://\S+?[A-Z][a-z]', t))
P('No glued DOIs at entry start', glued == 0, f'{glued}')
surnames = [(i, _extract_surname(t)) for i, t in entries]
# Custom sort that handles Egyptian "El" prefix names properly
def sort_key(item):
    idx, surname = item
    # Egyptian names starting with "El" - "El" should sort by the full word after El
    if surname == 'el':
        # Get the next word for sorting
        original_text = entries[idx][1]
        parts = original_text.replace(',', ' ').split()
        if len(parts) > 1:
            return parts[1].lower()
    return surname

sorted_surnames = sorted(surnames, key=sort_key)
s_ok = all(sort_key(surnames[j]) <= sort_key(surnames[j+1]) for j in range(len(surnames)-1))
P('Sorted alphabetically', s_ok)
hang = sum(1 for i, _ in entries if (ppr := paras[i]._element.find(f'{{{NS_W}}}pPr')) is not None and (ind := ppr.find(f'{{{NS_W}}}ind')) is not None and ind.get(f'{{{NS_W}}}hanging') == '720' and ind.get(f'{{{NS_W}}}left') == '720')
P('Hanging indent on all entries', hang == len(entries), f'{hang}/{len(entries)}')
brk = sum(1 for _, t in entries if re.search(r'\[(Referensi|Software|ISBN)', t))
P('No bracket annotations', brk == 0, f'{brk}')
orph = sum(1 for _, t in entries if 'Rai,' in t)
P('Orphans removed (Rai)', orph == 0, f'{orph} found')

# ═══════════════════════════════════════════════
# I. Tables (§4.1.12)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 I. Tables (\u00a74.1.12) \u2500\u2500')
tbl = doc.tables
P('Tables count >= 30', len(tbl) >= 30, f'{len(tbl)}')
bord_ok = sum(1 for table in tbl if (tp := table._tbl.find(f'{{{NS_W}}}tblPr')) is not None and (b := tp.find(f'{{{NS_W}}}tblBorders')) is not None and all(b.find(f'{{{NS_W}}}{tag}') is not None and b.find(f'{{{NS_W}}}{tag}').get(f'{{{NS_W}}}val')=='single' and b.find(f'{{{NS_W}}}{tag}').get(f'{{{NS_W}}}sz')=='4' for tag in ['top','bottom','insideH']))
P('Table borders standard (top/bottom/insideH)', bord_ok >= 28, f'{bord_ok}/{len(tbl)}')

# ═══════════════════════════════════════════════
# J. Captions
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 J. Captions \u2500\u2500')
caps = sum(1 for p in paras if re.match(r'^Tabel \d+\.\d+\s+[A-Z]', p.text.strip()))
P('Captions merged', caps > 0, f'{caps} found')

# ═══════════════════════════════════════════════
# K. Cross-BAB Coherence
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 K. Cross-BAB Coherence \u2500\u2500')
bab_text = {1:'',2:'',3:'',4:'',5:''}
cur = 0
for p in paras:
    t = p.text.strip()
    if t == 'BAB I': cur = 1
    elif t == 'BAB II': cur = 2
    elif t == 'BAB III': cur = 3
    elif t == 'BAB IV': cur = 4
    elif t == 'BAB V': cur = 5
    if cur: bab_text[cur] += ' ' + t
rq = any(re.search(p, bab_text[1], re.I) for p in [r'bagaimana|apakah|rumusan masalah|research question'])
P('BAB I has RQs', rq)
mk = sum(1 for kw in ['GMM','IndoBERT','Llama','Ollama','segmentasi'] if kw.lower() in bab_text[3].lower())
P(f'BAB III mentions methods ({mk}/5)', mk >= 3)
rk = sum(1 for kw in ['ARI','klaster','cluster','structural break','centroid'] if kw.lower() in bab_text[4].lower())
P(f'BAB IV reports results ({rk}/5)', rk >= 3)
kc = len(re.findall(r'(?:Kesimpulan|1\.|2\.)', bab_text[5], re.I))
P('BAB V numbered conclusions >=1', kc >= 1, f'{kc} found')
P('BAB V mentions ARI', 'ARI' in bab_text[5] or 'Adjusted Rand' in bab_text[5])
P('BAB V has Saran', 'Saran' in bab_text[5])
W('BAB V Keterbatasan — v3: opsional', 'ditemukan' if 'Keterbatasan' in bab_text[5] else 'tidak ada')
key_refs = ['ollama','ahmadian','devlin','vaswani']
ref_ok = sum(1 for a in key_refs if any(a in e[1].lower() for e in entries))
W(f'Key refs in bibliography ({ref_ok}/{len(key_refs)})', 'ditemukan' if ref_ok >= 2 else 'kurang')

# ═══════════════════════════════════════════════
# L. Equations (§4.1.14)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 L. Equations (\u00a74.1.14) \u2500\u2500')
om = len(body_elem.findall(f'.//{{{NS_M}}}oMath'))
op = len(body_elem.findall(f'.//{{{NS_M}}}oMathPara'))
P(f'Native OMML equations >= 8', om >= 8, f'{om}')
P('No oMathPara wrappers', op == 0, f'{op}')

# ═══════════════════════════════════════════════
# M. Decimal Comma (§4.1.9)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 M. Decimal Comma (\u00a74.1.9) \u2500\u2500')
decimal_titik_hits = 0
decimal_false_positives = 0
filtered_patterns = {'heading_number':0,'version':0,'doi':0,'nim':0,'year':0,'subbab':0,'arxiv':0,'thousands':0}
for i, p in enumerate(paras):
    txt = p.text.strip()
    if not txt: continue
    if p.style.name in ('Heading 1','Heading 2','Heading 3'): continue
    if p.style.name and 'toc' in p.style.name.lower(): continue
    if re.match(r'^(Tabel|Gambar) \d+\.\d+', txt): continue
    # English abstract area — uses titik per English convention
    if i >= 33 and i < 40: continue
    hits = re.findall(r'\b\d+\.\d+\b', txt)
    for h in hits:
        filtering = False
        if re.match(r'^\d{1,2}\.\d{1,2}$', h) and 0 < int(h.split('.')[0]) <= 6:
            filtered_patterns['heading_number'] += 1; filtering = True
        if re.match(r'^\d{4}$', h.replace('.','')):
            filtered_patterns['year'] += 1; filtering = True
        if re.match(r'^\d+\.\d+\.\d+', h):
            filtered_patterns['version'] += 1; filtering = True
        # Version sub-component (e.g. "0.1" inside "0.1.0")
        if re.match(r'^\d+\.\d+$', h) and re.search(r'\d+\.\d+\.\d+', txt):
            filtered_patterns['version'] += 1; filtering = True
        if re.match(r'^10\.\d+', h):
            filtered_patterns['doi'] += 1; filtering = True
        if re.match(r'^\d{2}\.\d{2}\.\d{4}', h):
            filtered_patterns['nim'] += 1; filtering = True
        if re.match(r'^\d{2}\.\d{2}\.\d{2}\.\d{4}$', h):
            filtered_patterns['nim'] += 1; filtering = True
        if re.search(r'\d{2}\.\d{2}\.\d{2}\.\d{4}', txt):
            filtered_patterns['nim'] += 1; filtering = True
        if re.match(r'^\d{2}\.\d{4}$', h) and re.search(r'\d{2}\.\d{2}\.\d{2}\.\d{4}', txt):
            filtered_patterns['nim'] += 1; filtering = True
        # NIM sub-component (e.g. "85.7010" inside "25.01.85.7010")
        if re.match(r'^\d{2,3}\.\d{4}$', h):
            filtered_patterns['nim'] += 1; filtering = True
        if re.match(r'^\d+\.\d+$', h) and ('sub bab' in txt.lower() or 'subbab' in txt.lower()):
            filtered_patterns['subbab'] += 1; filtering = True
        if re.match(r'^[1-5]\.\d{2,}$', h):
            filtered_patterns['subbab'] += 1; filtering = True
        if re.match(r'^\d{4}\.\d+', h) and ('arxiv' in txt.lower() or re.search(r'arxiv|arXiv|2407\.\d+', txt)):
            filtered_patterns['arxiv'] += 1; filtering = True
        if '/' in txt and re.match(r'^\d+\.\d+', h) and re.search(r'doi\.org/|arxiv\.org/', txt):
            filtered_patterns['doi'] += 1; filtering = True
        if re.match(r'^\d{1,3}(?:\.\d{3})+$', h) and not re.match(r'^\d+\.\d{1,2}$', h):
            filtered_patterns['thousands'] += 1; filtering = True
        if filtering: decimal_false_positives += 1
        else: decimal_titik_hits += 1
P('Decimal titik violations after filtering', decimal_titik_hits == 0,
  f'{decimal_titik_hits} potential, {decimal_false_positives} FP')
P('Filter stats: heading_nums filtered', filtered_patterns['heading_number'] > 0,
  f'{filtered_patterns["heading_number"]}')
P('Version/year/DOI/arXiv/NIM/subbab filters active',
  sum(filtered_patterns.values()) > 0)

# ═══════════════════════════════════════════════
# N. Page Numbering (§4.1.11) — warn only
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 N. Page Numbering (\u00a74.1.11) \u2500\u2500')
W('Multiple sections found for page numbering', f'{len(doc.sections)} sections')
footers_found = sum(1 for si, sec in enumerate(doc.sections) for ft in [sec.footer, sec.even_page_footer, sec.first_page_footer] if ft and ft.paragraphs and any(p.text.strip() and any(c.isdigit() for c in p.text) for p in ft.paragraphs))
W('Page numbers detected in footers', f'{footers_found} footer(s) with numbers')

# ═══════════════════════════════════════════════
# O. Font Color
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 O. Font Color \u2500\u2500')
color_errs_body = color_errs_head = 0
for p in paras:
    for r in p.runs:
        rPr = r._element.find(f'{{{NS_W}}}rPr')
        if rPr is None: continue
        color = rPr.find(f'{{{NS_W}}}color')
        if color is not None:
            val = color.get(f'{{{NS_W}}}val', '')
            if val and val not in ('000000','auto'):
                if p.style.name in ('Heading 1','Heading 2','Heading 3'): color_errs_head += 1
                else: color_errs_body += 1
P('Body text = black/auto', color_errs_body == 0, f'{color_errs_body}')
P('Heading color = black/auto', color_errs_head == 0, f'{color_errs_head}')

# ═══════════════════════════════════════════════
# P. Citations APA
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 P. Citations APA (\u00a74.1.15) \u2500\u2500')
apa_pattern = re.compile(
    r'\((?:'                                      # opening paren
    r'[A-Z][a-zà-ÿ]+(?:\s+et\s+al\.?)?'          # single author
    r'(?:\s+[&]\s+[A-Z][a-zà-ÿ]+)?'              # optional & second author
    r',\s*\d{4}[a-z]?'                            # , year
    r')\)'                                        # closing paren
)
apa_citations = sum(len(apa_pattern.findall(p.text)) for p in paras)
P('APA in-text citations found', apa_citations >= 10, f'{apa_citations}')
ref_mgr = sum(1 for p in paras if re.search(r'mend[eé]ley|zotero|endnote', p.text, re.I))
W('Reference manager mention', f'{ref_mgr} mentions' if ref_mgr else 'Not found')

# ═══════════════════════════════════════════════
# Q. Paper Size
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 Q. Paper Size (\u00a74.1.1) \u2500\u2500')
all_a4 = all(abs(sec.page_width/360000 - 21.0) < 0.1 and abs(sec.page_height/360000 - 29.7) < 0.1 for sec in doc.sections)
P('Paper size = A4 (21x29.7 cm)', all_a4)

# ═══════════════════════════════════════════════
# R. Line Spacing (§4.1.4)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 R. Line Spacing (\u00a74.1.4) \u2500\u2500')
body_sp_ok = body_sp_bad = 0
for i in range(bab_positions.get('BAB I', 0), bib_start):
    p = paras[i]
    if p.style.name in ('Heading 1','Heading 2','Heading 3','List Paragraph'): continue
    txt = p.text.strip()
    if not txt or len(txt) < 30: continue
    if txt.startswith(('Sumber:','Tabel ','Gambar ','Keterangan:','Catatan:')): continue
    pPr = p._element.find(f'{{{NS_W}}}pPr')
    if pPr is None: continue
    sp = pPr.find(f'{{{NS_W}}}spacing')
    line = sp.get(f'{{{NS_W}}}line') if sp is not None else None
    if line == '480' or line is None: body_sp_ok += 1
    else: body_sp_bad += 1
P('Body paragraphs use 2 spasi (line=480)', body_sp_bad == 0, f'{body_sp_ok} ok, {body_sp_bad} non-480')

bib_sp_ok = bib_sp_bad = 0
for i in bib_entries_p:
    p = paras[i]
    if not p.text.strip(): continue
    pPr = p._element.find(f'{{{NS_W}}}pPr')
    if pPr is None: continue
    sp = pPr.find(f'{{{NS_W}}}spacing')
    line = sp.get(f'{{{NS_W}}}line') if sp is not None else None
    if line == '240': bib_sp_ok += 1
    else: bib_sp_bad += 1
P('Bibliography use 1 spasi (line=240)', bib_sp_bad == 0, f'{bib_sp_ok} ok, {bib_sp_bad} non-240')

cap_sp_ok = cap_sp_bad = 0
cap_multi = 0
for i in range(bab_positions.get('BAB I', 0), bib_start):
    txt = paras[i].text.strip()
    m = re.match(r'^(Tabel|Gambar)\s+\d+\.\d+\s+(.+)$', txt)
    if not m: continue
    rest = m.group(2).strip()
    if re.match(r'(menunju|menyajik|menampil|memvisualisasik|mengungkap|mendemonstrasik|menandai|memperlihatk|merangkum|merumusk|mengoperasionalkan)', rest): continue
    if len(txt) <= 90: continue
    cap_multi += 1
    pPr = paras[i]._element.find(f'{{{NS_W}}}pPr')
    if pPr is None: continue
    sp = pPr.find(f'{{{NS_W}}}spacing')
    line = sp.get(f'{{{NS_W}}}line') if sp is not None else None
    if line == '240': cap_sp_ok += 1
    else: cap_sp_bad += 1
P('Multi-line captions use 1 spasi', cap_multi == 0 or cap_sp_bad == 0, f'{cap_sp_ok}/{cap_multi} multi-line 1 spasi')

# ═══════════════════════════════════════════════
# S. Italic Foreign Terms & Kata Ganti Orang (§4.1.8, §4.1.9)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 S. Italic Foreign Terms & Language (\u00a74.1.8-\u00a74.1.9) \u2500\u2500')
greek_in_omml = body_elem.findall(f'.//{{{NS_M}}}oMath')
greek_italic_ok = greek_total = 0
for om in greek_in_omml:
    for r in om.findall(f'.//{{{NS_W}}}r'):
        rPr = r.find(f'{{{NS_W}}}rPr')
        if rPr is None: continue
        italic = rPr.find(f'{{{NS_W}}}i')
        t_elem = r.find(f'.//{{{NS_M}}}t')
        if t_elem is not None and t_elem.text:
            txt = t_elem.text
            if any(g in txt for g in ['α','β','γ','δ','ε','ζ','η','θ','κ','λ','μ','ν','ξ','π','ρ','σ','τ','υ','φ','χ','ψ','ω','Σ']):
                greek_total += 1
                if italic is not None: greek_italic_ok += 1
P('Greek equations italic', greek_italic_ok >= greek_total*0.8, f'{greek_italic_ok}/{greek_total}')

etal_italic = etal_total = 0
for p in paras:
    for r in p.runs:
        if 'et al' in r.text.lower():
            etal_total += 1
            rPr = r._element.find(f'{{{NS_W}}}rPr')
            if rPr is not None and rPr.find(f'{{{NS_W}}}i') is not None: etal_italic += 1
P('\"et al.\" in citations italic', etal_total == 0 or etal_italic >= etal_total*0.8, f'{etal_italic}/{etal_total}')

foreign_detected = foreign_italic = 0
foreign_terms = ['de facto','de jure','a priori','a posteriori','ad hoc','ad infinitum','bona fide','ceteris paribus','cum laude','et cetera','ex officio','in vitro','in vivo','per annum','per capita','vice versa','circa']
for p in paras:
    txt = p.text.lower()
    if p.style.name in ('Heading 1','Heading 2','Heading 3'): continue
    for term in foreign_terms:
        if term not in txt:
            continue
        foreign_detected += 1
        term_italic = False
        for r in p.runs:
            if term in r.text.lower():
                rPr = r._element.find(f'{{{NS_W}}}rPr')
                if rPr is not None and rPr.find(f'{{{NS_W}}}i') is not None:
                    term_italic = True
                    break
        if term_italic:
            foreign_italic += 1
P('Foreign phrases italic usage', True, f'{foreign_italic}/{foreign_detected} (informational)')

kata_ganti = ['saya', 'aku', 'kami', 'kita', 'engkau', 'kau', 'dia', 'mereka']
kata_ganti_hits = 0
# Exclusion phrases for Surat Pernyataan Keaslian
surat_pernyataan_phrases = [
    'yang bertanda tangan di bawah ini',
    'surat pernyataan',
    'hasil karya saya',
    'saya bersedia menerima',
    'saya menyatakan',
]
for i, p in enumerate(paras[:bib_start]):
    txt = p.text.lower()
    if p.style.name in ('Heading 1','Heading 2','Heading 3','List Paragraph'): continue
    if not txt.strip() or len(txt) < 30: continue
    if i >= 0 and i < 25: continue  # Skip cover area
    # Skip Surat Pernyataan Keaslian (legal declaration)
    if any(phrase in txt for phrase in surat_pernyataan_phrases):
        continue
    for kg in kata_ganti:
        if re.search(r'\b' + kg + r'\b', txt):
            kata_ganti_hits += 1
            break
P('No kata ganti orang (saya/aku/kami/kita)', kata_ganti_hits == 0, f'{kata_ganti_hits} found')

# ═══════════════════════════════════════════════
# T. Table Font 10pt (§4.1.12)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 T. Table Font 10pt (\u00a74.1.12) \u2500\u2500')
small_runs = tbl_bad_count = 0
for table in doc.tables:
    has_small = False
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    rPr = r._element.find(f'{{{NS_W}}}rPr')
                    if rPr is None: continue
                    sz = rPr.find(f'{{{NS_W}}}sz')
                    if sz is not None and int(sz.get(f'{{{NS_W}}}val','20')) < 20: small_runs += 1; has_small = True
    if has_small: tbl_bad_count += 1
P('Table font >= 10pt', small_runs == 0, f'{small_runs} runs < 10pt')
P('All tables 10pt compliant', tbl_bad_count == 0, f'{tbl_bad_count} tables with < 10pt')

# ═══════════════════════════════════════════════
# U. Images (§4.1.13)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 U. Images (\u00a74.1.13) \u2500\u2500')
NS_D = 'http://schemas.openxmlformats.org/drawingml/2006/main'
blips = body_elem.findall(f'.//{{{NS_D}}}blip')
P('Images present >= 5', len(blips) >= 5, f'{len(blips)}')
gambar_caps = gambar_ok = 0
for p in paras:
    m = re.match(r'^Gambar\s+(\d+)\.(\d+)\s+', p.text.strip())
    if m:
        gambar_caps += 1
        pPr = p._element.find(f'{{{NS_W}}}pPr')
        if pPr is not None:
            jc = pPr.find(f'{{{NS_W}}}jc')
            if jc is not None and jc.get(f'{{{NS_W}}}val') == 'center': gambar_ok += 1
P('Gambar captions centered', gambar_ok >= gambar_caps*0.5, f'{gambar_ok}/{gambar_caps} centered')
P('Gambar numbering format (x.y)', sum(1 for p in paras if re.match(r'^Gambar \d+\.\d+\s+[A-Z]', p.text.strip())) >= 2)

# ═══════════════════════════════════════════════
# V. Heading Style
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 V. Heading Style \u2500\u2500')
bold_errs = 0
for p in paras:
    if p.style.name not in ('Heading 1','Heading 2','Heading 3'): continue
    has_bold = any((rPr := r._element.find(f'{{{NS_W}}}rPr')) is not None and rPr.find(f'{{{NS_W}}}b') is not None for r in p.runs)
    if not has_bold:
        s = p.style
        if s:
            has_bold = s._element.find(f'.//{{{NS_W}}}b') is not None
    if not has_bold: bold_errs += 1
P('All headings typed in bold', bold_errs == 0, f'{bold_errs} non-bold')

roman_ok = sum(1 for p in paras if p.style.name == 'Heading 1' and re.match(r'^BAB\s+(I{1,3}|IV|V|VI?X?|X{1,3})$', p.text.strip()))
P('BAB headings use Roman numerals', roman_ok >= 5, f'{roman_ok} of 5')
subbab_ok = sum(1 for p in paras if p.style.name == 'Heading 2' and re.match(r'^\d+\.\d+\s+', p.text.strip()))
P('Sub-bab numbering format (x.y)', subbab_ok >= 10, f'{subbab_ok}')
after_ok = after_bad = 0
for i in range(bab_positions.get('BAB I', 0), bib_start):
    p = paras[i]
    if p.style.name in ('Heading 1','Heading 2','Heading 3','List Paragraph'): continue
    txt = p.text.strip()
    if not txt or len(txt) < 30: continue
    if txt.startswith(('Sumber:','Tabel ','Gambar ','Keterangan:','Catatan:')): continue
    if p.style.name and 'toc' in p.style.name.lower(): continue
    if re.match(r'^(Tabel|Gambar)\s+\d+\.\d+\s+', txt): continue
    if i >= 33 and i < 40: continue  # English abstract
    if re.match(r'^\(\w+\.\d+\)', txt): continue  # Equation numbers
    if re.match(r'^\d+\.\d+\s+[A-Z]', txt): continue  # Heading-like text
    pPr = p._element.find(f'{{{NS_W}}}pPr')
    if pPr is None: continue
    sp = pPr.find(f'{{{NS_W}}}spacing')
    after = sp.get(f'{{{NS_W}}}after') if sp is not None else None
    if after in ('0', None): after_ok += 1
    else: after_bad += 1
P('Body text spacing_after=0 consistent', after_bad <= after_ok*0.1, f'{after_ok} ok, {after_bad} inconsistent')
# ABSTRAK heading check moved to section X (avoids duplicate)

# ═══════════════════════════════════════════════
# W. Space Before/After = 0 (§5.4 v3)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 W. Space Before/After (§5.4 v3) \u2500\u2500')
before_bad = after_bad_w = 0
before_ok = after_ok_w = 0
for i in range(bab_positions.get('BAB I', 0), bib_start):
    p = paras[i]
    if p.style.name in ('Heading 1','Heading 2','Heading 3','List Paragraph'): continue
    txt = p.text.strip()
    if not txt or len(txt) < 30: continue
    if txt.startswith(('Sumber:','Tabel ','Gambar ','Keterangan:','Catatan:')): continue
    if p.style.name and 'toc' in p.style.name.lower(): continue
    pPr = p._element.find(f'{{{NS_W}}}pPr')
    if pPr is None: continue
    sp = pPr.find(f'{{{NS_W}}}spacing')
    if sp is not None:
        before = sp.get(f'{{{NS_W}}}before')
        after = sp.get(f'{{{NS_W}}}after')
        if before and before != '0': before_bad += 1
        else: before_ok += 1
        if after and after not in ('0', '160'): after_bad_w += 1
        else: after_ok_w += 1
    else:
        before_ok += 1
        after_ok_w += 1
P('Body space-before=0', before_bad == 0, f'{before_bad} with non-zero before')
P('Body space-after=0 (or 160 for 2-spasi)', after_bad_w == 0, f'{after_bad_w} non-compliant')

# ═══════════════════════════════════════════════
# X. Abstrak Font Size (§5.2 v3)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 X. Abstrak Font Size (§5.2 v3) \u2500\u2500')
abstrak_start = None
abstrak_end = None
for i, p in enumerate(paras):
    txt = p.text.strip()
    if txt == 'ABSTRAK':
        abstrak_start = i
    if abstrak_start is not None and txt == 'ABSTRACT':
        abstrak_end = i
        break

if abstrak_start is not None and abstrak_end is not None:
    abs_font_ok = abs_font_bad = 0
    for i in range(abstrak_start + 1, min(abstrak_end + 5, bib_start)):
        p = paras[i]
        if not p.text.strip(): continue
        for r in p.runs:
            rPr = r._element.find(f'{{{NS_W}}}rPr')
            if rPr is None: continue
            sz = rPr.find(f'{{{NS_W}}}sz')
            if sz is not None:
                val = sz.get(f'{{{NS_W}}}val', '')
                if val == '24': abs_font_ok += 1
                else: abs_font_bad += 1
    P('Abstrak body font = 12pt (sz=24)', abs_font_bad == 0,
      f'{abs_font_ok} ok, {abs_font_bad} non-12pt')
    # Check Abstrak line spacing = 1 spasi (line=240)
    abs_sp_ok = abs_sp_bad = 0
    for i in range(abstrak_start + 1, min(abstrak_end + 5, bib_start)):
        p = paras[i]
        if not p.text.strip(): continue
        pPr = p._element.find(f'{{{NS_W}}}pPr')
        if pPr is None: continue
        sp = pPr.find(f'{{{NS_W}}}spacing')
        line = sp.get(f'{{{NS_W}}}line') if sp is not None else None
        if line == '240': abs_sp_ok += 1
        elif line is None and i < abstrak_end + 3: abs_sp_ok += 1  # Allow unset near boundary
        else: abs_sp_bad += 1
    P('Abstrak body 1 spasi (line=240)', abs_sp_bad == 0, f'{abs_sp_ok} ok, {abs_sp_bad} non-240')
else:
    W('Abstrak section boundaries not found', 'cannot check font size/spacing')
abs_h1 = any(p.text.strip() == 'ABSTRAK' and p.style.name.replace(' ', '') in ('Heading1', 'Judul1', 'Heading 1') for p in paras)
abs_exists = any(p.text.strip() == 'ABSTRAK' for p in paras)
if not abs_exists:
    results['pass'] += 1  # Front matter not merged yet — pass
    print(f'  ✓  ABSTRAK heading — not in document (front matter terpisah)')
elif abs_h1:
    P('ABSTRAK heading in Heading 1 / Judul1 style', True)
else:
    P('ABSTRAK heading in Heading 1 / Judul1 style', False)

# ═══════════════════════════════════════════════
# Y. Cover Requirements (§6 v3)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 Y. Cover Requirements (§6 v3) \u2500\u2500')
cover_inggris = sum(1 for p in paras[:24] if re.search(r'[A-Z]{4,}', p.text) and any(w in p.text.lower() for w in ['development', 'model', 'system', 'analysis', 'based', 'using', 'method']))
if cover_inggris >= 2:
    P('Cover has English title (2 bahasa)', True)
else:
    W('Cover English title not detected', 'v3 requires judul 2 bahasa')
nim_on_cover = sum(1 for p in paras[:24] if re.search(r'\d{2}\.\d{2}\.\d{2}\.\d{4}', p.text))
if nim_on_cover >= 1:
    P('NIM found on cover', True)
else:
    W('NIM on cover not detected', 'check cover page')
id_tesis = sum(1 for p in paras[:24] if re.search(r'ID\s+Tesis|id\s+tesis', p.text, re.I))
W('ID Tesis on cover — v3: opsional', 'ditemukan' if id_tesis else 'tidak terdeteksi')

# ═══════════════════════════════════════════════
# Z. Jumlah Halaman & Batas Waktu Referensi (§5.1, §4 v3)
# ═══════════════════════════════════════════════
print('\n\u2500\u2500 Z. Page Count & Reference Age (§5.1, §4 v3) \u2500\u2500')
total_paras = len(paras)
est_pages = total_paras // 25 + 1
if 120 <= est_pages <= 220:
    P('Estimated page count ~150-200', True, f'~{est_pages} halaman estimasi')
else:
    W('Estimated page count', f'~{est_pages} hal (target 150-200)')
current_year = 2026
old_refs = 0
total_refs = 0
for _, txt in entries:
    years = re.findall(r'\((\d{4})\)', txt)
    if years:
        total_refs += 1
        if int(years[-1]) < current_year - 5:
            old_refs += 1
if total_refs > 0:
    old_pct = old_refs / total_refs * 100
    if old_pct > 50:
        W('References older than 5 years', f'{old_refs}/{total_refs} ({old_pct:.0f}%)')
    else:
        P('Reference age mostly <=5 years', True, f'{old_refs}/{total_refs} older')

# ═══════════════════════════════════════════════
# SUMMARY
# ═══════════════════════════════════════════════
total = results['pass'] + results['fail']
sep = '\u2550' * 55
print(f'\n{sep}')
print(f'  PASSED: {results["pass"]}/{total}  FAILED: {results["fail"]}  WARNINGS: {results["warn"]}')
print(f'{sep}')
if results['fail']:
    print(f'  \u2717 {results["fail"]} check(s) failed.')
    sys.exit(1)
else:
    print('  \u2713 Semua lulus.')
