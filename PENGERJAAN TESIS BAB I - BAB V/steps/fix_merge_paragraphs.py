#!/usr/bin/env python3
"""
fix_merge_paragraphs.py — Menggabungkan paragraf pendek di BAB I

Berdasarkan analisis:
- 1.2 Rumusan Masalah: Par[36-39] (4 paragraf)  → 1 paragraf naratif
- 1.3 Batasan Masalah: Par[43-49] (7 paragraf)  → 3 paragraf grup
- 1.4 Hipotesis:       Par[54-56] (3 paragraf)  → 1 paragraf
- 1.5.1 Tujuan:        Par[59-62] (4 paragraf)  → 2 paragraf

Strategi:
  - Backup file sebelum modifikasi
  - Merge konten dengan narasi penghubung (konjungsi, transisi)
  - Paragraf yang sudah digabung dikosongkan teksnya (runs dihapus)
  - Print log perubahan secara detail

Usage:
  python3 fix_merge_paragraphs.py

Author: @explore mode
Date:   2026-06-08
"""

import shutil
import os
import sys
from datetime import datetime
from docx import Document

from pipeline import utils

DOC = str(utils.get_doc_path())


# ─── KONFIGURASI ────────────────────────────────────────────────────────────

FILE_PATH = DOC


# ============================================================================
# BUILDER FUNCTIONS — masing-masing menerima list paragraph object,
# mengembalikan string teks hasil merge dengan narasi penghubung.
# ============================================================================

def build_rms_text(pars):
    """
    1.2 Rumusan Masalah: gabung intro + 3 RQ jadi 1 paragraf naratif.

    Par[36]: "Berdasarkan latar belakang yang telah diuraikan, rumusan
              masalah dalam penelitian ini adalah sebagai berikut."
    Par[37]: "Bagaimana karakteristik segmentasi probabilistik ...?"
    Par[38]: "Bagaimana evolusi distribusi klaster ...?"
    Par[39]: "Bagaimana sistem otomasi analisis LLM ...?"
    """
    intro = pars[0].text.strip()
    rq1   = pars[1].text.strip()
    rq2   = pars[2].text.strip()
    rq3   = pars[3].text.strip()

    return (
        f"{intro} Penelitian ini bertujuan untuk menjawab tiga pertanyaan "
        f"berikut. Pertama, {rq1[0].lower()}{rq1[1:]} "
        f"Kedua, {rq2[0].lower()}{rq2[1:]} "
        f"Ketiga, {rq3[0].lower()}{rq3[1:]}"
    )


def build_batasan_text_group1(pars):
    """
    1.3 Batasan — Group 1: intro + batasan data.
    Par[43]: "Untuk memfokuskan penelitian, ditetapkan batasan masalah..."
    Par[44]: "Data penelitian terbatas pada data sekunder..."
    """
    intro = pars[0].text.strip()
    data  = pars[1].text.strip()
    return f"{intro} {data}"


def build_batasan_text_group2(pars):
    """
    1.3 Batasan — Group 2: LLM scope + GMM impl + jenis analisis.
    Par[45]: '"Otomasi analisis LLM" mencakup empat fungsi...'
    Par[46]: 'GMM diimplementasikan dengan Scikit-learn...'
    Par[47]: 'Analisis bersifat cross-sectional comparative longitudinal...'
    """
    llm  = pars[0].text.strip()
    gmm  = pars[1].text.strip()
    anal = pars[2].text.strip()
    return f"{llm} {gmm} {anal}"


def build_batasan_text_group3(pars):
    """
    1.3 Batasan — Group 3: geocoding + wilayah rekrutmen.
    Par[48]: 'Geocoding menggunakan GeoPy/Nominatim...'
    Par[49]: 'Wilayah rekrutmen terbatas pada data ITSNU Pekalongan...'
    """
    geo = pars[0].text.strip()
    wil = pars[1].text.strip()
    return f"{geo} {wil}"


def build_hipotesis_text(pars):
    """
    1.4 Hipotesis: gabung 3 hipotesis jadi 1 paragraf naratif.

    Par[54]: "Hipotesis 1 (H1): ARI < 0,30..."
    Par[55]: "Hipotesis 2 (H2): Cosine similarity..."
    Par[56]: "Hipotesis 3 (H3): GMM menghasilkan..."
    """
    h1 = pars[0].text.strip()
    h2 = pars[1].text.strip()
    h3 = pars[2].text.strip()
    return (
        "Penelitian ini mengajukan tiga hipotesis yang diuji secara empiris. "
        f"{h1} {h2} {h3}"
    )


def build_tujuan_text_group1(pars):
    """
    1.5.1 Tujuan — Group 1: Tujuan 1 + Tujuan 2.

    Par[59]: "Tujuan pertama menganalisis karakteristik segmentasi..."
    Par[60]: "Tujuan kedua mengukur stabilitas klaster..."
    """
    t1 = pars[0].text.strip()
    t2 = pars[1].text.strip()
    return f"{t1} {t2}"


def build_tujuan_text_group2(pars):
    """
    1.5.1 Tujuan — Group 2: Tujuan 3 + Tujuan 4.

    Par[61]: "Tujuan ketiga membangun sistem otomasi LLM..."
    Par[62]: "Tujuan keempat Menguji validitas hasil otomasi..."
             → perbaiki kapitalisasi "Menguji" menjadi "menguji"
    """
    t3 = pars[0].text.strip()
    t4 = pars[1].text.strip()
    # Perbaiki kapitalisasi: "Tujuan keempat Menguji" → "Tujuan keempat menguji"
    t4_fixed = t4.replace("Tujuan keempat Menguji", "Tujuan keempat menguji", 1)
    return f"{t3} {t4_fixed}"


# ============================================================================
# MERGE PLAN
# Struktur: {
#   "section_label": [
#       (start_par, end_par_exclusive, builder_fn, style_name),
#       ...
#   ],
#   ...
# }
# ============================================================================

MERGE_PLAN = {
    "1.2 Rumusan Masalah": [
        (36, 40, build_rms_text, "Normal"),
    ],
    "1.3 Batasan Masalah": [
        # Group 1: intro + data — par[43], par[44]  → 1 par Normal
        (43, 45, build_batasan_text_group1, "Normal"),
        # Group 2: LLM + GMM + analisis — par[45], par[46], par[47]  → 1 par List
        (45, 48, build_batasan_text_group2, "List Paragraph"),
        # Group 3: geocoding + wilayah — par[48], par[49]  → 1 par List
        (48, 50, build_batasan_text_group3, "List Paragraph"),
    ],
    "1.4 Hipotesis": [
        (54, 57, build_hipotesis_text, "Normal"),
    ],
    "1.5.1 Tujuan": [
        # Group 1: Tujuan 1 + Tujuan 2  — par[59], par[60]
        (59, 61, build_tujuan_text_group1, "Normal"),
        # Group 2: Tujuan 3 + Tujuan 4  — par[61], par[62]
        (61, 63, build_tujuan_text_group2, "Normal"),
    ],
}


# ─── UTILITIES ──────────────────────────────────────────────────────────────

def clear_paragraph(par):
    """Kosongkan semua runs dalam paragraph tanpa menghapus par itu sendiri."""
    for run in par.runs:
        run.text = ""
    # Fallback: hapus juga melalui XML langsung
    from docx.oxml.ns import qn
    for r_elem in par._element.findall(qn('w:r')):
        for t_elem in r_elem.findall(qn('w:t')):
            t_elem.text = ""


def backup_file(path):
    """Buat backup file dengan timestamp."""
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    bak = path.replace(".docx", f"_BACKUP_{ts}.docx")
    shutil.copy2(path, bak)
    print(f"[BACKUP] → {bak}")
    return bak


# ─── MAIN ───────────────────────────────────────────────────────────────────

def main():
    if not os.path.exists(FILE_PATH):
        print(f"[ERROR] File tidak ditemukan: {FILE_PATH}")
        return

    print("=" * 72)
    print("  fix_merge_paragraphs.py  —  Merge Paragraf Pendek BAB I")
    print("=" * 72)
    print()
    print(f"  Target file: {FILE_PATH}")
    print()

    # 1. Backup
    bak_path = backup_file(FILE_PATH)
    print()

    # 2. Buka dokumen
    doc = Document(FILE_PATH)
    total_groups = 0
    total_deleted = 0

    # 3. Proses setiap section
    for section_label, groups in MERGE_PLAN.items():
        print(f"\n  ── {section_label} ──")

        for (start, end, builder, keep_style) in groups:
            pars = [doc.paragraphs[i] for i in range(start, end)]
            original_texts = [(i, doc.paragraphs[i].text.strip()) for i in range(start, end)]

            # Skip jika semua kosong
            if not any(t for _, t in original_texts):
                print(f"     [SKIP] Par[{start}-{end-1}] — semua kosong")
                continue

            # Bangun teks hasil merge
            merged_text = builder(pars)

            # === Tulis ke paragraph pertama ===
            target_par = pars[0]
            clear_paragraph(target_par)
            target_par.text = merged_text

            # Set style
            try:
                target_par.style = doc.styles[keep_style]
            except Exception:
                pass

            # === Kosongkan paragraph sisanya ===
            for p in pars[1:]:
                clear_paragraph(p)

            # === Log ===
            n = len(pars)
            total_groups += 1
            total_deleted += (n - 1)

            print(f"     ✓ Par[{start}-{end-1}] ({n} par → 1 par)")
            for i, txt in original_texts:
                preview = txt[:90].replace("\n", " ")
                if preview:
                    print(f"       [{i}] {preview}...")
            print(f"       → {merged_text[:130]}...")
            print()

    # 4. Simpan
    doc.save(FILE_PATH)
    print()
    print("=" * 72)
    print("  SELESAI!")
    print(f"  - {total_groups} grup paragraf berhasil digabungkan")
    print(f"  - {total_deleted} paragraf dikosongkan (sisa merging)")
    print(f"  - File diperbarui: {FILE_PATH}")
    print(f"  - Backup:          {bak_path}")
    print("=" * 72)


if __name__ == "__main__":
    main()
