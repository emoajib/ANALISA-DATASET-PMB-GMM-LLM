#!/usr/bin/env python3
"""
Fix remaining compliance violations:
1. Bibliography entries: line=240 (1 spasi)
2. Multi-line captions only: line=240 (1 spasi)
3. "et al." in citations: make italic
4. Table font: sz=18 → sz=20 (9pt→10pt)
"""
import os, re, shutil, sys
from lxml import etree
import docx

DOC = sys.argv[1] if len(sys.argv) > 1 else os.environ.get('TESIS_DOC', 'Tesis_ITSNU_v10_Final.docx')
BACKUP = DOC.replace('.docx', '_BEFORE_FIX_REMAINING.docx')
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

shutil.copy2(DOC, BACKUP)
print(f"Backup: {BACKUP}")

doc = docx.Document(DOC)
paras = doc.paragraphs

# ── Locate bibliography ──
bib_start = None
for i, p in enumerate(paras):
    if p.text.strip() == 'DAFTAR PUSTAKA' or p.text.strip().startswith('DAFTAR PUSTAKA\n'):
        bib_start = i
        break

bib_entries = []
if bib_start:
    for i in range(bib_start + 1, len(paras)):
        if paras[i].text.strip():
            bib_entries.append(i)
        else:
            break

# ── Helper: ensure spacing ──
def set_spacing(p_elem, line_val, after_val=0):
    """Set spacing via direct XML (survives python-docx save)."""
    pPr = p_elem.find(f'{{{NS_W}}}pPr')
    if pPr is None:
        pPr = etree.SubElement(p_elem, f'{{{NS_W}}}pPr')
        p_elem.insert(0, pPr)
    # Remove existing spacing to avoid duplicates
    for sp in pPr.findall(f'{{{NS_W}}}spacing'):
        pPr.remove(sp)
    spacing = etree.SubElement(pPr, f'{{{NS_W}}}spacing')
    spacing.set(f'{{{NS_W}}}line', str(line_val))
    spacing.set(f'{{{NS_W}}}lineRule', 'auto')
    spacing.set(f'{{{NS_W}}}after', str(after_val))
    spacing.set(f'{{{NS_W}}}before', '0')

# ── Verb check for body paragraphs ──
def is_caption_txt(txt):
    m = re.match(r'^(Tabel|Gambar)\s+\d+\.\d+\s+(.+)$', txt)
    if not m:
        return False
    rest = m.group(2).strip()
    # Body paragraphs use verbs like menunjukkan, menyajikan, etc.
    if re.match(r'(menunju|menyajik|menampil|memvisualisasik|mengungkap|mendemonstrasik|menandai|memperlihatk|merangkum|merumusk|mengoperasionalkan)', rest):
        return False
    return True

# ── 1. Bibliography line=240 ──
for idx in bib_entries:
    set_spacing(paras[idx]._element, 240)
print(f"Fixed {len(bib_entries)} bibliography entries: line=240")

# ── 2. Multi-line captions line=240 ──
cap_fixed = 0
for i, p in enumerate(paras):
    txt = p.text.strip()
    if is_caption_txt(txt) and len(txt) > 90:
        set_spacing(p._element, 240)
        cap_fixed += 1
        print(f"  P{i} caption: '{txt[:60]}...'")

print(f"Fixed {cap_fixed} multi-line captions: line=240")

# ── 3. Italicize "et al." ──
etal_fixed = 0
for p in paras:
    for r in p.runs:
        if 'et al' in r.text.lower():
            rPr = r._element.find(f'{{{NS_W}}}rPr')
            if rPr is None:
                rPr = etree.SubElement(r._element, f'{{{NS_W}}}rPr')
                r._element.insert(0, rPr)
            if rPr.find(f'{{{NS_W}}}i') is None:
                etree.SubElement(rPr, f'{{{NS_W}}}i')
                etal_fixed += 1

print(f"Italicized {etal_fixed} 'et al.' occurrences")

# ── 3b. Italicize Latin phrases ──
latin_phrases = ['a priori', 'de facto', 'per se']
latin_fixed = 0
for p in paras:
    txt_lower = p.text.lower()
    for phrase in latin_phrases:
        if phrase in txt_lower:
            for r in p.runs:
                if phrase in r.text.lower():
                    rPr = r._element.find(f'{{{NS_W}}}rPr')
                    if rPr is None:
                        rPr = etree.SubElement(r._element, f'{{{NS_W}}}rPr')
                        r._element.insert(0, rPr)
                    if rPr.find(f'{{{NS_W}}}i') is None:
                        etree.SubElement(rPr, f'{{{NS_W}}}i')
                        latin_fixed += 1

print(f"Italicized {latin_fixed} Latin phrase occurrences")

# ── 4. Table font sz=18 → 20 ──
tbl_fixed = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    rPr = r._element.find(f'{{{NS_W}}}rPr')
                    if rPr is None: continue
                    sz = rPr.find(f'{{{NS_W}}}sz')
                    if sz is not None:
                        val = sz.get(f'{{{NS_W}}}val')
                        if val and int(val) < 20:
                            sz.set(f'{{{NS_W}}}val', '20')
                            tbl_fixed += 1

print(f"Fixed {tbl_fixed} table font runs to 10pt")

# ── 5. All captions: seragamkan line=240 + after=240 ──
def set_spacing_full(p_elem, line_val=None, after_val=None):
    pPr = p_elem.find(f'{{{NS_W}}}pPr')
    if pPr is None:
        pPr = etree.SubElement(p_elem, f'{{{NS_W}}}pPr')
        p_elem.insert(0, pPr)
    # Remove existing spacing
    for sp in pPr.findall(f'{{{NS_W}}}spacing'):
        pPr.remove(sp)
    spacing = etree.SubElement(pPr, f'{{{NS_W}}}spacing')
    if line_val is not None:
        spacing.set(f'{{{NS_W}}}line', str(line_val))
        spacing.set(f'{{{NS_W}}}lineRule', 'auto')
    if after_val is not None:
        spacing.set(f'{{{NS_W}}}after', str(after_val))
    spacing.set(f'{{{NS_W}}}before', '0')

cap_fixed2 = 0
for i, p in enumerate(paras):
    txt = p.text.strip()
    if re.match(r'^(Tabel|Gambar)\s+\d+\.\d+\s+[A-Z]', txt):
        set_spacing_full(p._element, line_val=240, after_val=240)
        cap_fixed2 += 1

print(f"Seragamkan {cap_fixed2} caption: line=240 after=240")

# ── 6. Body text spacing_after=0 and line=480 ──
bab1_idx = None
dp_idx = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if t == 'BAB I' and bab1_idx is None:
        bab1_idx = i
    if t.startswith('DAFTAR PUSTAKA'):
        dp_idx = i

if bab1_idx is not None and dp_idx is not None:
    body_fixed = 0
    for i in range(bab1_idx, dp_idx):
        p = paras[i]
        if p.style.name in ('Heading 1','Heading 2','Heading 3','List Paragraph'): continue
        txt = p.text.strip()
        if not txt or len(txt) < 30: continue
        if txt.startswith(('Sumber:','Tabel ','Gambar ','Keterangan:','Catatan:')): continue
        if p.style.name and 'toc' in p.style.name.lower(): continue
        if re.match(r'^(Tabel|Gambar)\s+\d+\.\d+\s+', txt): continue
        if re.match(r'^\(\w+\.\d+\)', txt): continue
        if re.match(r'^\d+\.\d+\s+[A-Z]', txt): continue
        
        # Enforce spacing
        set_spacing_full(p._element, line_val=480, after_val=0)
        body_fixed += 1
    print(f"Fixed {body_fixed} body paragraphs to line=480, after=0, before=0")

# ── 7. Abstrak Font Size and Spacing (12pt, 1 spasi) ──
abstrak_start = None
kp_idx = None
for i, p in enumerate(paras):
    txt = p.text.strip()
    if txt == 'ABSTRAK':
        abstrak_start = i
    if txt == 'KATA PENGANTAR' or txt.startswith('DAFTAR ISI'):
        kp_idx = i
        if abstrak_start is not None: break

if abstrak_start is not None and kp_idx is not None:
    abs_fixed = 0
    for i in range(abstrak_start + 1, kp_idx):
        p = paras[i]
        txt = p.text.strip()
        if not txt: continue
        if txt in ('ABSTRACT', 'ABSTRAK'): continue
        
        # Set 1 spasi
        set_spacing_full(p._element, line_val=240, after_val=0)
        
        # Set 12pt font
        for r in p.runs:
            rPr = r._element.find(f'{{{NS_W}}}rPr')
            if rPr is None:
                rPr = etree.SubElement(r._element, f'{{{NS_W}}}rPr')
                r._element.insert(0, rPr)
            sz = rPr.find(f'{{{NS_W}}}sz')
            if sz is None:
                sz = etree.SubElement(rPr, f'{{{NS_W}}}sz')
            sz.set(f'{{{NS_W}}}val', '24')
        abs_fixed += 1
    print(f"Fixed {abs_fixed} Abstrak body paragraphs to 12pt, 1 spasi")

doc.save(DOC)
print(f"\nSaved: {DOC}")
