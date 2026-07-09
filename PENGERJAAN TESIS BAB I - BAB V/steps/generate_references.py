#!/usr/bin/env python3
"""
generate_references.py — Generate & insert APA 7 Daftar Pustaka ke DOCX.

Fungsi:
1. Baca semua paragraf dari docx
2. Ekstrak semua pola sitasi dari teks
3. Cocokkan dengan database referensi yang sudah ditentukan
4. Buat daftar pustaka format APA 7 (hanging indent, single spacing)
5. Insert setelah heading "DAFTAR PUSTAKA"
6. Urut alfabetis

Usage:
    python3 generate_references.py [path_to_docx]

Default:
    BAB I - BAB IV.docx
"""

import re
import sys
import os
import copy
from lxml import etree

# ─── Namespace ───
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

# ─── File paths ───
DEFAULT_DOC = 'BAB I - BAB IV.docx'

# ─── Database Referensi Lengkap (APA 7) ───
# Format: (key, entry_text)
# key = lowercase, stripped of punctuation, untuk pencocokan
REFERENCE_DB = [
    {
        'key': 'ahmadian 2024',
        'authors': 'Ahmadian, M.',
        'text': 'Ahmadian, M., Bateni, M., Esfandiari, H., Lattanzi, S., Monemizadeh, M., & Norouzi-Fard, A. (2024). Resilient k-clustering. In R. Baeza-Yates & F. Bonchi (Eds.), <i>Proceedings of the 30th ACM SIGKDD Conference on Knowledge Discovery and Data Mining (KDD 2024)</i> (pp. 29–38). ACM. https://doi.org/10.1145/3637528.3671909'
    },
    {
        'key': 'aristovnik 2021',
        'authors': 'Aristovnik, A.',
        'text': 'Aristovnik, A., Keržič, D., Ravšelj, D., Tomaževič, N., & Umek, L. (2021). Impacts of the COVID-19 pandemic on life of higher education students: A global perspective. <i>Data in Brief</i>, <i>39</i>, 107594. https://doi.org/10.1016/j.dib.2021.107594'
    },
    {
        'key': 'cahyadi 2021',
        'authors': 'Cahyadi, A.',
        'text': 'Cahyadi, A., Hendryadi, H., Widyastuti, S., & Suryani, S. (2021). COVID-19 emergency remote teaching and learning: A systematic literature review. <i>Education and Information Technologies</i>, <i>26</i>(6), 6995–7023. https://doi.org/10.1007/s10639-021-10625-8'
    },
    {
        'key': 'dempster 1977',
        'authors': 'Dempster, A. P.',
        'text': 'Dempster, A. P., Laird, N. M., & Rubin, D. B. (1977). Maximum likelihood from incomplete data via the EM algorithm. <i>Journal of the Royal Statistical Society: Series B (Methodological)</i>, <i>39</i>(1), 1–22. https://doi.org/10.1111/j.2517-6161.1977.tb01600.x'
    },
    {
        'key': 'devlin 2019',
        'authors': 'Devlin, J.',
        'text': 'Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. In J. Burstein, C. Doran, & T. Solorio (Eds.), <i>Proceedings of the 2019 Conference of the North American Chapter of the Association for Computational Linguistics: Human Language Technologies (NAACL-HLT 2019)</i> (pp. 4171–4186). Association for Computational Linguistics. https://doi.org/10.18653/v1/N19-1423'
    },
    {
        'key': 'el said 2021',
        'authors': 'El Said, G. R.',
        'text': 'El Said, G. R. (2021). How did the COVID-19 pandemic affect higher education learning experience? A case study. In C. Stephanidis, M. Antona, & S. Ntoa (Eds.), <i>Proceedings of the 2021 International Conference on Human-Computer Interaction (HCII 2021)</i> (pp. 311–324). Springer. https://doi.org/10.1007/978-3-030-78645-8_49'
    },
    {
        'key': 'george 2023',
        'authors': 'George, L.',
        'text': 'George, L., & Sumathy, S. (2023). An integrated clustering and BERT framework for improved sentiment analysis and topic modeling. <i>International Journal of Information Technology</i>, <i>15</i>(4), 2015–2027. https://doi.org/10.1007/s41870-023-01256-4'
    },
    {
        'key': 'grattafiori 2024',
        'authors': 'Grattafiori, A.',
        'text': 'Grattafiori, A., Dubey, A., Jauhri, A., Pandey, A., Kadian, A., Al-Dahle, A., ... & Meta AI Team. (2024). The Llama 3 herd of models. <i>arXiv preprint</i> arXiv:2407.21783. https://arxiv.org/abs/2407.21783'
    },
    {
        'key': 'hubert 1985',
        'authors': 'Hubert, L.',
        'text': 'Hubert, L., & Arabie, P. (1985). Comparing partitions. <i>Journal of Classification</i>, <i>2</i>(1), 193–218. https://doi.org/10.1007/BF01908075'
    },
    {
        'key': 'hossler 1987',
        'authors': 'Hossler, D.',
        'text': 'Hossler, D., & Gallagher, K. S. (1987). Studying student college choice: A three-phase model and the implications for policymakers. <i>College and University</i>, <i>62</i>(3), 207–221.'
    },
    {
        'key': 'imaduddin 2023',
        'authors': 'Imaduddin, M.',
        'text': 'Imaduddin, M., Wibowo, A. T., & Adiwijaya. (2023). Sentiment analysis in Indonesian healthcare using IndoBERT and feature extraction optimization. <i>International Journal of Advanced Computer Science and Applications (IJACSA)</i>, <i>14</i>(10), 10–19. https://doi.org/10.14569/IJACSA.2023.0141010'
    },
    {
        'key': 'jolliffe 2016',
        'authors': 'Jolliffe, I. T.',
        'text': 'Jolliffe, I. T., & Cadima, J. (2016). Principal component analysis: A review and recent developments. <i>Philosophical Transactions of the Royal Society A: Mathematical, Physical and Engineering Sciences</i>, <i>374</i>(2065), 20150202. https://doi.org/10.1098/rsta.2015.0202'
    },
    {
        'key': 'koto 2020',
        'authors': 'Koto, F.',
        'text': 'Koto, F., Rahimi, A., Lau, J. H., & Baldwin, T. (2020). IndoLEM and IndoBERT: A benchmark dataset and pre-trained language model for Indonesian NLP. In D. Scott, N. Bel, & C. Zong (Eds.), <i>Proceedings of the 28th International Conference on Computational Linguistics (COLING 2020)</i> (pp. 5579–5590). International Committee on Computational Linguistics. https://doi.org/10.18653/v1/2020.coling-main.490'
    },
    {
        'key': 'kotler 2016',
        'authors': 'Kotler, P.',
        'text': 'Kotler, P., & Keller, K. L. (2016). <i>Marketing management</i> (15th ed.). Pearson.'
    },
    {
        'key': 'macqueen 1967',
        'authors': 'MacQueen, J.',
        'text': 'MacQueen, J. (1967). Some methods for classification and analysis of multivariate observations. In L. M. Le Cam & J. Neyman (Eds.), <i>Proceedings of the Fifth Berkeley Symposium on Mathematical Statistics and Probability</i> (Vol. 1, pp. 281–297). University of California Press.'
    },
    {
        'key': 'mariscal 2010',
        'authors': 'Mariscal, G.',
        'text': 'Mariscal, G., Marbán, Ó., & Fernández, C. (2010). A survey of data mining and knowledge discovery process models and methodologies. <i>The Knowledge Engineering Review</i>, <i>25</i>(2), 137–166. https://doi.org/10.1017/S0269888910000032'
    },
    {
        'key': 'parker 2024',
        'authors': 'Parker, M. J.',
        'text': 'Parker, M. J., Anderson, C., Stone, C., & Oh, Y. R. (2024). A large language model approach to educational survey feedback analysis. <i>International Journal of Artificial Intelligence in Education</i>, <i>35</i>, 444–481. https://doi.org/10.1007/s40593-024-00414-0'
    },
    {
        'key': 'pedregosa 2011',
        'authors': 'Pedregosa, F.',
        'text': 'Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., ... & Duchesnay, É. (2011). Scikit-learn: Machine learning in Python. <i>Journal of Machine Learning Research</i>, <i>12</i>, 2825–2830.'
    },
    {
        'key': 'reimers 2019',
        'authors': 'Reimers, N.',
        'text': 'Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. In K. Inui, J. Jiang, V. Ng, & X. Wan (Eds.), <i>Proceedings of the 2019 Conference on Empirical Methods in Natural Language Processing and the 9th International Joint Conference on Natural Language Processing (EMNLP-IJCNLP 2019)</i> (pp. 3982–3992). Association for Computational Linguistics. https://doi.org/10.18653/v1/D19-1410'
    },
    {
        'key': 'scrucca 2024',
        'authors': 'Scrucca, L.',
        'text': 'Scrucca, L., Fop, M., Murphy, T. B., & Raftery, A. E. (2024). Model-based clustering in education via latent profile analysis: A review and practical guide. <i>Springer</i>. https://doi.org/10.1007/978-3-031-60012-8'
    },
    {
        'key': 'shearer 2000',
        'authors': 'Shearer, C.',
        'text': 'Shearer, C. (2000). The CRISP-DM model: The new blueprint for data mining. <i>Journal of Data Warehousing</i>, <i>5</i>(4), 13–22.'
    },
    {
        'key': 'shmueli 2016',
        'authors': 'Shmueli, G.',
        'text': 'Shmueli, G., Patel, N. R., & Bruce, P. C. (2016). <i>Data mining for business analytics: Concepts, techniques, and applications in R</i>. John Wiley & Sons.'
    },
    {
        'key': 'subakti 2022',
        'authors': 'Subakti, A.',
        'text': 'Subakti, A., Murfi, H., & Hariadi, N. (2022). The performance of BERT as data representation of text clustering. <i>Journal of Big Data</i>, <i>9</i>(1), 15. https://doi.org/10.1186/s40537-022-00572-7'
    },
    {
        'key': 'vaswani 2017',
        'authors': 'Vaswani, A.',
        'text': 'Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., ... & Polosukhin, I. (2017). Attention is all you need. In I. Guyon, U. V. Luxburg, S. Bengio, H. Wallach, R. Fergus, S. Vishwanathan, & R. Garnett (Eds.), <i>Proceedings of the 31st International Conference on Neural Information Processing Systems (NeurIPS 2017)</i> (pp. 5998–6008). Curran Associates.'
    },
    {
        'key': 'wei 2022',
        'authors': 'Wei, J.',
        'text': 'Wei, J., Wang, X., Schuurmans, D., Bosma, M., Ichter, B., Xia, F., ... & Zhou, D. (2022). Chain-of-thought prompting elicits reasoning in large language models. In S. Koyejo, S. Mohamed, A. Agarwal, D. Belgrave, K. Cho, & A. Oh (Eds.), <i>Proceedings of the 36th International Conference on Neural Information Processing Systems (NeurIPS 2022)</i> (pp. 24824–24837). Curran Associates.'
    },
    {
        'key': 'dubey 2024',
        'authors': 'Dubey, A.',
        'text': 'Dubey, A., Jauhri, A., Pandey, A., Kadian, A., Al-Dahle, A., Letman, A., ... & Meta AI Team. (2024). The Llama 3 model card. <i>Meta AI Research</i>. https://ai.meta.com/research/publications/the-llama-3-herd-of-models/'
    },
    {
        'key': 'ollama 2024',
        'authors': 'Ollama',
        'text': 'Ollama. (2024). <i>Ollama: Get up and running with large language models</i> [Documentation]. https://ollama.ai'
    },
]


def load_docx(path):
    """Load a DOCX file using python-docx."""
    import docx
    return docx.Document(path)


def extract_citations(text):
    """
    Extract citation patterns from text.
    Returns set of (author_key, year) tuples.
    """
    citations = set()

    # Pattern 1: (Author et al., Tahun) atau (Author & Author, Tahun) atau (Author, Tahun)
    # More comprehensive pattern for parenthetical citations
    paren_pattern = re.compile(r'\(([A-Z][a-zA-Zà-ÿ&.,\s]+?)\s*,\s*(\d{4})\)')
    for m in paren_pattern.finditer(text):
        name_part = m.group(1).strip()
        year = m.group(2)
        # Filter out non-citation patterns (numbers, common words)
        if re.match(r'^[A-Z][a-zà-ÿ]+(?:\s+et\s+al\.?)?(?:\s+[&]\s+[A-Z][a-zà-ÿ]+)?$', name_part):
            citations.add((name_part.lower(), year))
        # Also catch multi-author with &
        elif re.match(r'^[A-Z][a-zà-ÿ]+\s+[&]\s+[A-Z][a-zà-ÿ]+$', name_part):
            citations.add((name_part.lower(), year))

    # Pattern 2: Author et al. (Tahun) — narrative
    narr_pattern = re.compile(r'([A-Z][a-zà-ÿ]+(?:\s+et\s+al\.?)?)\s*\((\d{4})\)')
    for m in narr_pattern.finditer(text):
        name_part = m.group(1).strip()
        year = m.group(2)
        if len(name_part) > 2 and name_part[0].isupper():
            citations.add((name_part.lower(), year))

    # Pattern 3: Author & Author (Tahun) — narrative
    narr2_pattern = re.compile(r'([A-Z][a-zà-ÿ]+\s+[&]\s+[A-Z][a-zà-ÿ]+)\s*\((\d{4})\)')
    for m in narr2_pattern.finditer(text):
        citations.add((m.group(1).strip().lower(), m.group(2)))

    # Pattern 4: Author and Author (Tahun) — narrative
    narr3_pattern = re.compile(r'([A-Z][a-zà-ÿ]+\s+and\s+[A-Z][a-zà-ÿ]+)\s*\((\d{4})\)')
    for m in narr3_pattern.finditer(text):
        citations.add((m.group(1).strip().lower(), m.group(2)))

    # Pattern 5: Multi-citation: (Author, Year; Author, Year)
    multi_paren = re.compile(r'\(([A-Z][a-zA-Zà-ÿ&.,\s;]+?)\s*,\s*(\d{4})\s*;\s*([A-Z][a-zA-Zà-ÿ&.,\s]+?)\s*,\s*(\d{4})\)')
    for m in multi_paren.finditer(text):
        citations.add((m.group(1).strip().lower(), m.group(2)))
        citations.add((m.group(3).strip().lower(), m.group(4)))

    return citations


def match_citations_to_references(citations):
    """
    Match extracted citations to the reference database.
    Returns list of matched reference entries in alphabetical order.
    """
    matched = []
    matched_keys = set()

    for cite_name, cite_year in citations:
        cite_clean = cite_name.replace('et al.', '').replace('&', '').strip()
        # Try to find matching reference
        for ref in REFERENCE_DB:
            ref_first_author = ref['authors'].split(',')[0].lower()
            ref_year = re.search(r'(\d{4})', ref['text'])
            ref_year_val = ref_year.group(1) if ref_year else ''

            # Check if this citation matches this reference
            name_match = (cite_clean.startswith(ref_first_author.split()[0]) or
                          ref_first_author.split()[0].startswith(cite_clean.split()[0]))
            year_match = abs(int(cite_year) - int(ref_year_val)) <= 1 if ref_year_val else False

            if name_match and year_match and ref['key'] not in matched_keys:
                matched.append(ref)
                matched_keys.add(ref['key'])
                break

    # Add any remaining references from the database that weren't matched
    # (include all references from the prescribed list)
    for ref in REFERENCE_DB:
        if ref['key'] not in matched_keys:
            matched.append(ref)
            matched_keys.add(ref['key'])

    # Sort alphabetically by entry text (APA 7 style)
    matched.sort(key=lambda r: r['text'].lower())

    return matched


def create_reference_paragraph(doc, ref_text):
    """
    Create a new paragraph with hanging indent and single spacing for a reference entry.
    Returns the paragraph object.
    """
    # Add a new paragraph
    p = doc.add_paragraph()

    # Parse the reference text and handle <i> tags for italics
    parts = re.split(r'(<i>|</i>)', ref_text)
    is_italic = False
    for part in parts:
        if part == '<i>':
            is_italic = True
        elif part == '</i>':
            is_italic = False
        else:
            run = p.add_run(part)
            if is_italic:
                run.italic = True

    # Set hanging indent via XML
    pPr = p._element.find(f'{{{NS_W}}}pPr')
    if pPr is None:
        pPr = etree.SubElement(p._element, f'{{{NS_W}}}pPr')
        p._element.insert(0, pPr)

    # Remove existing indentation
    for ind in pPr.findall(f'{{{NS_W}}}ind'):
        pPr.remove(ind)

    # Set hanging indent: left=720 twips, hanging=720 twips (0.5 inch)
    ind_elem = etree.SubElement(pPr, f'{{{NS_W}}}ind')
    ind_elem.set(f'{{{NS_W}}}left', '720')
    ind_elem.set(f'{{{NS_W}}}hanging', '720')

    # Remove existing spacing
    for sp in pPr.findall(f'{{{NS_W}}}spacing'):
        pPr.remove(sp)

    # Set single spacing (line=240 twips)
    sp_elem = etree.SubElement(pPr, f'{{{NS_W}}}spacing')
    sp_elem.set(f'{{{NS_W}}}line', '240')
    sp_elem.set(f'{{{NS_W}}}lineRule', 'auto')
    sp_elem.set(f'{{{NS_W}}}after', '0')
    sp_elem.set(f'{{{NS_W}}}before', '0')

    return p


def find_daftar_pustaka_paragraph(paras):
    """Find the paragraph index of DAFTAR PUSTAKA heading."""
    for i, p in enumerate(paras):
        if p.text.strip().upper() == 'DAFTAR PUSTAKA' or p.text.strip().upper().startswith('DAFTAR PUSTAKA\n'):
            return i
    return None


def insert_references(doc, ref_paragraphs):
    """
    Insert reference paragraphs into the document right after the DAFTAR PUSTAKA heading.
    Uses direct XML manipulation to insert at the correct position.
    """
    body = doc.element.body
    children = list(body)

    # Find DAFTAR PUSTAKA paragraph element
    dp_index = None
    dp_element = None
    for i, child in enumerate(children):
        if child.tag == f'{{{NS_W}}}p':
            # Check text
            texts = [t.text for t in child.findall(f'.//{{{NS_W}}}t') if t.text]
            full_text = ''.join(texts).strip()
            if full_text.upper() == 'DAFTAR PUSTAKA' or full_text.upper().startswith('DAFTAR PUSTAKA\n'):
                dp_index = i
                dp_element = child
                break

    if dp_index is None:
        print("ERROR: DAFTAR PUSTAKA heading not found!")
        return False

    # Remove any existing paragraphs between DAFTAR PUSTAKA and next section
    # (or between DAFTAR PUSTAKA and end of body)
    next_sectpr = None
    for j in range(dp_index + 1, len(children)):
        child = children[j]
        if child.tag == f'{{{NS_W}}}sectPr':
            next_sectpr = child
            break
        # Remove existing paragraph (old/broken references)
        if child.tag == f'{{{NS_W}}}p':
            body.remove(child)

    # Insert new reference paragraphs before the sectPr (or at end)
    insert_before = next_sectpr if next_sectpr is not None else None

    for ref_p in ref_paragraphs:
        if insert_before is not None and len(body) > 0:
            body.insert(list(body).index(insert_before), ref_p._element)
        else:
            body.append(ref_p._element)

    return True


def main():
    # ─── Parse args ───
    doc_path = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_DOC

    if not os.path.exists(doc_path):
        print(f"ERROR: File tidak ditemukan: {doc_path}")
        sys.exit(1)

    print(f"📄 Memproses: {doc_path}")

    # ─── Load document ───
    import docx as docx_lib
    doc = docx_lib.Document(doc_path)
    paras = doc.paragraphs

    # ─── Step 1: Ekstrak semua sitasi ───
    all_text = '\n'.join([p.text for p in paras])
    citations = extract_citations(all_text)

    print(f"\n📚 Sitasi terdeteksi: {len(citations)}")
    for cite_name, cite_year in sorted(citations):
        print(f"   • {cite_name.title()}, {cite_year}")

    # ─── Step 2: Cocokkan dengan database referensi ───
    matched_refs = match_citations_to_references(citations)

    print(f"\n📖 Referensi yang akan dimasukkan: {len(matched_refs)}")
    for i, ref in enumerate(matched_refs, 1):
        # Extract first 80 chars for display
        display = re.sub(r'<[^>]+>', '', ref['text'])[:80]
        print(f"   {i:2d}. {display}...")

    # ─── Step 3: Buat paragraf referensi ───
    ref_paragraphs = []
    for ref in matched_refs:
        p = create_reference_paragraph(doc, ref['text'])
        ref_paragraphs.append(p)

    # ─── Step 4: Insert ke dokumen ───
    success = insert_references(doc, ref_paragraphs)

    if not success:
        print("❌ Gagal menyisipkan daftar pustaka!")
        sys.exit(1)

    # ─── Step 5: Simpan ───
    doc.save(doc_path)
    print(f"\n✅ Daftar Pustaka berhasil disisipkan!")
    print(f"   {len(ref_paragraphs)} entries dalam format APA 7")
    print(f"   Hanging indent + single spacing diterapkan")
    print(f"   Disimpan ke: {doc_path}")

    # ─── Step 6: Verifikasi ───
    verify_doc = docx_lib.Document(doc_path)
    dp_idx = find_daftar_pustaka_paragraph(verify_doc.paragraphs)

    if dp_idx is not None:
        ref_count = 0
        for j in range(dp_idx + 1, len(verify_doc.paragraphs)):
            if verify_doc.paragraphs[j].text.strip():
                ref_count += 1
            else:
                break
        print(f"\n🔍 Verifikasi: {ref_count} entri ditemukan setelah DAFTAR PUSTAKA")
        print(f"   Total paragraf: {len(verify_doc.paragraphs)}")
    else:
        print("⚠️  DAFTAR PUSTAKA heading tidak ditemukan setelah saving!")


if __name__ == '__main__':
    main()
