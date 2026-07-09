#!/usr/bin/env python3
"""
fix_pakar_label.py — Perbaiki duplikasi label "Pakar 2" di BAB IV

MASALAH:
  Di Section 4.9 Hasil Validasi Pakar, terdapat dua paragraf yang sama-sama
  diberi label "Pakar 2" tetapi merujuk pada dua orang yang berbeda:

    Par[425]: Pakar 2 (Andri Nur Cahyo, S.Sn., M.Sn.)   ← seharusnya Pakar 1
    Par[426]: Pakar 2 (Muhammad Kholilurrahman)           ← Pakar 2 (benar)

  Tabel 4.19 (Table index 31) sudah mendefinisikan kolom "Pakar 1" dan
  "Pakar 2" secara terpisah. Namun di narasi teks, kedua pakar sama-sama
  disebut "Pakar 2", menyebabkan inkonsistensi.

PERBAIKAN:
  Ubah Par[425] dari "Pakar 2 (Andri Nur Cahyo...)" menjadi
  "Pakar 1 (Andri Nur Cahyo...)".

  Alasan: (a) Andri Nur Cahyo disebut lebih dulu dalam urutan narasi;
          (b) Beliau adalah kepala PMB (lebih senior); 
          (c) Tabel 4.19 membedakan Pakar 1 dan Pakar 2.

Usage:
  python3 fix_pakar_label.py

Author: @explore mode
Date:   2026-06-08
"""

import shutil
import os
import re
import sys
from datetime import datetime
from docx import Document

from pipeline import utils


# ─── KONFIGURASI ────────────────────────────────────────────────────────────

DOC = str(utils.get_doc_path())
FILE_PATH = DOC

# old_substring digunakan untuk deteksi (print log) dan replacement.
# new_substring adalah teks pengganti.
TARGET_SUBSTRING = "Pakar 2 (Andri Nur Cahyo, S.Sn., M.Sn.)"
REPLACEMENT_SUBSTRING = "Pakar 1 (Andri Nur Cahyo, S.Sn., M.Sn.)"


# ─── UTILITIES ──────────────────────────────────────────────────────────────

def backup_file(path):
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    bak = path.replace(".docx", f"_BACKUP_{ts}.docx")
    shutil.copy2(path, bak)
    print(f"[BACKUP] → {bak}")
    return bak


def get_full_text(par):
    """Kembalikan teks lengkap dari paragraph."""
    return par.text


# ─── MAIN ───────────────────────────────────────────────────────────────────

def main():
    if not os.path.exists(FILE_PATH):
        print(f"[ERROR] File tidak ditemukan: {FILE_PATH}")
        return

    print("=" * 72)
    print("  fix_pakar_label.py  —  Perbaiki Duplikasi Label 'Pakar 2'")
    print("=" * 72)
    print()

    # 1. Backup
    bak_path = backup_file(FILE_PATH)
    print()

    # 2. Buka dokumen
    doc = Document(FILE_PATH)

    # 3. Verifikasi: scan seluruh dokumen untuk "Pakar 2" dan "Pakar 1"
    print("  ── Pra-scan: Semua referensi Pakar ──")
    for i, p in enumerate(doc.paragraphs):
        if 'Pakar' in p.text:
            snippet = p.text[:150].replace('\n', ' ')
            print(f"     Par[{i:4d}]: {snippet}...")
    print()

    # 4. Lakukan replacement (secara dinamis tanpa hardcoded index)
    print("  ── Eksekusi Perbaikan ──")
    changes = []
    for par_idx, p in enumerate(doc.paragraphs):
        full = get_full_text(p)
        if TARGET_SUBSTRING in full or TARGET_SUBSTRING.lower() in full.lower():
            # Lakukan replacement case-insensitive
            new_text = re.sub(re.escape(TARGET_SUBSTRING), REPLACEMENT_SUBSTRING, full, flags=re.IGNORECASE)
            # Hapus konten lama
            for run in p.runs:
                run.text = ""
            p.text = new_text
            changes.append((par_idx, TARGET_SUBSTRING, REPLACEMENT_SUBSTRING))
            print(f"     ✓ Par[{par_idx}]: {TARGET_SUBSTRING}")
            print(f"       → {REPLACEMENT_SUBSTRING}")
            print(f"       Teks: {p.text[:130]}...")

    print()

    # 5. Verifikasi post-fix
    if changes:
        print("  ── Post-scan: verifikasi ──")
        for i, p in enumerate(doc.paragraphs):
            if 'Pakar 2' in p.text:
                snippet = p.text[:150].replace('\n', ' ')
                print(f"     Par[{i:4d}]: {snippet}...")
        print()

    # 6. Simpan
    doc.save(FILE_PATH)

    print("=" * 72)
    if changes:
        print(f"  SELESAI! {len(changes)} perubahan dilakukan:")
        for idx, old, new in changes:
            print(f"    - Par[{idx}]: {old} → {new}")
    else:
        print("  SELESAI! Tidak ada perubahan.")
    print(f"  Backup: {bak_path}")
    print("=" * 72)


if __name__ == "__main__":
    main()
