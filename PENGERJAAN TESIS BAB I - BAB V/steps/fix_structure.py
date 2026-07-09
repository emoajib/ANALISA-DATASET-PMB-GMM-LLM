#!/usr/bin/env python3
"""Fix headings + BAB V body + citations."""
import sys, os, re, shutil
from lxml import etree
import docx

NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

DOC = sys.argv[1] if len(sys.argv) > 1 else os.environ.get('TESIS_DOC', 'Tesis_ITSNU_v10_Final.docx')
BACKUP = DOC.replace('.docx', '_BEFORE_STRUCTFIX.docx')


def set_outline_level(pPr, level):
    """Set or replace outline level."""
    existing = pPr.find(f'{{{NS_W}}}outlineLvl')
    if existing is not None:
        pPr.remove(existing)
    ol = etree.SubElement(pPr, f'{{{NS_W}}}outlineLvl')
    ol.set(f'{{{NS_W}}}val', str(level))


def set_alignment(pPr, val):
    """Set paragraph justification."""
    jc = pPr.find(f'{{{NS_W}}}jc')
    if jc is None:
        jc = etree.SubElement(pPr, f'{{{NS_W}}}jc')
    jc.set(f'{{{NS_W}}}val', val)


def set_spacing_after(pPr, val):
    """Set space after in twips."""
    spacing = pPr.find(f'{{{NS_W}}}spacing')
    if spacing is None:
        spacing = etree.SubElement(pPr, f'{{{NS_W}}}spacing')
    spacing.set(f'{{{NS_W}}}after', str(val))


def set_first_line_indent(pPr, val):
    """Set first-line indent in twips."""
    ind = pPr.find(f'{{{NS_W}}}ind')
    if ind is None:
        ind = etree.SubElement(pPr, f'{{{NS_W}}}ind')
    ind.set(f'{{{NS_W}}}first', str(val))


def make_run(text, bold=False, size=24):
    """Create a w:r element."""
    r = etree.SubElement(etree.Element('dummy'), f'{{{NS_W}}}r')
    rPr = etree.SubElement(r, f'{{{NS_W}}}rPr')
    rFonts = etree.SubElement(rPr, f'{{{NS_W}}}rFonts')
    rFonts.set(f'{{{NS_W}}}ascii', 'Times New Roman')
    rFonts.set(f'{{{NS_W}}}hAnsi', 'Times New Roman')
    rFonts.set(f'{{{NS_W}}}cs', 'Times New Roman')
    sz = etree.SubElement(rPr, f'{{{NS_W}}}sz')
    sz.set(f'{{{NS_W}}}val', str(size))
    szCs = etree.SubElement(rPr, f'{{{NS_W}}}szCs')
    szCs.set(f'{{{NS_W}}}val', str(size))
    if bold:
        etree.SubElement(rPr, f'{{{NS_W}}}b')
    t = etree.SubElement(r, f'{{{NS_W}}}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    return r


def fix_heading(p_elem, new_style, alignment, outline_level, text):
    """Convert a paragraph to a proper heading."""
    # Find and remove all existing runs
    runs = p_elem.findall(f'{{{NS_W}}}r')
    for r in runs:
        p_elem.remove(r)
    
    pPr = p_elem.find(f'{{{NS_W}}}pPr')
    if pPr is None:
        pPr = etree.SubElement(p_elem, f'{{{NS_W}}}pPr')
        p_elem.insert(0, pPr)
    
    # Set alignment
    set_alignment(pPr, alignment)
    # Set outline level
    set_outline_level(pPr, outline_level)
    
    # Create new bold run
    new_run = make_run(text, bold=True)
    p_elem.append(new_run)


def main():
    shutil.copy2(DOC, BACKUP)
    doc = docx.Document(DOC)
    
    changes = 0
    
    # === C1: BAB V headings (dynamic detection) ===
    heading1_targets = {'BAB V', 'KESIMPULAN DAN SARAN'}
    heading2_targets = {'5.1 Kesimpulan', '5.3 Saran'}
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text in heading1_targets:
            fix_heading(p._element, 'Heading 1', 'center', 1, text)
            p.style = doc.styles['Heading 1']
            changes += 1
            print(f'  ✓ Heading 1 [{i}]: {text}')
        elif text in heading2_targets:
            fix_heading(p._element, 'Heading 2', 'left', 2, text)
            p.style = doc.styles['Heading 2']
            changes += 1
            print(f'  ✓ Heading 2 [{i}]: {text}')
    
    # === C2: Add outline levels to ALL existing headings ===
    for i, p in enumerate(doc.paragraphs):
        pPr = p._element.find(f'{{{NS_W}}}pPr')
        if pPr is None:
            continue
        style = p.style.name if p.style else ''
        ol = pPr.find(f'{{{NS_W}}}outlineLvl')
        if style == 'Heading 1' and (ol is None or ol.get(f'{{{NS_W}}}val') != '1'):
            set_outline_level(pPr, 1)
            changes += 1
        elif style == 'Heading 2' and (ol is None or ol.get(f'{{{NS_W}}}val') != '2'):
            set_outline_level(pPr, 2)
            changes += 1
        elif style == 'Heading 3' and (ol is None or ol.get(f'{{{NS_W}}}val') != '3'):
            set_outline_level(pPr, 3)
            changes += 1
    
    print(f'  ✓ Outline levels set on all headings')
    
    # === M2: Seragamkan alignment headings ===
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ''
        pPr = p._element.find(f'{{{NS_W}}}pPr')
        if pPr is None:
            continue
        if style == 'Heading 1':
            set_alignment(pPr, 'center')
        elif style in ('Heading 2', 'Heading 3'):
            set_alignment(pPr, 'left')
    
    print(f'  ✓ Heading alignment seragam: H1=center, H2/H3=left')
    
    # === L1 + L2: BAB V body spacing + indent (dynamic BAB V area) ===
    bab5_start = None
    bab5_end = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == '5.1 Kesimpulan':
            bab5_start = i
        elif text == '5.3 Saran':
            bab5_end = i
        elif text == 'BAB V':
            if bab5_start is None:
                # Use as fallback marker
                pass
    
    if bab5_start is not None and bab5_end is not None:
        for i in range(bab5_start + 1, bab5_end):
            p = doc.paragraphs[i]
            style = p.style.name if p.style else ''
            if style == 'Normal' and p.text.strip():
                pPr = p._element.find(f'{{{NS_W}}}pPr')
                if pPr is None:
                    pPr = etree.SubElement(p._element, f'{{{NS_W}}}pPr')
                    p._element.insert(0, pPr)
                set_spacing_after(pPr, 0)  # v3: 0 space after
                set_first_line_indent(pPr, 720)  # 0.5in = 720 twips
                changes += 1
        print(f'  ✓ BAB V body (P{bab5_start+1}–P{bab5_end-1}): spacing_after=0, first_line_indent=0.5in')
    else:
        print(f'  ⚠ BAB V body not found (start={bab5_start}, end={bab5_end})')
    
    # === H2: Replace "dan" with "and" in citations ===
    author_pairs = [
        ('Purcell', 'Lumbreras'),
        ('Romero', 'Ventura'),
        ('George', 'Sumathy'),
        ('Reimers', 'Gurevych'),
        ('Jolliffe', 'Cadima'),
        ('Hubert', 'Arabie'),
        ('Kotler', 'Keller'),
    ]
    
    dan_fixes = 0
    for i, p in enumerate(doc.paragraphs):
        if i > 455:  # Skip Daftar Pustaka
            break
        for a1, a2 in author_pairs:
            pattern = f'{a1} dan {a2}'
            replacement = f'{a1} and {a2}'
            for r in p.runs:
                if pattern in r.text:
                    r.text = r.text.replace(pattern, replacement)
                    dan_fixes += 1
    
    print(f'  ✓ \"dan\" → \"and\": {dan_fixes} fixes')
    
    changes += dan_fixes
    
    doc.save(DOC)
    print(f'\nSaved: {DOC} ({changes} total changes)')


if __name__ == '__main__':
    main()
