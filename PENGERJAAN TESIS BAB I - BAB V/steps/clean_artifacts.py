#!/usr/bin/env python3
"""
clean_artifacts.py — Hapus artifacts: empty Heading 1 + duplikat ABSTRACT.
Jalankan SEBELUM pipeline.
"""
import os, sys, shutil
from lxml import etree
from docx import Document
from docx.oxml.ns import qn

DOC = sys.argv[1] if len(sys.argv) > 1 else os.environ.get(
    'TESIS_DOC',
    'Tesis_ITSNU_v10_Final.docx'
)
BACKUP = DOC.replace('.docx', '_BEFORE_CLEAN_ARTIFACTS.docx')


def para_text(elem):
    texts = [t.text for t in elem.findall('.//' + qn('w:t')) if t.text]
    return ''.join(texts).strip()


def get_style_val(elem):
    pPr = elem.find(qn('w:pPr'))
    if pPr is None: return ''
    ps = pPr.find(qn('w:pStyle'))
    return ps.get(qn('w:val'), '') if ps is not None else ''


def is_heading1(style_val):
    return style_val in ('Judul1', 'Heading1', 'heading1', '1')


def main():
    shutil.copy2(DOC, BACKUP)
    print(f'Backup: {BACKUP}')

    doc = Document(DOC)
    body = doc.element.body
    children = list(body)

    to_remove = []

    for i, c in enumerate(children):
        if c.tag != qn('w:p'):
            continue
        t = para_text(c)
        style = get_style_val(c)

        # 1a-1b: Empty Heading 1 paragraphs
        if is_heading1(style) and not t:
            to_remove.append((i, c, f'EMPTY H1 DELETE P{i}'))

        # 1c: Duplicate ABSTRACT (beyond first two: P19 ABSTRAK, P34 ABSTRACT)
        if i > 34 and t.upper() == 'ABSTRACT':
            to_remove.append((i, c, f'DUP ABSTRACT DELETE P{i}'))

    to_remove.sort(key=lambda x: x[0], reverse=True)
    print(f'Menghapus {len(to_remove)} artifact(s):')
    for idx, elem, desc in to_remove:
        print(f'  {desc}')
        try:
            body.remove(elem)
        except Exception as e:
            print(f'    ⚠ Gagal: {e}')

    doc.save(DOC)
    print(f'\nSaved: {DOC}')

    # Verify
    doc2 = Document(DOC)
    empty_h1 = 0
    dup_abstract = 0
    for i, p in enumerate(doc2.paragraphs):
        t = p.text.strip()
        s = p.style.name if p.style else ''
        if s == 'Heading 1' and not t:
            empty_h1 += 1
        if i > 2 and t == 'ABSTRACT' and s == 'Heading 1':
            dup_abstract += 1
    print(f'\nRemaining empty Heading 1: {empty_h1}')
    print(f'Remaining ABSTRACT (beyond first): {dup_abstract}')
    if empty_h1 == 0 and dup_abstract == 0:
        print('✅ Dokumen bersih dari artifacts.')
    else:
        print('⚠  Masih ada artifacts — periksa manual.')


if __name__ == '__main__':
    main()
