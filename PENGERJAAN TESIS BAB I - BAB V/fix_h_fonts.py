import docx
from lxml import etree
doc = docx.Document('../FULL TESIS/FULL TESIS FINAL.docx')
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Fix headings
for p in doc.paragraphs:
    if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
        for r in p.runs:
            rPr = r._element.find(f'{{{NS_W}}}rPr')
            if rPr is None:
                rPr = etree.SubElement(r._element, f'{{{NS_W}}}rPr')
                r._element.insert(0, rPr)
            sz = rPr.find(f'{{{NS_W}}}sz')
            if sz is None:
                sz = etree.SubElement(rPr, f'{{{NS_W}}}sz')
            sz.set(f'{{{NS_W}}}val', '24')

# Fix Abstrak
abstrak_start = None
kp_idx = None
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt == 'ABSTRAK':
        abstrak_start = i
    if txt == 'KATA PENGANTAR' or txt.startswith('DAFTAR ISI'):
        kp_idx = i
        if abstrak_start is not None: break

if abstrak_start is not None and kp_idx is not None:
    for i in range(abstrak_start + 1, kp_idx):
        p = doc.paragraphs[i]
        for r in p.runs:
            rPr = r._element.find(f'{{{NS_W}}}rPr')
            if rPr is None:
                rPr = etree.SubElement(r._element, f'{{{NS_W}}}rPr')
                r._element.insert(0, rPr)
            sz = rPr.find(f'{{{NS_W}}}sz')
            if sz is None:
                sz = etree.SubElement(rPr, f'{{{NS_W}}}sz')
            sz.set(f'{{{NS_W}}}val', '24')

doc.save('../FULL TESIS/FULL TESIS FINAL.docx')
