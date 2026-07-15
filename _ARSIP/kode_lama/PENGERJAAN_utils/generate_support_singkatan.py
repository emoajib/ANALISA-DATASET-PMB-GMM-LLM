import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# --- DAFTAR SINGKATAN ---
doc.add_heading("DAFTAR SINGKATAN", level=1)
p1 = doc.add_paragraph("Berikut adalah daftar singkatan yang digunakan di dalam penulisan tesis ini:")
p1.paragraph_format.space_after = Pt(12)

singkatan_data = [
    ("AIC", "Akaike Information Criterion"),
    ("API", "Application Programming Interface"),
    ("ARI", "Adjusted Rand Index"),
    ("BIC", "Bayesian Information Criterion"),
    ("CRISP-DM", "Cross-Industry Standard Process for Data Mining"),
    ("CSV", "Comma Separated Values"),
    ("GMM", "Gaussian Mixture Model"),
    ("IndoBERT", "Indonesian Bidirectional Encoder Representations from Transformers"),
    ("ITSNU", "Institut Teknologi dan Sains Nahdlatul Ulama"),
    ("LLM", "Large Language Model"),
    ("PCA", "Principal Component Analysis"),
    ("PMB", "Penerimaan Mahasiswa Baru"),
    ("TF-IDF", "Term Frequency - Inverse Document Frequency")
]

table_singkatan = doc.add_table(rows=len(singkatan_data), cols=3)
for i, (abbr, desc) in enumerate(singkatan_data):
    table_singkatan.cell(i, 0).text = abbr
    table_singkatan.cell(i, 1).text = "="
    table_singkatan.cell(i, 2).text = desc
    
    # Set widths
    table_singkatan.cell(i, 0).width = Cm(3.0)
    table_singkatan.cell(i, 1).width = Cm(0.5)
    table_singkatan.cell(i, 2).width = Cm(12.0)

doc.add_page_break()

# --- DAFTAR SIMBOL ---
doc.add_heading("DAFTAR SIMBOL", level=1)
p2 = doc.add_paragraph("Berikut adalah daftar simbol matematis maupun statistik yang digunakan di dalam penulisan tesis ini:")
p2.paragraph_format.space_after = Pt(12)

simbol_data = [
    ("K", "Jumlah klaster pembagian wilayah/segmen mahasiswa"),
    ("N", "Jumlah total data (observasi) calon mahasiswa baru"),
    ("x_i", "Vektor data (embedding) dari mahasiswa ke-i"),
    ("μ_k", "Vektor rata-rata (centroid/mean) dari klaster ke-k"),
    ("Σ_k", "Matriks kovarians (covariance matrix) dari klaster ke-k"),
    ("π_k", "Probabilitas prior (bobot campuran) dari klaster ke-k"),
    ("N(x|μ,Σ)", "Fungsi kepadatan probabilitas (PDF) distribusi Gaussian Multivariat"),
    ("γ(z_{nk})", "Responsibilitas / probabilitas posterior sampel n berada di klaster k"),
    ("S_i", "Silhouette Coefficient dari sampel ke-i")
]

table_simbol = doc.add_table(rows=len(simbol_data), cols=3)
for i, (sym, desc) in enumerate(simbol_data):
    table_simbol.cell(i, 0).text = sym
    table_simbol.cell(i, 1).text = "="
    table_simbol.cell(i, 2).text = desc
    
    # Set widths
    table_simbol.cell(i, 0).width = Cm(3.0)
    table_simbol.cell(i, 1).width = Cm(0.5)
    table_simbol.cell(i, 2).width = Cm(12.0)

doc.save("FULL TESIS/Bahan_Singkatan_Simbol.docx")
print("Bahan_Singkatan_Simbol.docx created successfully!")
