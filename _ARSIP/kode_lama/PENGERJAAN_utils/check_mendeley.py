import docx
doc = docx.Document("PENGERJAAN TESIS BAB I - BAB V/BAB I - BAB IV.docx")
found = False
for p in doc.paragraphs:
    if 'Mendeley' in p._element.xml:
        found = True
        break
print(f"Mendeley fields in source: {found}")

doc2 = docx.Document("FULL TESIS/FULL TESIS FINAL.docx")
found2 = False
for p in doc2.paragraphs:
    if 'Mendeley' in p._element.xml:
        found2 = True
        break
print(f"Mendeley fields in final: {found2}")
