import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# --- LAMPIRAN 2 ---
doc.add_heading("Lampiran 2 – Lembar Bimbingan Tesis", level=1)

doc.add_paragraph("LEMBAR BIMBINGAN TESIS").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Create a borderless table for Biodata
info_table = doc.add_table(rows=5, cols=3)
# To remove borders, we just don't set a style, default is no borders
info_data = [
    ("Nama Mahasiswa", "Mujibul Hakim"),
    ("NIM", "25.01.85.7010"),
    ("Program Studi", "Magister Teknologi Informasi"),
    ("Dosen Pembimbing", "Dr. Drs. Eri Zuliarso, M.Kom., MCF"),
    ("Judul Tesis", "Strategi Segmentasi Probabilistik Calon Mahasiswa Menggunakan Gaussian Mixture Model Dan Otomasi Analisis Large Language Model Untuk Optimalisasi Rekrutmen Di Itsnu Pekalongan")
]

for i, (label, value) in enumerate(info_data):
    info_table.cell(i, 0).text = label
    info_table.cell(i, 1).text = ":"
    info_table.cell(i, 2).text = value
    
    # Adjust widths (approximate for UI)
    info_table.cell(i, 0).width = Cm(4.0)
    info_table.cell(i, 1).width = Cm(0.5)
    info_table.cell(i, 2).width = Cm(11.0)

doc.add_paragraph() # spacing

p = doc.add_paragraph("Berikut adalah catatan perkembangan dan progres bimbingan tesis yang telah dilaksanakan minimal 8 (delapan) kali pertemuan:")
p.paragraph_format.space_after = Pt(12)

# Main Table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Tanggal Bimbingan'
hdr_cells[2].text = 'Materi / Pokok Bahasan'
hdr_cells[3].text = 'Tanda Tangan Pembimbing'

for row in table.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(4.0)
    row.cells[2].width = Cm(8.0)
    row.cells[3].width = Cm(3.5)

for i in range(1, 9):
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    for j in range(1, 4):
        row_cells[j].text = ""

for row in table.rows[1:]:
    row.height = Cm(1.5)

doc.save("FULL TESIS/Bahan_Lampiran_2_Final.docx")
print("Bahan_Lampiran_2_Final.docx created successfully!")
