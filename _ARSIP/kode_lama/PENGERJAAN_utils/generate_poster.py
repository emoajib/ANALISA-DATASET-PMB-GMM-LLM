import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

doc.add_heading("DRAF KONTEN POSTER PUBLIKASI TESIS", level=1)
p_intro = doc.add_paragraph("Panduan Pembuatan di Canva/Photoshop:\n1. Gunakan ukuran A1 (59.4 x 84.1 cm) posisi Portrait (Berdiri) atau Landscape (Mendatar).\n2. Bagi poster menjadi 3 kolom utama.\n3. Copy-Paste teks pendek di bawah ini ke dalam kotak-kotak desain poster Anda.\n4. Masukkan gambar dari folder 'outputs' sesuai instruksi di bawah ini.")
p_intro.style = 'Intense Quote'

doc.add_heading("BAGIAN ATAS (HEADER) - SEPANJANG LEBAR POSTER", level=2)
doc.add_paragraph("[Letakkan Logo ITSNU dan UNISBANK di sudut kiri dan kanan]")
doc.add_paragraph("STRATEGI SEGMENTASI PROBABILISTIK CALON MAHASISWA MENGGUNAKAN GAUSSIAN MIXTURE MODEL DAN OTOMASI ANALISIS LARGE LANGUAGE MODEL").bold = True
doc.add_paragraph("Mujibul Hakim (25.01.85.7010) | Pembimbing: Dr. Drs. Eri Zuliarso, M.Kom., MCF")
doc.add_paragraph("Magister Teknologi Informasi - Fakultas Teknologi Informasi dan Industri, UNISBANK")

doc.add_heading("KOLOM 1: PENDAHULUAN & METODOLOGI", level=2)
doc.add_heading("1. Latar Belakang & Tujuan", level=3)
doc.add_paragraph("Strategi marketing PMB ITSNU Pekalongan masih menggunakan pendekatan deskriptif konvensional. Penelitian ini memadukan NLP (IndoBERT), Soft-Clustering (GMM), dan Generative AI (LLM) untuk memetakan persona probabilitas calon mahasiswa guna strategi rekrutmen berbasis data.")
doc.add_heading("2. Metodologi (Hybrid Cognitive Pipeline)", level=3)
doc.add_paragraph("Tahapan CRISP-DM diimplementasikan dengan 3 pilar utama:")
doc.add_paragraph("- Text Embedding: IndoBERT (768 dimensi)\n- Clustering: Gaussian Mixture Model (GMM)\n- Reasoning: LLM Llama 3.3 (OpenRouter)")

doc.add_heading("KOLOM 2: HASIL & ANALISIS", level=2)
doc.add_heading("3. Hasil Segmentasi (K-Optimal & Time Series)", level=3)
doc.add_paragraph("[TEMPATKAN GAMBAR: gambar_4_3a_silhouette.png dan gambar_4_3c_ari.png DI SINI]")
doc.add_paragraph("Evaluasi menunjukkan jumlah klaster optimal berubah tiap periode (2-6 klaster). Analisis Adjusted Rand Index (ARI) mendeteksi adanya 'Structural Break' ekstrem (ARI = -0.0036) saat transisi pandemi COVID-19 (2019-2020), membuktikan bahwa segmentasi bersifat sangat dinamis.")
doc.add_heading("4. Generasi Persona LLM", level=3)
doc.add_paragraph("[TEMPATKAN GAMBAR: tabel_4_18_perbandingan.csv (Atau screenshot tabel dari tesis) DI SINI]")
doc.add_paragraph("LLM secara otomatis meringkas cluster matematika GMM menjadi persona humanis yang tervalidasi pakar, seperti 'Mahasiswa Teknis Lokal' atau 'Pencari Stabilitas Karir'.")

doc.add_heading("KOLOM 3: KESIMPULAN & REKOMENDASI", level=2)
doc.add_heading("5. Kesimpulan", level=3)
doc.add_paragraph("1. GMM & IndoBERT akurat menangani tumpang tindih probabilitas profil pendaftar.\n2. Pasar PMB bersifat sangat dinamis dan rentan terhadap disrupsi eksternal.\n3. LLM tervalidasi sangat kuat dalam mentransformasi big data menjadi narasi marketing.")
doc.add_heading("6. Rekomendasi 2025", level=3)
doc.add_paragraph("Kampus disarankan menerapkan dashboard otomasi PMB ini untuk melakukan penyesuaian promosi secara real-time dan hyper-personalized pada setiap klaster baru yang muncul di masa depan.")

doc.save("FULL TESIS/Draf_Poster_Publikasi.docx")
print("Draf_Poster_Publikasi.docx created successfully!")
