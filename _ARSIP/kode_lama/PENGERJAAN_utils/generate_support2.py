import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# --- LAMPIRAN 2 ---
doc.add_heading("Lampiran 2 – Lembar Bimbingan Tesis", level=1)

# Header Information
p_info = doc.add_paragraph()
runs = [
    ("Nama Mahasiswa", "Mujibul Hakim"),
    ("NIM", "25.01.85.7010"),
    ("Program Studi", "Magister Teknologi Informasi"),
    ("Dosen Pembimbing", "Dr. Drs. Eri Zuliarso, M.Kom., MCF"),
    ("Judul Tesis", "Strategi Segmentasi Probabilistik Calon Mahasiswa Menggunakan Gaussian Mixture Model Dan Otomasi Analisis Large Language Model Untuk Optimalisasi Rekrutmen Di Itsnu Pekalongan")
]

for label, value in runs:
    # Use tab stops to align the colons
    r = p_info.add_run(f"{label.ljust(20)} : {value}\n")
    if label == "Judul Tesis":
        # Remove trailing newline for the last item
        r.text = f"{label.ljust(20)} : {value}"

p_info.paragraph_format.space_after = Pt(12)

# Narrative (Optional, but good to have)
p = doc.add_paragraph("Berikut adalah catatan perkembangan dan progres bimbingan tesis yang telah dilaksanakan minimal 8 (delapan) kali pertemuan:")
p.paragraph_format.space_after = Pt(12)

# Table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Tanggal Bimbingan'
hdr_cells[2].text = 'Materi / Pokok Bahasan'
hdr_cells[3].text = 'Tanda Tangan Pembimbing'

# Set column widths (approximate)
for row in table.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(4.0)
    row.cells[2].width = Cm(8.0)
    row.cells[3].width = Cm(3.5)

# Add 8 empty rows for the student to fill
for i in range(1, 9):
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    for j in range(1, 4):
        row_cells[j].text = ""
        
# Set row heights so there's room to sign
for row in table.rows[1:]:
    row.height = Cm(1.5)

doc.save("FULL TESIS/Bahan_Lampiran_2_Revisi.docx")
print("Bahan_Lampiran_2_Revisi.docx created successfully!")
