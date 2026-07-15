import docx
import sys

def update_daftar(path):
    print("Updating Daftar lists and Lampiran...")
    doc = docx.Document(path)
    
    # 1. Find DAFTAR ISI and ABSTRAK
    di_idx = -1
    ab_idx = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().upper()
        if text == "DAFTAR ISI" and di_idx == -1:
            di_idx = i
        elif text == "ABSTRAK" and ab_idx == -1:
            ab_idx = i
            
    if di_idx != -1 and ab_idx != -1:
        # Clear paragraphs between DAFTAR ISI and ABSTRAK
        for i in range(di_idx + 1, ab_idx):
            doc.paragraphs[i].text = ""
            
        # The new content to insert after DAFTAR ISI
        new_content = [
            "",
            "DAFTAR TABEL",
            "",
            "DAFTAR GAMBAR",
            "",
            "DAFTAR SINGKATAN",
            "",
            "DAFTAR SIMBOL",
            "",
            "DAFTAR LAMPIRAN",
            "Lampiran 1 – Kumpulan Kode Pemrograman Python",
            "Lampiran 2 – Lembar Bimbingan Tesis",
            "Lampiran 3 – Surat Keterangan Bebas Plagiarisme",
            "Lampiran 4 – Hasil Cek Plagiarisme",
            "Lampiran 5 – LOA Jurnal Penelitian",
            ""
        ]
        
        current_p = doc.paragraphs[di_idx]
        for text in reversed(new_content):
            p = current_p.insert_paragraph_before(text)
            if text.startswith("DAFTAR "):
                p.style = 'Heading 1' if 'Heading 1' in [s.name for s in doc.styles] else 'Heading1'
        
        # We inserted before DAFTAR ISI, so DAFTAR ISI is pushed down.
        # Wait, if we want them AFTER DAFTAR ISI, we should insert before ABSTRAK!
        # Let's delete the newly inserted ones and insert before ABSTRAK instead.
        # A better way is just setting the text on the cleared paragraphs and adding new ones if needed.
        pass

    # Let's do it cleanly by recreating the Document or just manipulating the XML.
    # Actually, python-docx can insert_paragraph_before. We just find ABSTRAK and insert before it.
    
def clean_update(path):
    print("Clean updating Front Matter Lists...")
    doc = docx.Document(path)
    
    ab_idx = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().upper()
        if text == "ABSTRAK" and ab_idx == -1:
            ab_idx = i
            break
            
    if ab_idx != -1:
        # Clear everything between DAFTAR ISI and ABSTRAK
        di_idx = -1
        for i in range(ab_idx):
            if doc.paragraphs[i].text.strip().upper() == "DAFTAR ISI":
                di_idx = i
                break
                
        if di_idx != -1:
            for i in range(di_idx + 1, ab_idx):
                doc.paragraphs[i].text = ""
                
        # Now insert before ABSTRAK
        target_p = doc.paragraphs[ab_idx]
        new_content = [
            "",
            "DAFTAR TABEL",
            "",
            "DAFTAR GAMBAR",
            "",
            "DAFTAR SINGKATAN",
            "",
            "DAFTAR SIMBOL",
            "",
            "DAFTAR LAMPIRAN",
            "Lampiran 1 – Kumpulan Kode Pemrograman Python",
            "Lampiran 2 – Lembar Bimbingan Tesis",
            "Lampiran 3 – Surat Keterangan Bebas Plagiarisme",
            "Lampiran 4 – Hasil Cek Plagiarisme",
            "Lampiran 5 – LOA Jurnal Penelitian",
            ""
        ]
        
        for text in new_content:
            p = target_p.insert_paragraph_before(text)
            if text.startswith("DAFTAR "):
                p.style = 'Heading 1' if 'Heading 1' in [s.name for s in doc.styles] else 'Normal'
                # Attempt to set it, if style doesn't exist it might fail, fallback to Normal
                try:
                    p.style = 'Heading 1'
                except:
                    pass

    # Fix ABSTRAK Kata Kunci
    for p in doc.paragraphs:
        if "Kata Kunci:" in p.text:
            p.text = "Kata kunci: segmentasi, GMM, IndoBERT, LLM, rekrutmen, mahasiswa baru"
        elif "Keywords:" in p.text or "Keywords :" in p.text:
            p.text = "Keywords: segmentation, GMM, IndoBERT, LLM, recruitment, new students"

    doc.save(path)
    print("Done updating lists.")

if __name__ == "__main__":
    clean_update(sys.argv[1])
