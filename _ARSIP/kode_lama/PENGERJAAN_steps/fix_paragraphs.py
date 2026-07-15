#!/usr/bin/env python3
"""
fix_paragraphs.py — Unified paragraph formatting fixer.
Merged from: fix_bold_headings, fix_double_spaces, fix_indent,
fix_spacing, fix_outline_captions, fix_bib_indent,
fix_final_cleanup, fix_warnings.

Usage:
    python3 steps/fix_paragraphs.py <docx_path> --step all
    python3 steps/fix_paragraphs.py <docx_path> --step bold
    python3 steps/fix_paragraphs.py <docx_path> --step spacing
    python3 steps/fix_paragraphs.py <docx_path> --step indent
    python3 steps/fix_paragraphs.py <docx_path> --step cleanup
    python3 steps/fix_paragraphs.py <docx_path> --step warnings

Steps:
    bold        - Bold all headings + captions
    spaces      - Remove double spaces
    indent      - First-line indent 1.27cm body text
    spacing     - Line spacing 2.0 body, 1.0 bib/captions
    outline     - Heading outline levels + caption spacing
    bib_indent  - Hanging indent for bibliography
    cleanup     - Full cleanup (FLI, Normal style, spacing, empty paras)
    warnings    - ABSTRAK style, page numbering, placeholders
"""
import os, sys, re, shutil, zipfile
from datetime import datetime
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pipeline import utils
from pipeline import config

NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

doc_path = None
doc = None
backup_path = None


def load(path):
    global doc_path, doc, backup_path
    doc_path = utils.get_doc_path(path)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup_path = doc_path.parent / f'{doc_path.stem}_BEFORE_PARAFIX_{timestamp}{doc_path.suffix}'
    shutil.copy2(str(doc_path), str(backup_path))
    doc = utils.load_docx(doc_path)
    print(f'File: {doc_path}')
    print(f'Backup: {backup_path}')


def save():
    utils.save_docx(doc, doc_path)
    print(f'Saved: {doc_path}')


# ═══════════════════════════════════════════════════
# STEP: Bold headings
# ═══════════════════════════════════════════════════
def step_bold():
    total = 0
    for i, p in enumerate(doc.paragraphs):
        if p.style.name not in ('Heading 1', 'Heading 2', 'Heading 3'):
            continue
        changed = False
        for r in p.runs:
            if r.text.strip() and not r.bold:
                r.bold = True
                changed = True
        if changed:
            print(f'  [BOLD] Par[{i}] "{p.text[:60]}..."')
            total += 1
    print(f'Total heading di-bold: {total}')


# ═══════════════════════════════════════════════════
# STEP: Double spaces
# ═══════════════════════════════════════════════════
def step_spaces():
    total = 0
    paras = 0
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading'):
            continue
        fixed = 0
        for r in p.runs:
            if not r.text:
                continue
            new = re.sub(r'  +', ' ', r.text)
            if new != r.text:
                r.text = new
                fixed += 1
        if fixed:
            print(f'  [SPACES] Par[{i}] +{fixed} run(s)')
            total += fixed
            paras += 1
    print(f'Double spaces removed: {total} runs, {paras} paragraphs')


# ═══════════════════════════════════════════════════
# STEP: First-line indent
# ═══════════════════════════════════════════════════
def step_indent():
    total = 0
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading'):
            continue
        parent = p._p.getparent()
        in_table = False
        while parent is not None:
            if parent.tag == qn('w:tbl'):
                in_table = True
                break
            parent = parent.getparent()
        if in_table:
            continue
        if p._p.findall('.//' + qn('w:drawing')):
            continue
        if not p.text.strip():
            continue
        p.paragraph_format.first_line_indent = Cm(1.27)
        total += 1
    print(f'First-line indent applied: {total} paragraphs')


# ═══════════════════════════════════════════════════
# STEP: Line spacing
# ═══════════════════════════════════════════════════
def step_spacing():
    total = 0
    skip = 0
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'DAFTAR PUSTAKA' or p.text.strip().startswith('DAFTAR PUSTAKA\n'): break
        if p.style.name.startswith('Heading'):
            skip += 1
            continue
        parent = p._p.getparent()
        in_table = False
        while parent is not None:
            if parent.tag == qn('w:tbl'):
                in_table = True
                break
            parent = parent.getparent()
        if in_table:
            skip += 1
            continue
        sn = p.style.name.lower()
        if any(k in sn for k in ('caption', 'bibliography', 'daftar pustaka')):
            skip += 1
            continue
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = etree.SubElement(p._p, qn('w:pPr'))
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:line'), '480')
        sp.set(qn('w:lineRule'), 'auto')
        sp.set(qn('w:after'), '0')
        sp.set(qn('w:before'), '0')
        total += 1
    print(f'Line spacing 2.0: {total} paragraphs (skipped: {skip}) — after=0 before=0')


# ═══════════════════════════════════════════════════
# STEP: Outline levels + caption spacing (via ZIP/XML)
# ═══════════════════════════════════════════════════
def step_outline():
    tmp_backup = str(doc_path).replace('.docx', f'_BEFORE_OUTLINE_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    shutil.copy2(str(doc_path), tmp_backup)

    entries = {}
    with zipfile.ZipFile(str(doc_path), 'r') as z:
        for name in z.namelist():
            entries[name] = z.read(name)

    tree = etree.fromstring(entries['word/document.xml'])
    body = tree.find(f'{{{NS_W}}}body')
    paras = body.findall(f'{{{NS_W}}}p')

    style_outline = {'Judul1': '1', 'Judul2': '2', 'Judul3': '3'}
    fixed = {'Judul1': 0, 'Judul2': 0, 'Judul3': 0}

    for p in paras:
        pPr = p.find(f'{{{NS_W}}}pPr')
        if pPr is None:
            continue
        pStyle = pPr.find(f'{{{NS_W}}}pStyle')
        if pStyle is None:
            continue
        style_val = pStyle.get(f'{{{NS_W}}}val', '')
        if style_val not in style_outline:
            continue
        expected_lvl = style_outline[style_val]
        ol = pPr.find(f'{{{NS_W}}}outlineLvl')
        if ol is None:
            ol = etree.SubElement(pPr, f'{{{NS_W}}}outlineLvl')
        ol.set(f'{{{NS_W}}}val', expected_lvl)
        fixed[style_val] += 1

    print(f'Outline levels: H1={fixed["Judul1"]} H2={fixed["Judul2"]} H3={fixed["Judul3"]}')

    cap_fixed = 0
    for p in paras:
        text = ''.join(t.text or '' for t in p.iter(f'{{{NS_W}}}t')).replace('&amp;', '&').strip()
        if not re.match(r'^(Tabel|Gambar)\s+\d+\.\d+\s+', text):
            continue
        if len(text) <= 90:
            continue
        pPr = p.find(f'{{{NS_W}}}pPr')
        if pPr is None:
            pPr = etree.SubElement(p, f'{{{NS_W}}}pPr')
            p.insert(0, pPr)
        spacing = pPr.find(f'{{{NS_W}}}spacing')
        if spacing is None:
            spacing = etree.SubElement(pPr, f'{{{NS_W}}}spacing')
            pPr.append(spacing)
        spacing.set(f'{{{NS_W}}}line', '240')
        spacing.set(f'{{{NS_W}}}lineRule', 'auto')
        cap_fixed += 1

    print(f'Multi-line captions line=240: {cap_fixed}')

    entries['word/document.xml'] = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)
    with zipfile.ZipFile(str(doc_path), 'w', zipfile.ZIP_DEFLATED) as z:
        for name, data in entries.items():
            z.writestr(name, data)
    print(f'Outline + captions fixed via XML')


# ═══════════════════════════════════════════════════
# STEP: Bibliography hanging indent
# ═══════════════════════════════════════════════════
def step_bib_indent():
    dp_idx, _ = utils.find_daftar_pustaka(doc)
    if dp_idx is None:
        print('DAFTAR PUSTAKA heading not found')
        return
    fixed = 0
    for j in range(dp_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[j]
        if not p.text.strip():
            continue
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = etree.Element(qn('w:pPr'))
            p._element.insert(0, pPr)
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = etree.Element(qn('w:ind'))
            pPr.append(ind)
        ind.set(qn('w:hanging'), '720')
        ind.set(qn('w:left'), '720')
        fixed += 1
    print(f'Hanging indent applied: {fixed} bibliography entries')


# ═══════════════════════════════════════════════════
# STEP: Full cleanup
# ═══════════════════════════════════════════════════
def step_cleanup():
    paras = doc.paragraphs
    fli_removed = 0

    for i, p in enumerate(paras):
        text = p.text.strip()
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            continue
        jc = pPr.find(qn('w:jc'))
        is_centered = jc is not None and jc.get(qn('w:val')) == 'center'
        is_caption_text = bool(re.match(r'^(Tabel|Gambar)\s+\d+\.\d+\s+', text))
        is_cover_page = i < 24
        if not (is_centered or is_caption_text or is_cover_page):
            continue
        ind = pPr.find(qn('w:ind'))
        if ind is not None and ind.get(qn('w:firstLine')) is not None:
            ind.attrib.pop(qn('w:firstLine'), None)
            if len(ind.attrib) == 0:
                pPr.remove(ind)
            fli_removed += 1
    print(f'FLI removed from centered/captions/cover: {fli_removed}')

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.bold = False
    normal_style.font.italic = False
    rPr = normal_style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(normal_style.element, qn('w:rPr'))
        normal_style.element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = etree.SubElement(rPr, qn('w:sz'))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = etree.SubElement(rPr, qn('w:szCs'))
    sz.set(qn('w:val'), '24')
    szCs.set(qn('w:val'), '24')
    print('Normal style: Times New Roman 12pt')

    spacing_fixed = 0
    for i, p in enumerate(paras):
        text = p.text.strip()
        if not text or p.style.name.startswith('Heading'):
            continue
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            continue
        is_caption = bool(re.match(r'^(Tabel|Gambar)\s+\d+\.\d+\s+', text))
        is_bibliography = p.style.name.lower() in ('bibliography', 'daftar pustaka')
        is_daftar_pustaka_heading = (text == 'DAFTAR PUSTAKA' or text.startswith('DAFTAR PUSTAKA\n'))
        is_sumber = text.startswith('Sumber:')

        if is_caption or is_sumber:
            target_line = 240
            target_after = 240  # v3: 1 spasi antara caption dan teks
        elif is_bibliography or is_daftar_pustaka_heading:
            target_line = 240
            target_after = 0
        else:
            target_line = 480
            target_after = 0

        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:line'), str(target_line))
        sp.set(qn('w:lineRule'), 'auto')
        sp.set(qn('w:after'), str(target_after))
        sp.set(qn('w:before'), '0')
        spacing_fixed += 1
    print(f'Spacing standardized: {spacing_fixed} paragraphs')

    page_breaks_added = 0
    for i, p in enumerate(paras):
        text = p.text.strip()
        if re.match(r'^BAB\s+(I|II|III|IV|V|VI|VII|VIII|IX|X|XI|XII)\s*$', text) and i > 0:
            pPr = p._p.find(qn('w:pPr'))
            has_break = False
            if pPr is not None:
                for br in pPr.findall(qn('w:pageBreakBefore')):
                    has_break = True
                    break
            if not has_break:
                if pPr is None:
                    pPr = etree.SubElement(p._p, qn('w:pPr'))
                    p._p.insert(0, pPr)
                pb = etree.SubElement(pPr, qn('w:pageBreakBefore'))
                page_breaks_added += 1
    print(f'Page breaks added before BAB: {page_breaks_added}')

    indices_to_remove = []
    for i, p in enumerate(paras):
        text = p.text.strip()
        has_drawing = p._p.findall('.//' + qn('w:drawing'))
        if not text and not has_drawing:
            parent = p._p.getparent()
            in_table = False
            while parent is not None:
                if parent.tag == qn('w:tbl'):
                    in_table = True
                    break
                parent = parent.getparent()
            if not in_table:
                indices_to_remove.append(i)
    for idx in reversed(indices_to_remove):
        p_elem = paras[idx]._p
        p_elem.getparent().remove(p_elem)
    print(f'Empty paragraphs removed: {len(indices_to_remove)}')


# ═══════════════════════════════════════════════════
# STEP: Warnings fix (ABSTRAK style, page numbering)
# ═══════════════════════════════════════════════════
def step_warnings():
    from docx.oxml import OxmlElement
    body = doc.element.body
    children = list(body)
    changes = []

    def para_text(elem):
        texts = [t.text for t in elem.findall('.//' + qn('w:t')) if t.text]
        return ''.join(texts).strip()

    def ensure_pPr(para):
        pPr = para.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para.insert(0, pPr)
        return pPr

    def set_heading1_judul(para):
        pPr = ensure_pPr(para)
        for ps in pPr.findall(qn('w:pStyle')):
            pPr.remove(ps)
        ps = OxmlElement('w:pStyle')
        ps.set(qn('w:val'), 'Judul1')
        pPr.append(ps)
        for ol in pPr.findall(qn('w:outlineLvl')):
            pPr.remove(ol)
        ol = OxmlElement('w:outlineLvl')
        ol.set(qn('w:val'), '1')
        pPr.append(ol)
        for r in para.findall('.//' + qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            sz.set(qn('w:val'), '24')
            szCs = rPr.find(qn('w:szCs'))
            if szCs is None:
                szCs = OxmlElement('w:szCs')
                rPr.append(szCs)
            szCs.set(qn('w:val'), '24')

    def set_center(para):
        pPr = ensure_pPr(para)
        for jc in pPr.findall(qn('w:jc')):
            pPr.remove(jc)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)

    for i, c in enumerate(children):
        if c.tag != qn('w:p'):
            continue
        t = para_text(c)
        if t.upper() in ('ABSTRAK', 'ABSTRACT'):
            pPr = c.find(qn('w:pPr'))
            style_val = ''
            if pPr is not None:
                ps = pPr.find(qn('w:pStyle'))
                style_val = ps.get(qn('w:val'), '') if ps is not None else ''
            if style_val != 'Judul1':
                set_heading1_judul(c)
                set_center(c)
                changes.append(f'[P{i}] "{t}" → Judul1 + 12pt (was "{style_val}")')

    abstrak_positions = []
    for i, c in enumerate(children):
        if c.tag != qn('w:p'):
            continue
        t = para_text(c)
        if t.upper() == 'ABSTRAK':
            abstrak_positions.append(i)
        elif t.upper() == 'ABSTRACT':
            abstrak_positions.append(i)

    if len(abstrak_positions) > 2:
        to_remove = []
        for i, c in enumerate(children):
            if c.tag != qn('w:p'):
                continue
            t = para_text(c)
            if t.upper() in ('ABSTRAK', 'ABSTRACT', 'KATA KUNCI', 'KEYWORDS'):
                first_occurrence = abstrak_positions[0] if t.upper() == 'ABSTRAK' else abstrak_positions[1] if len(abstrak_positions) > 1 else -1
                if i != first_occurrence:
                    to_remove.append((i, c, f'DUP DELETE P{i}: {t[:60]}'))
            if '[Isi Abstract' in t and i >= (abstrak_positions[1] if len(abstrak_positions) > 1 else 0):
                to_remove.append((i, c, f'DUP DELETE P{i}: {t[:60]}'))
        if to_remove:
            to_remove.sort(key=lambda x: x[0], reverse=True)
            for idx, elem, desc in to_remove:
                body.remove(elem)
                changes.append(desc)

    sections = doc.sections
    for si, sec in enumerate(sections):
        if si == 0:
            sectPr = sec._sectPr
            pnt = sectPr.find(qn('w:pgNumType'))
            if pnt is None:
                pnt = OxmlElement('w:pgNumType')
                sectPr.append(pnt)
            pnt.set(qn('w:fmt'), 'romanLower')
            pnt.set(qn('w:start'), '1')
        else:
            sectPr = sec._sectPr
            pnt = sectPr.find(qn('w:pgNumType'))
            if pnt is None:
                pnt = OxmlElement('w:pgNumType')
                sectPr.append(pnt)
            pnt.set(qn('w:fmt'), 'decimal')

        footer = sec.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = 1
        run = p.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._element.append(fld_begin)
        run2 = p.add_run()
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run2._element.append(instr)
        run3 = p.add_run()
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        run3._element.append(fld_sep)
        run4 = p.add_run()
        run4.text = '1'
        run5 = p.add_run()
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run5._element.append(fld_end)
        for r in p.runs:
            rPr = r._element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r._element.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            sz.set(qn('w:val'), '20')
            szCs = rPr.find(qn('w:szCs'))
            if szCs is None:
                szCs = OxmlElement('w:szCs')
                rPr.append(szCs)
            szCs.set(qn('w:val'), '20')

    changes.append(f'Page number format set: {len(sections)} sections')

    # ── v3: Force abstrak body font to 12pt ──
    abstrak_paras = []
    in_abstrak = False
    for c in children:
        if c.tag != qn('w:p'):
            continue
        t = para_text(c)
        if t.upper() == 'ABSTRAK':
            in_abstrak = True
            continue
        if t.upper() == 'ABSTRACT':
            in_abstrak = False
            continue
        if in_abstrak and t.strip():
            abstrak_paras.append(c)

    abs_font_fixed = 0
    for para in abstrak_paras:
        for r in para.findall('.//' + qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            sz.set(qn('w:val'), '24')
            szCs = rPr.find(qn('w:szCs'))
            if szCs is None:
                szCs = OxmlElement('w:szCs')
                rPr.append(szCs)
            szCs.set(qn('w:val'), '24')
        abs_font_fixed += 1
    if abs_font_fixed:
        changes.append(f'Abstrak body font set to 12pt: {abs_font_fixed} paragraphs')

    print(f'{len(changes)} changes:')
    for ch in changes:
        print(f'  {ch}')


# ═══════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════
def main():
    import argparse
    parser = argparse.ArgumentParser(description='Paragraph formatting fixer (merged)')
    parser.add_argument('file', nargs='?', help='Path ke DOCX')
    parser.add_argument('--step', default='all',
                        choices=['all', 'bold', 'spaces', 'indent', 'spacing',
                                 'outline', 'bib_indent', 'cleanup', 'warnings'],
                        help='Step to run (default: all)')
    args = parser.parse_args()

    load(args.file)
    step = args.step

    if step in ('all', 'bold'):
        print('\n--- Step: Bold Headings ---')
        step_bold()
        save()
    if step in ('all', 'spaces'):
        print('\n--- Step: Double Spaces ---')
        step_spaces()
        save()
    if step in ('all', 'indent'):
        print('\n--- Step: First-line Indent ---')
        step_indent()
        save()
    if step in ('all', 'spacing'):
        print('\n--- Step: Line Spacing ---')
        step_spacing()
        save()
    if step in ('all', 'outline'):
        print('\n--- Step: Outline + Captions ---')
        step_outline()
    if step in ('all', 'bib_indent'):
        print('\n--- Step: Bib Hanging Indent ---')
        step_bib_indent()
        save()
    if step in ('all', 'cleanup'):
        print('\n--- Step: Full Cleanup ---')
        step_cleanup()
        save()
    if step in ('all', 'warnings'):
        print('\n--- Step: Warnings Fix ---')
        step_warnings()
        save()

    print(f'\nDone. Backup: {backup_path}')


if __name__ == '__main__':
    main()
