#!/usr/bin/env python3
"""
clean_bab5_duplicates.py — Hapus heading BAB V palsu + TOC corrupt.

Dynamic & idempotent: mendeteksi heading palsu berdasarkan posisi relatif,
bukan indeks absolut. Cocok untuk pipeline otomatis.

Strategi:
  1. Cari semua Heading 1 "BAB V" — keep yang TERAKHIR (real), hapus sisanya.
  2. Hapus juga "KESIMPULAN DAN SARAN" palsu yang mengikuti.
  3. Hapus entri TOC corrupt (style toc 1/2/3) antara DAFTAR ISI dan BAB I.
  4. Hapus empty Heading 1 artifacts.
"""
import os
import shutil
import sys
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

DOC = sys.argv[1] if len(sys.argv) > 1 else os.environ.get(
    'TESIS_DOC',
    'Tesis_ITSNU_v10_Final.docx'
)
BACKUP = DOC.replace('.docx', '_BEFORE_CLEAN_BAB5.docx')


def para_text(elem):
    texts = [t.text for t in elem.findall('.//' + qn('w:t')) if t.text]
    return ''.join(texts).strip()


def get_style_val(elem):
    pPr = elem.find(qn('w:pPr'))
    if pPr is None:
        return ''
    ps = pPr.find(qn('w:pStyle'))
    return ps.get(qn('w:val'), '') if ps is not None else ''


def is_heading1(style_val):
    return any(x in style_val for x in ('Heading1', 'heading1', 'Judul1', 'judul1', '1'))


def is_toc(style_val):
    return style_val.startswith('toc') or 'TOC' in style_val


def main():
    shutil.copy2(DOC, BACKUP)
    print(f'Backup: {BACKUP}')

    doc = Document(DOC)
    body = doc.element.body
    children = list(body)

    # ── Step 1: Find all BAB V heading positions ──
    bab5_positions = []
    daftar_isi_idx = None
    bab1_idx = None

    for i, c in enumerate(children):
        if c.tag != qn('w:p'):
            continue
        t = para_text(c)
        style = get_style_val(c)
        h1 = is_heading1(style)

        if h1 and t == 'BAB V':
            bab5_positions.append(i)
        if t == 'DAFTAR ISI':
            daftar_isi_idx = i
        if h1 and t == 'BAB I' and bab1_idx is None:
            bab1_idx = i

    # ── Step 2: Keep last BAB V, remove all earlier ones ──
    if len(bab5_positions) <= 1:
        print('Hanya 1 BAB V ditemukan — tidak ada duplikasi.')
        # Still clean TOC if needed
    else:
        real_bab5 = bab5_positions[-1]
        fake_bab5 = bab5_positions[:-1]
        print(f'Ditemukan {len(bab5_positions)} BAB V heading '
              f'(real=P{real_bab5}, {len(fake_bab5)} fake)')

    # ── Collect elements to remove ──
    to_remove = []
    keep_until = None  # Track spurious BAB V ranges to delete all content inside

    for i, c in enumerate(children):
        if c.tag != qn('w:p'):
            continue
        t = para_text(c)
        style = get_style_val(c)
        h1 = is_heading1(style)

        # ── Step 3: Delete old TOC entries ──
        if daftar_isi_idx is not None and bab1_idx is not None:
            if daftar_isi_idx < i < bab1_idx and is_toc(style):
                to_remove.append((i, c, f'TOC DELETE P{i} [{style}]: {t[:60]}'))

        # ── Step 4: Empty Heading 1 artifacts ──
        if h1 and not t:
            to_remove.append((i, c, f'EMPTY H1 DELETE P{i}'))

        # ── Step 5: Delete spurious BAB V + KESIMPULAN DAN SARAN ──
        if len(bab5_positions) > 1:
            is_fake_bab5 = (h1 and t == 'BAB V' and i != bab5_positions[-1])
            if is_fake_bab5:
                to_remove.append((i, c, f'FAKE BABV DELETE P{i}'))
            elif h1 and t == 'KESIMPULAN DAN SARAN':
                # Only delete if it immediately follows a fake BAB V
                for fi in fake_bab5:
                    if i == fi + 1:
                        to_remove.append((i, c, f'FAKE KDS DELETE P{i}'))
                        break

    # ── Remove from highest to lowest ──
    to_remove.sort(key=lambda x: x[0], reverse=True)
    print(f'Menghapus {len(to_remove)} element(s):')
    for idx, elem, desc in to_remove:
        print(f'  {desc}')
        body.remove(elem)

    doc.save(DOC)
    print(f'\nSaved: {DOC}')

    # ── Verify ──
    doc2 = Document(DOC)
    print(f'\n=== VERIFIKASI ===')
    print(f'Total paragraphs: {len(doc2.paragraphs)}')
    print(f'Total sections: {len(doc2.sections)}')

    # Count remaining BAB V
    bab5_count = 0
    for i, p in enumerate(doc2.paragraphs):
        t = p.text.strip()
        s = p.style.name if p.style else ''
        if s == 'Heading 1' and t:
            print(f'  H1 P{i}: {t}')
        if t == 'BAB V' and s == 'Heading 1':
            bab5_count += 1

    print(f'\nBAB V heading count: {bab5_count}')
    toc_remain = sum(1 for p in doc2.paragraphs
                     if p.style and p.style.name and 'toc' in p.style.name.lower())
    print(f'TOC entries remaining: {toc_remain}')
    if bab5_count == 1:
        print('✅ Struktur BAB V OK')
    else:
        print(f'⚠️  Masih ada {bab5_count} BAB V — periksa manual')


if __name__ == '__main__':
    main()
