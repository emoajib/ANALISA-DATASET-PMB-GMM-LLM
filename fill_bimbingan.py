import docx

def fill_bimbingan():
    doc = docx.Document('/Volumes/WORK/MTI UNSIBANK/TESIS/FULL TESIS/Bahan_Lampiran_2_Final.docx')
    
    bimbingan_data = [
        ("05 April 2026", "Pengajuan Outline Tesis dan Diskusi Latar Belakang Masalah (Bab 1)."),
        ("20 April 2026", "Review Kajian Literatur, Penentuan Dataset PMB, dan Kajian IndoBERT (Bab 2)."),
        ("05 Mei 2026", "Perancangan Arsitektur Hybrid Cognitive Pipeline dan Metodologi CRISP-DM (Bab 3)."),
        ("15 Mei 2026", "Revisi Proposal Tesis dan Persiapan Paparan Ujian Workshop 1."),
        ("05 Juni 2026", "Ekstraksi Fitur Teks, Reduksi Dimensi (PCA), dan Eksperimen GMM (Bab 4)."),
        ("18 Juni 2026", "Analisis Stabilitas Time Series (ARI) dan Prompting Generasi Persona LLM (Bab 4)."),
        ("01 Juli 2026", "Perumusan Kesimpulan, Implikasi Manajerial, dan Draf Artikel Jurnal (Bab 5)."),
        ("08 Juli 2026", "Finalisasi Naskah Tesis Lengkap, Video, dan Persiapan Sidang Workshop 2.")
    ]
    
    table = doc.tables[1]
    
    for i, data in enumerate(bimbingan_data):
        row = table.rows[i + 1]
        
        # Tanggal Bimbingan
        cell_date = row.cells[1]
        cell_date.text = data[0]
        # Preserve font (if possible, but usually just setting text is fine in basic tables)
        for paragraph in cell_date.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = docx.shared.Pt(12)
                
        # Materi
        cell_materi = row.cells[2]
        cell_materi.text = data[1]
        for paragraph in cell_materi.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = docx.shared.Pt(12)

    # Save to a new file to preserve the original as backup just in case
    doc.save('/Volumes/WORK/MTI UNSIBANK/TESIS/FULL TESIS/Bahan_Lampiran_2_Terisi.docx')
    print("Berhasil mengisi Lembar Bimbingan pada file Bahan_Lampiran_2_Terisi.docx")

if __name__ == '__main__':
    fill_bimbingan()
